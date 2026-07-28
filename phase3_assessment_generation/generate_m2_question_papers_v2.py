"""
M2 (Mid-Semester II) Question Paper Generator — Dynamic Systems Modeling
=======================================================================
Generates 4 sets of QP + 4 sets of MODEL ANSWERS in Word (.docx) format.

Syllabus Scope : U3 L4–L5, U4 (complete), U5 L1–L4
Marking Scheme : Q1 compulsory (4 marks) + Solve any 2 from Q2/Q3/Q4 (3 marks each) = 10 marks
Duration       : 45 minutes

Course Outcomes
  CO1 — Describe theoretical concepts in mathematical modelling
  CO2 — Analyze dynamic systems using constitutive laws
  CO3 — Synthesize dynamic systems via modelling, simulation, and control design

Revision History
  v2 — Comprehensive review pass:
        • Q1(a) Set 4: replaced out-of-scope D'Alembert with holonomic constraint (U4 L1)
        • Q2: unified to Atwood Machine across all 4 sets (same structure, different masses)
        • Q3/Q4: rebalanced difficulty across sets
        • Fixed yo-yo moment of inertia (2.25e-5, not 1.125e-5)
        • Fixed Set 4 Q4 log-decrement solvability issue
        • Added full step-by-step MODEL ANSWER generation
        • All nomenclature verified against course notebooks
"""

import os, math
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ═══════════════════════ CONFIGURATION ═══════════════════════

INSTITUTION = "Mukesh Patel School of Technology Management & Engineering"
PROGRAM     = "B.Tech Mechatronics Engineering, Semester IV"
COURSE      = "Dynamic Systems Modeling (702MH0C023)"
EXAM_TITLE  = "Mid Semester Examination – II"
EXAM_DATE   = "16-04-2026"
DURATION    = "45 minutes"
TOTAL_MARKS = "10"
OUTPUT_DIR  = os.path.join(os.path.dirname(os.path.abspath(__file__)), "question_papers_m2")

g = 9.81  # gravitational acceleration used throughout


# ═══════════════════════ QUESTION + SOLUTION BANKS ═══════════

# ─── Q1(a) Pure Theory ──────────────────────────────────────
# Bloom L1–L2 (Remember / Understand)  |  CO1
# 4 sub-questions × 0.5 marks = 2 marks

Q1A_BANKS = {
    1: {
        "questions": [
            # U3 L4
            "Define the Coriolis acceleration and state when it becomes significant in engineering analysis.",
            # U4 L2
            "What is an Atwood machine? State the underlying principle on which it operates.",
            # U5 L1
            "Define the moment of inertia of a rigid body and state its SI unit.",
            # U5 L4
            "Differentiate between undamped and damped free vibrations.",
        ],
        "solutions": [
            "The Coriolis acceleration is the term 2Ω × v_rel that arises when an object "
            "moves with velocity v_rel in a reference frame rotating at angular velocity Ω. "
            "It becomes significant when the product Ω·v_rel is comparable to other acceleration "
            "terms, e.g. in large-scale meteorology, long-range ballistics, or fast-rotating machinery.",

            "An Atwood machine consists of two masses connected by a light inextensible string "
            "passing over a frictionless pulley. It works on Newton's second law applied to "
            "a constrained two-body system: the heavier mass accelerates downward while the "
            "lighter mass accelerates upward, with a common acceleration determined by the "
            "mass difference.",

            "The moment of inertia (I) of a rigid body about a given axis is the sum (or "
            "integral) of the products of each mass element and the square of its distance "
            "from the axis: I = Σ mᵢrᵢ² (discrete) or I = ∫ r² dm (continuous). "
            "SI unit: kg·m².",

            "Undamped free vibration: the system oscillates indefinitely at its natural "
            "frequency ω_n with constant amplitude (no energy loss). "
            "Damped free vibration: a dissipative element (damper, friction) causes the "
            "amplitude to decay over time; the oscillation frequency ω_d is slightly lower "
            "than ω_n, and the system eventually comes to rest.",
        ],
    },
    2: {
        "questions": [
            # U5 L1
            "State the parallel axis theorem and give its mathematical expression.",
            # U5 L2
            "Define the term 'gear ratio' and explain its physical significance.",
            # U3 L5
            "What is a slider-crank mechanism? Name two practical applications.",
            # U5 L4
            "Define critical damping and state the mathematical condition for it.",
        ],
        "solutions": [
            "The parallel axis theorem states that the moment of inertia about any axis "
            "parallel to an axis through the center of mass is: I = I_cm + md², "
            "where I_cm = MOI about the CM axis, m = total mass, d = perpendicular distance "
            "between the two axes. This is the minimum-shift property: I is always smallest "
            "about the CM.",

            "The gear ratio (GR) is the ratio of the number of teeth (or radius) of the "
            "driven gear to the driving gear: GR = N_driven / N_driver. Physically, it "
            "determines the speed and torque transformation: output speed = input speed / GR, "
            "and output torque = input torque × GR (ignoring losses). A GR > 1 reduces speed "
            "but amplifies torque.",

            "A slider-crank mechanism converts rotary motion (crank) into reciprocating linear "
            "motion (slider/piston) via a connecting rod. The kinematic constraint is "
            "r sin θ = L sin φ, relating the crank angle θ to the connecting rod angle φ. "
            "Applications: (1) Internal combustion engines (piston–crankshaft), "
            "(2) Reciprocating compressors / pumps.",

            "Critical damping is the boundary between oscillatory (underdamped) and "
            "non-oscillatory (overdamped) response. Mathematical condition: damping ratio "
            "ζ = 1, i.e. c = c_cr = 2√(km), where c = damping coefficient, k = stiffness, "
            "m = mass. The system returns to equilibrium in the shortest time without "
            "overshooting.",
        ],
    },
    3: {
        "questions": [
            # U5 L1
            "Explain the concept of rolling without slipping and state the velocity constraint.",
            # U4 L3
            "What is a trapezoidal velocity profile in CNC motion planning? Sketch it qualitatively.",
            # U5 L4
            "Define logarithmic decrement in vibration analysis and state its physical meaning.",
            # U4 L2
            "State the relationship between mechanical advantage and the number of supporting rope segments in a block-and-tackle system.",
        ],
        "solutions": [
            "Rolling without slipping means the contact point of the rolling body has zero "
            "velocity relative to the surface. The velocity constraint is v_cm = ωR, where "
            "v_cm = linear velocity of the center of mass, ω = angular velocity, R = radius. "
            "Equivalently, a_cm = αR for accelerations.",

            "A trapezoidal velocity profile has three phases: (1) constant acceleration "
            "(velocity ramps up linearly), (2) constant velocity (cruise phase), and "
            "(3) constant deceleration (velocity ramps down). The velocity vs. time graph "
            "forms a trapezoid shape. It is used in CNC and servo drives because it limits "
            "the maximum acceleration (jerk control) while achieving a desired feedrate.",

            "The logarithmic decrement δ is the natural log of the ratio of two successive "
            "peak amplitudes in a damped oscillation: δ = ln(x_n / x_{n+1}). It is related "
            "to the damping ratio by δ = 2πζ / √(1 − ζ²). Physically, it quantifies the "
            "rate of amplitude decay per cycle — a larger δ means faster energy dissipation.",

            "In an ideal block-and-tackle system, the mechanical advantage (MA) equals the "
            "number of rope segments supporting the load: MA = n. This means the required "
            "pulling force is F = W/n, but the rope must be pulled n times the load "
            "displacement (energy conservation). In a real system, friction reduces the "
            "effective MA.",
        ],
    },
    4: {
        "questions": [
            # U4 L1
            "Define a holonomic constraint and give one example from a mechanical system.",
            # U5 L4
            "What is the physical significance of the natural frequency of a vibrating system?",
            # U3 L4
            "Define the centripetal and Coriolis components of acceleration in a rotating reference frame.",
            # U5 L1
            "What is a physical (compound) pendulum? How does it differ from a simple pendulum?",
        ],
        "solutions": [
            "A holonomic constraint is one that can be expressed as an equation relating "
            "only the coordinates (and possibly time), not velocities: f(q₁, q₂, ..., t) = 0. "
            "Example: A bead constrained to move on a circular wire of radius R satisfies "
            "x² + y² = R² (holonomic). This reduces the degrees of freedom by one.",

            "The natural frequency ω_n = √(k/m) is the frequency at which the undamped "
            "system oscillates freely when displaced from equilibrium. It depends only on "
            "system properties (stiffness and mass), not on initial conditions. At "
            "resonance, when an external forcing frequency ω ≈ ω_n, the response amplitude "
            "becomes very large — this is critical in vibration engineering design.",

            "In a frame rotating at angular velocity Ω: "
            "• Centripetal acceleration = Ω × (Ω × r), magnitude Ω²r, directed radially "
            "inward toward the rotation axis. "
            "• Coriolis acceleration = 2Ω × v_rel, magnitude 2Ωv_rel, perpendicular to both "
            "the rotation axis and the relative velocity. It arises only when the object "
            "moves within the rotating frame.",

            "A physical (compound) pendulum is an extended rigid body (e.g. a rod, disk) "
            "that oscillates about a fixed pivot point. Unlike a simple pendulum (point mass "
            "on a massless string, T = 2π√(L/g)), the physical pendulum's period depends on "
            "the moment of inertia about the pivot and the distance from pivot to CM: "
            "T = 2π√(I_pivot / (mgd)).",
        ],
    },
}


# ─── Q1(b) Mixed Theory + Short Numerical ──────────────────
# Bloom L2–L3 (Understand / Apply)  |  CO1, CO2
# 4 sub-questions × 0.5 marks = 2 marks

Q1B_BANKS = {
    1: {
        "questions": [
            # U5 L1 — short numerical
            "A solid disk of mass 5 kg and radius 0.2 m rolls without slipping. Calculate its moment of inertia about the contact point using the parallel axis theorem.",
            # U5 L1 — conceptual
            "Explain why a solid sphere reaches the bottom of an incline before a hollow cylinder of the same mass and radius.",
            # U4 L2 — conceptual
            "In an Atwood machine with masses 5 kg and 3 kg, state which mass accelerates downward and why.",
            # U4 L2 — conceptual
            "State two assumptions commonly made when analyzing an ideal pulley system.",
        ],
        "solutions": [
            "I_cm = (1/2)mR² = 0.5 × 5 × 0.2² = 0.10 kg·m². "
            "By parallel axis theorem: I_contact = I_cm + mR² = 0.10 + 5 × 0.04 = 0.10 + 0.20 "
            "= 0.30 kg·m².",

            "The rolling acceleration is a = g sinθ / (1 + β), where β = I/(mR²). "
            "Solid sphere: β = 2/5 = 0.4, hollow cylinder: β = 1.0. Since the sphere has a "
            "smaller β, it has a larger acceleration and reaches the bottom first. The "
            "hollow cylinder puts more energy into rotation.",

            "The 5 kg mass accelerates downward because it is heavier. In an Atwood machine, "
            "the net downward force is (m₁ − m₂)g = (5 − 3) × 9.81 = 19.62 N, and this "
            "causes the heavier mass to descend while the lighter mass rises.",

            "(1) The pulley is massless and frictionless (no rotational inertia, no axle "
            "friction). (2) The string/rope is massless, inextensible, and has uniform "
            "tension throughout its length.",
        ],
    },
    2: {
        "questions": [
            # U5 L2 — short numerical
            "A gear train has an input speed of 100 RPM and a gear ratio of 15 : 1 (speed reduction). Calculate the output shaft speed.",
            # U5 L4 — conceptual
            "Explain the physical meaning of the damping ratio ζ and what happens when ζ = 1.",
            # U4 L3 — short numerical
            "A CNC axis must traverse 200 mm with v_max = 100 mm/s and a_max = 500 mm/s². Estimate the total traverse time using a trapezoidal profile.",
            # U4 L1 — conceptual
            "State the condition for a block to remain stationary on an inclined plane with friction.",
        ],
        "solutions": [
            "Output speed = Input speed / GR = 100 / 15 = 6.67 RPM.",

            "The damping ratio ζ = c / c_cr is the ratio of actual damping to critical "
            "damping. It determines the system's response type: ζ < 1 → underdamped "
            "(oscillates with decaying amplitude), ζ = 1 → critically damped (fastest return "
            "to equilibrium without overshoot), ζ > 1 → overdamped (slow, non-oscillatory "
            "return).",

            "t_accel = v_max / a_max = 100/500 = 0.2 s. "
            "d_accel = v_max² / (2 a_max) = 10000/1000 = 10 mm. "
            "d_cruise = 200 − 2 × 10 = 180 mm. "
            "t_cruise = 180/100 = 1.8 s. "
            "Total time = 0.2 + 1.8 + 0.2 = 2.2 s.",

            "The block remains stationary when the component of gravity along the incline "
            "does not exceed the maximum static friction: mg sinθ ≤ μ_s mg cosθ, i.e. "
            "tanθ ≤ μ_s, where μ_s is the coefficient of static friction.",
        ],
    },
    3: {
        "questions": [
            # U3 L4 — short numerical
            "A platform rotates at Ω = 3 rad/s. A particle moves radially outward at v_rel = 2 m/s. Calculate the magnitude of the Coriolis acceleration.",
            # U5 L4 — conceptual
            "Write the equation of motion for an undamped spring-mass system and identify each term.",
            # U4 L2 — conceptual
            "Define mechanical advantage of a pulley system. State its ideal value for a single movable pulley.",
            # U3 L4 — conceptual
            "Explain why fictitious forces (pseudo-forces) appear in a rotating (non-inertial) reference frame.",
        ],
        "solutions": [
            "a_Coriolis = 2Ω v_rel = 2 × 3 × 2 = 12 m/s². "
            "Direction: perpendicular to v_rel, opposing the sense of rotation.",

            "Equation: mẍ + kx = 0, where m = mass (kg), ẍ = d²x/dt² = acceleration (m/s²), "
            "k = spring stiffness (N/m), x = displacement from equilibrium (m). "
            "Term mẍ = inertia force; term kx = restoring spring force.",

            "Mechanical advantage (MA) = Load / Effort = W / F. For an ideal single movable "
            "pulley, MA = 2, since two rope segments support the load. This means the effort "
            "is halved, but the rope must be pulled twice the load displacement.",

            "Newton's laws are valid only in inertial (non-accelerating) frames. When we "
            "write F = ma in a rotating frame, extra terms appear because the frame itself "
            "is accelerating. These terms (centrifugal force = −mΩ×(Ω×r), Coriolis force = "
            "−2mΩ×v_rel) are called fictitious or pseudo-forces because they have no "
            "physical agent — they arise purely from the non-inertial reference frame.",
        ],
    },
    4: {
        "questions": [
            # U3 L5 — short numerical
            "A slider-crank has crank radius r = 50 mm and connecting rod length L = 150 mm. State the ratio L/r and explain its significance in mechanism design.",
            # U5 L4 — short numerical
            "For a spring-mass system with m = 2 kg and k = 200 N/m, calculate the natural frequency ω_n in rad/s.",
            # U4 L1 — conceptual
            "Give an example of a constraint force and explain why it does no work in the direction of the constraint.",
            # U5 L2 — conceptual
            "State two applications of gear trains in mechatronics systems.",
        ],
        "solutions": [
            "L/r = 150/50 = 3. This ratio determines how closely the piston motion "
            "approximates simple harmonic motion. Higher L/r ratios (typically 3–5) reduce "
            "higher-order harmonics in the piston displacement, making the motion smoother "
            "and reducing vibrations.",

            "ω_n = √(k/m) = √(200/2) = √100 = 10 rad/s.",

            "Example: The normal force N on a block resting on a surface. The block is "
            "constrained not to penetrate the surface. The normal force acts perpendicular "
            "to the surface, and since the block cannot move in that direction (the direction "
            "of the constraint), the work done by N is zero: W = N · displacement_⊥ = 0.",

            "(1) Robotic arm joints — gear reducers provide high torque at low speed for "
            "precise positioning. (2) CNC spindle drives — gear trains match motor speed to "
            "required cutting speed. Other examples: conveyor belt drives, stepper motor "
            "micro-stepping gearboxes.",
        ],
    },
}


# ─── Q2: COMMON TOPIC — Atwood Machine (U4 L2) ────────────
# Bloom L3 (Apply)  |  CO2
# 6 sub-questions × 0.5 marks = 3 marks
# Same structure, different mass values for each set

Q2_BANKS = {
    1: {
        "stem": (
            "Two masses m₁ = 5 kg and m₂ = 3 kg are connected by a light, "
            "inextensible string passing over a massless, frictionless pulley. "
            "The system is released from rest. Take g = 9.81 m/s²."
        ),
        "bloom": "BL 3 (Apply)", "co": "CO2",
        "m1": 5, "m2": 3, "t": 2,
        "subs": [
            "Draw the free-body diagram for each mass and write the equation of motion for m₁.",
            "Write the equation of motion for m₂.",
            "Find the acceleration of the system.",
            "Calculate the tension in the string.",
            "Find the velocity of the masses after 2 seconds.",
            "Calculate the distance traveled by each mass in 2 seconds.",
        ],
    },
    2: {
        "stem": (
            "Two masses m₁ = 8 kg and m₂ = 5 kg are connected by a light, "
            "inextensible string passing over a massless, frictionless pulley. "
            "The system is released from rest. Take g = 9.81 m/s²."
        ),
        "bloom": "BL 3 (Apply)", "co": "CO2",
        "m1": 8, "m2": 5, "t": 3,
        "subs": [
            "Draw the free-body diagram for each mass and write the equation of motion for m₁.",
            "Write the equation of motion for m₂.",
            "Find the acceleration of the system.",
            "Calculate the tension in the string.",
            "Find the velocity of the masses after 3 seconds.",
            "Calculate the distance traveled by each mass in 3 seconds.",
        ],
    },
    3: {
        "stem": (
            "Two masses m₁ = 6 kg and m₂ = 4 kg are connected by a light, "
            "inextensible string passing over a massless, frictionless pulley. "
            "The system is released from rest. Take g = 9.81 m/s²."
        ),
        "bloom": "BL 3 (Apply)", "co": "CO2",
        "m1": 6, "m2": 4, "t": 2,
        "subs": [
            "Draw the free-body diagram for each mass and write the equation of motion for m₁.",
            "Write the equation of motion for m₂.",
            "Find the acceleration of the system.",
            "Calculate the tension in the string.",
            "Find the velocity of the masses after 2 seconds.",
            "Calculate the distance traveled by each mass in 2 seconds.",
        ],
    },
    4: {
        "stem": (
            "Two masses m₁ = 10 kg and m₂ = 6 kg are connected by a light, "
            "inextensible string passing over a massless, frictionless pulley. "
            "The system is released from rest. Take g = 9.81 m/s²."
        ),
        "bloom": "BL 3 (Apply)", "co": "CO2",
        "m1": 10, "m2": 6, "t": 3,
        "subs": [
            "Draw the free-body diagram for each mass and write the equation of motion for m₁.",
            "Write the equation of motion for m₂.",
            "Find the acceleration of the system.",
            "Calculate the tension in the string.",
            "Find the velocity of the masses after 3 seconds.",
            "Calculate the distance traveled by each mass in 3 seconds.",
        ],
    },
}


def solve_q2(set_num):
    """Compute Atwood machine solutions for given set."""
    b = Q2_BANKS[set_num]
    m1, m2, t = b["m1"], b["m2"], b["t"]
    a = (m1 - m2) * g / (m1 + m2)
    T = m2 * (g + a)
    v = a * t
    s = 0.5 * a * t**2
    return [
        f"FBD for m₁: Weight m₁g = {m1}×9.81 = {m1*g:.2f} N downward, Tension T upward.\n"
        f"Equation of motion (taking downward positive for m₁): m₁g − T = m₁a\n"
        f"→ {m1*g:.2f} − T = {m1}a",

        f"FBD for m₂: Weight m₂g = {m2}×9.81 = {m2*g:.2f} N downward, Tension T upward.\n"
        f"Equation of motion (taking upward positive for m₂): T − m₂g = m₂a\n"
        f"→ T − {m2*g:.2f} = {m2}a",

        f"Adding both equations: m₁g − m₂g = (m₁ + m₂)a\n"
        f"a = (m₁ − m₂)g / (m₁ + m₂) = ({m1} − {m2}) × 9.81 / ({m1} + {m2})\n"
        f"a = {(m1-m2)*g:.2f} / {m1+m2} = {a:.3f} m/s²",

        f"From equation for m₂: T = m₂(g + a) = {m2} × (9.81 + {a:.3f})\n"
        f"T = {m2} × {g + a:.3f} = {T:.2f} N\n"
        f"Check via m₁: T = m₁(g − a) = {m1} × {g - a:.3f} = {m1*(g-a):.2f} N ✓",

        f"v = u + at = 0 + {a:.3f} × {t} = {v:.3f} m/s",

        f"s = ½at² = 0.5 × {a:.3f} × {t}² = 0.5 × {a:.3f} × {t**2} = {s:.3f} m",
    ]


# ─── Q3: Medium-Hard Numerical (varying topics) ─────────────
# Bloom L3–L4 (Apply / Analyze)  |  CO2
# 6 sub-questions × 0.5 marks = 3 marks

Q3_BANKS = {
    # Set 1 — Coriolis on Rotating Platform (U3 L4)
    1: {
        "stem": (
            "A platform rotates at constant angular velocity Ω = 3 rad/s about a "
            "vertical axis. An object at r = 2 m from the center moves radially outward "
            "at v_rel = 1.5 m/s with no radial acceleration (r̈ = 0). "
            "The angular acceleration of the platform is zero (α = 0)."
        ),
        "bloom": "BL 3–4 (Apply / Analyze)", "co": "CO2",
        "subs": [
            "Calculate the centripetal acceleration magnitude.",
            "Calculate the Coriolis acceleration magnitude.",
            "State the direction of the Coriolis acceleration relative to the radial velocity.",
            "Find the total acceleration magnitude in the rotating frame.",
            "Determine the tangential velocity of the object as observed from the fixed (inertial) frame.",
            "If the object continues at the same radial speed, find the centripetal acceleration when r = 4 m.",
        ],
    },
    # Set 2 — Rough Incline with Friction (U4 L1 / P1)
    2: {
        "stem": (
            "A block of mass m = 5 kg is placed on a rough inclined plane at angle "
            "θ = 30° to the horizontal. The coefficient of kinetic friction is "
            "μ_k = 0.2. The block is released from rest. Take g = 9.81 m/s²."
        ),
        "bloom": "BL 3–4 (Apply / Analyze)", "co": "CO2",
        "subs": [
            "Draw the free-body diagram showing all forces on the block.",
            "Calculate the normal reaction force N.",
            "Calculate the kinetic friction force f.",
            "Determine the net force along the incline and find the acceleration.",
            "Find the velocity of the block after sliding 5 m from rest.",
            "Calculate the time taken to slide 5 m from rest.",
        ],
    },
    # Set 3 — Slider-Crank Kinematics (U3 L5)
    3: {
        "stem": (
            "A slider-crank mechanism has crank radius r = 50 mm, connecting rod "
            "length L = 150 mm, and operates at 3000 RPM. Analyze the mechanism at "
            "crank angle θ = 30°."
        ),
        "bloom": "BL 3–4 (Apply / Analyze)", "co": "CO2",
        "subs": [
            "Convert the crank speed from RPM to rad/s.",
            "Calculate the connecting rod angle φ using the kinematic relation r sin θ = L sin φ.",
            "Determine the piston displacement from top dead center (TDC).",
            "Find the piston velocity.",
            "Identify the crank angles at which the piston velocity is zero.",
            "State the type of motion executed by the piston and name one engineering application.",
        ],
    },
    # Set 4 — Rolling Sphere on Incline (U5 L1)
    4: {
        "stem": (
            "A solid sphere of mass m = 5 kg and radius R = 0.2 m rolls without "
            "slipping down an incline of angle θ = 30°. Take g = 9.81 m/s²."
        ),
        "bloom": "BL 3–4 (Apply / Analyze)", "co": "CO2",
        "subs": [
            "Calculate the moment of inertia of the solid sphere about its center.",
            "State the no-slip constraint relating linear and angular quantities.",
            "Write the expression for total kinetic energy (translational + rotational) of the rolling sphere.",
            "Find the linear acceleration of the sphere down the incline.",
            "Compare this acceleration to that of a frictionless sliding block on the same incline.",
            "Find the velocity of the sphere after rolling 3 m from rest.",
        ],
    },
}


def solve_q3(set_num):
    """Compute Q3 solutions."""
    if set_num == 1:
        # Coriolis
        omega, r, v_rel = 3.0, 2.0, 1.5
        a_cent = omega**2 * r
        a_cor = 2 * omega * v_rel
        a_total = math.sqrt(a_cent**2 + a_cor**2)
        v_tang = omega * r
        a_cent4 = omega**2 * 4.0
        return [
            f"a_centripetal = Ω²r = {omega}² × {r} = {a_cent:.1f} m/s²",
            f"a_Coriolis = 2Ωv_rel = 2 × {omega} × {v_rel} = {a_cor:.1f} m/s²",
            "The Coriolis acceleration is perpendicular to v_rel (the radial direction), "
            "directed tangentially in the sense opposite to the platform rotation "
            "(by the right-hand rule for 2Ω × v_rel).",
            f"|a_total| = √(a_cent² + a_Cor²) = √({a_cent:.1f}² + {a_cor:.1f}²) "
            f"= √({a_cent**2:.0f} + {a_cor**2:.0f}) = √{a_cent**2+a_cor**2:.0f} "
            f"= {a_total:.2f} m/s²",
            f"v_tangential = Ωr = {omega} × {r} = {v_tang:.1f} m/s",
            f"At r = 4 m: a_centripetal = Ω²r = {omega}² × 4 = {a_cent4:.1f} m/s²",
        ]
    elif set_num == 2:
        # Rough incline
        m, theta_deg = 5.0, 30.0
        theta = math.radians(theta_deg)
        mu_k = 0.2
        N = m * g * math.cos(theta)
        f_k = mu_k * N
        F_net = m * g * math.sin(theta) - f_k
        a = F_net / m
        s = 5.0
        v = math.sqrt(2 * a * s)
        t = math.sqrt(2 * s / a)
        return [
            f"Forces: (1) Weight W = mg = {m*g:.2f} N downward, "
            f"(2) Normal force N perpendicular to incline (upward from surface), "
            f"(3) Kinetic friction f = μ_k N opposing motion (up the incline).",
            f"N = mg cos θ = {m} × 9.81 × cos 30° = {m*g:.2f} × {math.cos(theta):.4f} "
            f"= {N:.2f} N",
            f"f = μ_k × N = {mu_k} × {N:.2f} = {f_k:.2f} N",
            f"F_net = mg sinθ − f = {m*g*math.sin(theta):.2f} − {f_k:.2f} = {F_net:.2f} N\n"
            f"a = F_net / m = {F_net:.2f} / {m} = {a:.2f} m/s²",
            f"v = √(2as) = √(2 × {a:.2f} × {s}) = √{2*a*s:.2f} = {v:.2f} m/s",
            f"t = √(2s/a) = √(2 × {s} / {a:.2f}) = √{2*s/a:.2f} = {t:.2f} s",
        ]
    elif set_num == 3:
        # Slider-crank
        r_mm, L_mm = 50, 150
        N_rpm = 3000
        theta_deg = 30.0
        omega_rad = 2 * math.pi * N_rpm / 60
        theta = math.radians(theta_deg)
        sin_phi = (r_mm / L_mm) * math.sin(theta)
        phi = math.asin(sin_phi)
        phi_deg = math.degrees(phi)
        cos_phi = math.cos(phi)
        x_pos = r_mm * math.cos(theta) + L_mm * cos_phi
        x_TDC = r_mm + L_mm
        disp = x_TDC - x_pos
        # Velocity
        v_piston = -r_mm * omega_rad * (math.sin(theta) +
                    (r_mm * math.sin(2*theta)) / (2 * L_mm))
        v_piston_m = v_piston / 1000
        return [
            f"ω = 2π × N / 60 = 2π × 3000 / 60 = {omega_rad:.2f} rad/s",
            f"sin φ = (r/L) sin θ = ({r_mm}/{L_mm}) × sin 30° = {1/3:.4f} × 0.5 "
            f"= {sin_phi:.4f}\n"
            f"φ = arcsin({sin_phi:.4f}) = {phi_deg:.2f}°",
            f"Piston position: x = r cos θ + L cos φ = {r_mm} cos 30° + {L_mm} cos {phi_deg:.2f}°\n"
            f"= {r_mm*math.cos(theta):.2f} + {L_mm*cos_phi:.2f} = {x_pos:.2f} mm\n"
            f"TDC position: x_TDC = r + L = {x_TDC} mm\n"
            f"Displacement from TDC = {x_TDC} − {x_pos:.2f} = {disp:.2f} mm",
            f"v_piston ≈ −rω(sin θ + (r sin 2θ)/(2L))\n"
            f"= −{r_mm} × {omega_rad:.2f} × (sin 30° + ({r_mm} × sin 60°)/(2 × {L_mm}))\n"
            f"= −{r_mm} × {omega_rad:.2f} × (0.5 + {r_mm*math.sin(2*theta)/(2*L_mm):.4f})\n"
            f"= {v_piston:.0f} mm/s = {v_piston_m:.2f} m/s",
            "v_piston = 0 at Top Dead Center (θ = 0°) and Bottom Dead Center (θ = 180°), "
            "where the crank and connecting rod are collinear.",
            "The piston executes reciprocating (oscillatory linear) motion. "
            "Application: Internal combustion engines.",
        ]
    else:  # set_num == 4
        # Rolling sphere
        m, R = 5.0, 0.2
        theta_deg = 30.0
        theta = math.radians(theta_deg)
        I_cm = (2/5) * m * R**2
        beta = 2/5
        a_roll = g * math.sin(theta) / (1 + beta)
        a_slide = g * math.sin(theta)
        s = 3.0
        v = math.sqrt(2 * a_roll * s)
        return [
            f"I_cm = (2/5)mR² = (2/5) × {m} × {R}² = (2/5) × {m} × {R**2} "
            f"= {I_cm:.3f} kg·m²",
            "No-slip constraint: v_cm = ωR and a_cm = αR, where ω = angular velocity, "
            "α = angular acceleration, R = radius.",
            f"KE_total = ½mv² + ½Iω² = ½mv² + ½(2/5)mR²(v/R)²\n"
            f"= ½mv² + (1/5)mv² = (7/10)mv²",
            f"Newton's second law:\n"
            f"  Translation: mg sinθ − f = ma\n"
            f"  Rotation: fR = Iα = (2/5)mR²(a/R) → f = (2/5)ma\n"
            f"Substituting: mg sinθ − (2/5)ma = ma → mg sinθ = (7/5)ma\n"
            f"a = (5/7)g sinθ = (5/7) × 9.81 × sin 30° = (5/7) × 4.905 = {a_roll:.3f} m/s²",
            f"Frictionless sliding: a_slide = g sinθ = 9.81 × 0.5 = {a_slide:.3f} m/s².\n"
            f"Rolling acceleration ({a_roll:.3f}) < sliding acceleration ({a_slide:.3f}) "
            f"because part of the gravitational PE goes into rotational KE.",
            f"v = √(2as) = √(2 × {a_roll:.3f} × {s}) = √{2*a_roll*s:.3f} = {v:.3f} m/s",
        ]


# ─── Q4: Hard Numerical (varying topics) ────────────────────
# Bloom L4 (Analyze)  |  CO2, CO3
# 6 sub-questions × 0.5 marks = 3 marks

Q4_BANKS = {
    # Set 1 — Physical Pendulum (U5 L1)
    1: {
        "stem": (
            "A uniform rod of mass m = 2 kg and length L = 1 m is pivoted at one "
            "end and oscillates as a physical pendulum under gravity. "
            "Take g = 9.81 m/s²."
        ),
        "bloom": "BL 4–5 (Analyze / Evaluate)", "co": "CO2, CO3",
        "subs": [
            "Calculate the moment of inertia of the rod about the pivot using the parallel axis theorem.",
            "Find the distance of the center of mass from the pivot.",
            "Write the equation of motion for small-angle oscillations.",
            "Determine the natural frequency ω_n.",
            "Calculate the time period of oscillation.",
            "Compare this period with that of a simple pendulum of the same length (L = 1 m).",
        ],
    },
    # Set 2 — Undamped + Damped Vibrations (U5 L4)
    2: {
        "stem": (
            "A mass m = 2 kg is attached to a spring of stiffness k = 200 N/m. "
            "The system is displaced x₀ = 0.1 m from equilibrium and released "
            "from rest (v₀ = 0)."
        ),
        "bloom": "BL 4–5 (Analyze / Evaluate)", "co": "CO2, CO3",
        "subs": [
            "Write the equation of motion for the undamped system.",
            "Calculate the natural frequency ω_n and the time period T.",
            "Write the displacement as a function of time x(t) for the given initial conditions.",
            "A viscous damper with c = 4 N·s/m is now added. Calculate the critical damping coefficient c_cr and the damping ratio ζ.",
            "State whether the damped system is underdamped, critically damped, or overdamped. Calculate the damped natural frequency ω_d.",
            "Calculate the logarithmic decrement δ for the damped system.",
        ],
    },
    # Set 3 — Yo-Yo Unwinding (U5 L2)
    3: {
        "stem": (
            "A yo-yo of mass m = 0.05 kg has outer radius R = 0.03 m, axle radius "
            "r = 0.01 m, and moment of inertia I = 2.25 × 10⁻⁵ kg·m² (modeled as "
            "a solid disk of radius R). The string unwinds from the axle. "
            "Take g = 9.81 m/s²."
        ),
        "bloom": "BL 4–5 (Analyze / Evaluate)", "co": "CO2, CO3",
        "subs": [
            "Write the translational equation of motion for the yo-yo (vertical forces).",
            "Write the rotational equation of motion about the center of mass.",
            "State the constraint relating linear acceleration a and angular acceleration α.",
            "Solve for the linear acceleration of the yo-yo.",
            "Calculate the tension in the string.",
            "Express the acceleration as a fraction of g and explain why the yo-yo descends slower than free fall.",
        ],
    },
    # Set 4 — 2-DOF Spring-Mass Eigenvalue Problem (U5 L3)
    4: {
        "stem": (
            "A two-degree-of-freedom spring-mass system has m₁ = m₂ = 1 kg, "
            "k₁ = k₃ = 100 N/m (grounding springs), and coupling spring "
            "k₂ = 50 N/m connecting the two masses."
        ),
        "bloom": "BL 4–5 (Analyze / Evaluate)", "co": "CO2, CO3",
        "subs": [
            "Write the equations of motion for mass m₁ and mass m₂.",
            "Express the equations in matrix form [M]{ẍ} + [K]{x} = {0}.",
            "Write the mass matrix [M] and the stiffness matrix [K].",
            "Set up the characteristic equation det([K] − ω²[M]) = 0 and expand it.",
            "Solve for the two natural frequencies ω₁ and ω₂.",
            "Describe the physical meaning of the two mode shapes.",
        ],
    },
}


def solve_q4(set_num):
    """Compute Q4 solutions."""
    if set_num == 1:
        # Physical pendulum
        m, L = 2.0, 1.0
        I_cm = m * L**2 / 12
        I_piv = m * L**2 / 3
        d = L / 2
        omega_n = math.sqrt(m * g * d / I_piv)
        T_phys = 2 * math.pi / omega_n
        T_simp = 2 * math.pi * math.sqrt(L / g)
        return [
            f"I_cm = (1/12)mL² = (1/12) × {m} × {L}² = {I_cm:.4f} kg·m²\n"
            f"I_pivot = I_cm + m(L/2)² = {I_cm:.4f} + {m} × {(L/2)**2} "
            f"= {I_cm:.4f} + {m*(L/2)**2:.4f} = {I_piv:.4f} kg·m²\n"
            f"(Or directly: I_pivot = (1/3)mL² = {I_piv:.4f} kg·m²)",
            f"d = L/2 = {L}/2 = {d} m",
            f"Restoring torque: τ = −mgd sinθ ≈ −mgdθ (small angle)\n"
            f"Equation: I_pivot θ̈ + mgd θ = 0  →  θ̈ + (mgd/I_pivot) θ = 0",
            f"ω_n = √(mgd / I_pivot) = √({m} × {g} × {d} / {I_piv:.4f})\n"
            f"= √({m*g*d:.3f} / {I_piv:.4f}) = √{m*g*d/I_piv:.3f} "
            f"= {omega_n:.3f} rad/s",
            f"T = 2π / ω_n = 2π / {omega_n:.3f} = {T_phys:.3f} s",
            f"Simple pendulum (same length L = {L} m): T_simple = 2π√(L/g) "
            f"= 2π√({L}/{g}) = {T_simp:.3f} s.\n"
            f"Physical pendulum period ({T_phys:.3f} s) < Simple pendulum period "
            f"({T_simp:.3f} s) because the physical pendulum's mass is distributed "
            f"closer to the pivot, reducing the effective pendulum length.",
        ]
    elif set_num == 2:
        # Damped vibrations
        m, k, c = 2.0, 200.0, 4.0
        x0 = 0.1
        omega_n = math.sqrt(k / m)
        T_n = 2 * math.pi / omega_n
        c_cr = 2 * math.sqrt(k * m)
        zeta = c / c_cr
        omega_d = omega_n * math.sqrt(1 - zeta**2)
        delta = 2 * math.pi * zeta / math.sqrt(1 - zeta**2)
        return [
            f"Equation of motion: mẍ + kx = 0  →  {m}ẍ + {k}x = 0",
            f"ω_n = √(k/m) = √({k}/{m}) = √{k/m:.0f} = {omega_n:.2f} rad/s\n"
            f"T = 2π/ω_n = 2π/{omega_n:.2f} = {T_n:.4f} s ≈ {T_n:.3f} s",
            f"With x(0) = {x0} m, ẋ(0) = 0:\n"
            f"x(t) = x₀ cos(ω_n t) = {x0} cos({omega_n:.0f}t) m\n"
            f"(Amplitude {x0} m, purely cosinusoidal, no decay)",
            f"c_cr = 2√(km) = 2√({k} × {m}) = 2√{k*m:.0f} = 2 × {math.sqrt(k*m):.1f} "
            f"= {c_cr:.1f} N·s/m\n"
            f"ζ = c / c_cr = {c} / {c_cr:.1f} = {zeta:.2f}",
            f"Since ζ = {zeta:.2f} < 1 → Underdamped (oscillatory with decay).\n"
            f"ω_d = ω_n√(1 − ζ²) = {omega_n:.2f}√(1 − {zeta}²) "
            f"= {omega_n:.2f} × {math.sqrt(1-zeta**2):.4f} = {omega_d:.2f} rad/s",
            f"δ = 2πζ / √(1 − ζ²) = 2π × {zeta} / √(1 − {zeta}²)\n"
            f"= {2*math.pi*zeta:.4f} / {math.sqrt(1-zeta**2):.4f} = {delta:.4f}\n"
            f"This means each successive peak is e^(−δ) = e^(−{delta:.4f}) "
            f"= {math.exp(-delta):.4f} times the previous peak.",
        ]
    elif set_num == 3:
        # Yo-yo
        m_yo = 0.05
        R_yo, r_yo = 0.03, 0.01
        I_yo = 2.25e-5  # = (1/2)mR²
        beta = I_yo / (m_yo * r_yo**2)
        a_yo = g / (1 + beta)
        T_str = m_yo * (g - a_yo)
        ratio = a_yo / g
        return [
            f"Translation (positive downward): mg − T = ma\n"
            f"→ {m_yo} × {g} − T = {m_yo}a  →  {m_yo*g:.4f} − T = {m_yo}a",
            f"Rotation about CM: Tr = Iα\n"
            f"→ T × {r_yo} = {I_yo} × α",
            f"String unwinds from the axle of radius r, so: a = αr\n"
            f"→ α = a/r = a/{r_yo} = {1/r_yo:.0f}a",
            f"Substituting α = a/r into the rotational equation:\n"
            f"T × {r_yo} = {I_yo} × (a/{r_yo})  →  T = {I_yo}/{r_yo**2} × a = {I_yo/r_yo**2}a\n"
            f"Substituting into the translational equation:\n"
            f"{m_yo*g:.4f} − {I_yo/r_yo**2}a = {m_yo}a\n"
            f"{m_yo*g:.4f} = ({m_yo} + {I_yo/r_yo**2})a = {m_yo + I_yo/r_yo**2}a\n"
            f"a = {m_yo*g:.4f} / {m_yo + I_yo/r_yo**2} = {a_yo:.3f} m/s²",
            f"T = m(g − a) = {m_yo}({g} − {a_yo:.3f}) = {m_yo} × {g - a_yo:.3f} "
            f"= {T_str:.4f} N ≈ {T_str:.3f} N",
            f"a/g = {a_yo:.3f} / {g} = {ratio:.3f} ≈ {ratio*100:.1f}% of free fall.\n"
            f"The yo-yo descends slower than free fall because its rotational inertia "
            f"(I = {I_yo} kg·m²) resists angular acceleration. Part of the gravitational "
            f"potential energy is converted to rotational kinetic energy, leaving less "
            f"for translational kinetic energy.",
        ]
    else:  # set_num == 4
        # 2-DOF
        k1, k2, k3 = 100, 50, 100
        # [K] = [[k1+k2, -k2], [-k2, k2+k3]] = [[150, -50], [-50, 150]]
        K11, K12, K22 = k1+k2, -k2, k2+k3
        # det([K] - w^2 [I]) = (150-w^2)^2 - 2500 = 0
        w1_sq = K11 - abs(K12)  # 100
        w2_sq = K11 + abs(K12)  # 200
        w1 = math.sqrt(w1_sq)
        w2 = math.sqrt(w2_sq)
        return [
            f"For m₁: m₁ẍ₁ + k₁x₁ + k₂(x₁ − x₂) = 0  →  ẍ₁ + (k₁+k₂)x₁ − k₂x₂ = 0\n"
            f"  → ẍ₁ + {K11}x₁ − {abs(K12)}x₂ = 0\n"
            f"For m₂: m₂ẍ₂ + k₃x₂ + k₂(x₂ − x₁) = 0  →  ẍ₂ − k₂x₁ + (k₂+k₃)x₂ = 0\n"
            f"  → ẍ₂ − {abs(K12)}x₁ + {K22}x₂ = 0",

            f"[M]{{ẍ}} + [K]{{x}} = {{0}} where {{x}} = [x₁, x₂]ᵀ",

            f"[M] = [[1, 0], [0, 1]]  (since m₁ = m₂ = 1 kg)\n"
            f"[K] = [[{K11}, {K12}], [{K12}, {K22}]]  (N/m)",

            f"det([K] − ω²[M]) = det([[{K11}−ω², {K12}], [{K12}, {K22}−ω²]]) = 0\n"
            f"({K11} − ω²)({K22} − ω²) − ({K12})² = 0\n"
            f"({K11} − ω²)² − {K12**2} = 0   (since K11 = K22 = {K11})",

            f"({K11} − ω²)² = {K12**2}\n"
            f"{K11} − ω² = ±{abs(K12)}\n"
            f"ω² = {K11} − {abs(K12)} = {w1_sq}  or  ω² = {K11} + {abs(K12)} = {w2_sq}\n"
            f"ω₁ = √{w1_sq} = {w1:.2f} rad/s\n"
            f"ω₂ = √{w2_sq} = {w2:.2f} rad/s",

            f"Mode 1 (ω₁ = {w1:.2f} rad/s): Substituting ω₁² = {w1_sq} into "
            f"({K11}−{w1_sq})X₁ + ({K12})X₂ = 0 → {abs(K12)}X₁ = {abs(K12)}X₂ → X₁ = X₂. "
            f"Both masses move in the SAME direction (in-phase). The coupling spring is "
            f"not stretched.\n"
            f"Mode 2 (ω₂ = {w2:.2f} rad/s): ({K11}−{w2_sq})X₁ + ({K12})X₂ = 0 → "
            f"−{abs(K12)}X₁ = {abs(K12)}X₂ → X₁ = −X₂. "
            f"Masses move in OPPOSITE directions (out-of-phase). The coupling spring is "
            f"maximally compressed/extended, giving a higher frequency.",
        ]


# ═══════════════════════ WORD DOCUMENT BUILDERS ══════════════

def _add_header(doc, set_num, is_solution=False):
    """Add institution header, info table, and instructions to a document."""
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.3)
        section.right_margin = Cm(1.3)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(9.5)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # Institution
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(INSTITUTION)
    run.bold = True
    run.font.size = Pt(11)

    # Program
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(PROGRAM)
    run.font.size = Pt(9.5)

    # Title
    title = EXAM_TITLE + (" — MODEL ANSWERS" if is_solution else "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)

    # Info table
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info = [
        ("Course:", COURSE, "Date:", EXAM_DATE),
        ("Duration:", DURATION, "Total Marks:", TOTAL_MARKS),
        ("Set:", str(set_num), "Program:", "B.Tech Mechatronics Engg."),
    ]
    for i, row_data in enumerate(info):
        for j, text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ""
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            cp.paragraph_format.space_before = Pt(0)
            run = cp.add_run(text)
            run.font.size = Pt(8.5)
            if j % 2 == 0:
                run.bold = True

    if not is_solution:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("Instructions: ")
        run.bold = True
        run.underline = True
        run.font.size = Pt(8.5)
        run2 = p.add_run(
            "Q1 compulsory [4M]. Solve any 2 from Q2/Q3/Q4 [3M each = 6M]. "
            "All subs 0.5M. Assume standard values. Draw diagrams where needed."
        )
        run2.font.size = Pt(8.5)

    # separator
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("―" * 60)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _add_q_block(doc, label, marks, bloom, co, stem, subs, solutions=None):
    """Add one question block. If solutions is given, interleave answers."""
    # Heading + Bloom on same line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label} {marks}")
    run.bold = True
    run.font.size = Pt(9.5)
    run2 = p.add_run(f"    Bloom: {bloom}  |  {co}")
    run2.italic = True
    run2.font.size = Pt(7.5)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Stem
    if stem:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        run3 = p3.add_run(stem)
        run3.font.size = Pt(9)

    for idx, sub in enumerate(subs):
        letter = chr(ord("a") + idx)
        # Question line
        p4 = doc.add_paragraph()
        p4.paragraph_format.left_indent = Cm(0.7)
        p4.paragraph_format.space_after = Pt(0)
        run4 = p4.add_run(f"({letter})  {sub}")
        run4.font.size = Pt(9)
        mark_run = p4.add_run("  [0.5]")
        mark_run.font.size = Pt(8)
        mark_run.bold = True

        # Solution line (if in solution mode)
        if solutions and idx < len(solutions):
            p5 = doc.add_paragraph()
            p5.paragraph_format.left_indent = Cm(1.2)
            p5.paragraph_format.space_after = Pt(0)
            run5 = p5.add_run("Ans: ")
            run5.bold = True
            run5.font.size = Pt(9)
            run5.font.color.rgb = RGBColor(0x00, 0x40, 0x80)
            for line in solutions[idx].split("\n"):
                run_sol = p5.add_run(line)
                run_sol.font.size = Pt(9)
                run_sol.font.color.rgb = RGBColor(0x00, 0x40, 0x80)
                if line != solutions[idx].split("\n")[-1]:
                    p5 = doc.add_paragraph()
                    p5.paragraph_format.left_indent = Cm(1.2)
                    p5.paragraph_format.space_after = Pt(0)


# ─── Template Utilities ──────────────────────────────────────

def _set_no_borders(table):
    """Remove all borders from a Word table (matching template style)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        from docx.oxml import OxmlElement
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for child in list(tblPr):
        if child.tag == qn('w:tblBorders'):
            tblPr.remove(child)
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    ))


def _cell_run(cell, text, bold=False, size=Pt(11), first=True, align=None):
    """Add a formatted run to a cell paragraph. first=True uses paragraphs[0]."""
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = "Times New Roman"
    return p, run


# ─── QP Builder (NEW institutional template format) ─────────

def _build_qp_docx(set_num):
    """Build a QP .docx matching the SVKM'S NMIMS template layout."""
    doc = Document()

    # Page setup: A4, 1-inch (2.54 cm) margins
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # ═══ TABLE 0 — Header (6 rows × 2 cols, no borders) ═══
    ht = doc.add_table(rows=6, cols=2)
    _set_no_borders(ht)
    for row in ht.rows:
        row.cells[0].width = Cm(8.3)
        row.cells[1].width = Cm(7.65)

    # R0 — merged: institution block
    ht.cell(0, 0).merge(ht.cell(0, 1))
    c = ht.cell(0, 0)
    c.text = ""
    p, r = _cell_run(c, "SVKM'S NMIMS", bold=True, size=Pt(12),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run("MUKESH PATEL SCHOOL OF TECHNOLOGY MANAGEMENT & ENGINEERING")
    r.bold = True; r.font.size = Pt(11); r.font.name = "Times New Roman"
    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(0)
    r = p3.add_run("Academic Year: 2025-2026")
    r.font.size = Pt(11); r.font.name = "Times New Roman"

    # R1 — Program / Year
    ht.cell(1, 0).text = ""
    _cell_run(ht.cell(1, 0), "Program: B.Tech    Stream: Mechatronics")
    ht.cell(1, 1).text = ""
    _cell_run(ht.cell(1, 1), "Year: 2nd   Semester: IV")

    # R2 — Subject / Time
    ht.cell(2, 0).text = ""
    _cell_run(ht.cell(2, 0), f"Subject: {COURSE}")
    ht.cell(2, 1).text = ""
    _cell_run(ht.cell(2, 1), "Time: 45 Minutes")

    # R3 — Date / Pages
    ht.cell(3, 0).text = ""
    _cell_run(ht.cell(3, 0), f"Date: {EXAM_DATE}")
    ht.cell(3, 1).text = ""
    _cell_run(ht.cell(3, 1), "No. of Pages: 1")

    # R4 — Marks / Set
    ht.cell(4, 0).text = ""
    _cell_run(ht.cell(4, 0), f"Marks: {TOTAL_MARKS}")
    ht.cell(4, 1).text = ""
    _cell_run(ht.cell(4, 1), f"Set: {set_num}")

    # R5 — merged: "Test-II"
    ht.cell(5, 0).merge(ht.cell(5, 1))
    c = ht.cell(5, 0)
    c.text = ""
    _cell_run(c, "Test-II", bold=True, size=Pt(14),
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # ═══ INSTRUCTIONS (between tables) ═══
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Instructions: Candidates should read carefully the instructions.")
    r.bold = True; r.font.size = Pt(11); r.font.name = "Times New Roman"

    for txt in [
        "1) Figures in brackets on the right hand side indicate full marks.",
        "3) Assume Suitable data if necessary.",
        "4) Question 1 is compulsory.",
        "5) Answer any 2 from the remaining questions.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(txt)
        r.font.size = Pt(11); r.font.name = "Times New Roman"

    # ═══ TABLE 1 — Questions (6 rows × 4 cols, no borders) ═══
    # Columns: C0 ~9% (Q#/CO/SO/BL) | C1 ~4% (letter) | C2 ~80% (text) | C3 ~7% (marks)
    qt = doc.add_table(rows=6, cols=4)
    _set_no_borders(qt)
    for row in qt.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(0.65)
        row.cells[2].width = Cm(12.75)
        row.cells[3].width = Cm(1.05)

    q1a = Q1A_BANKS[set_num]
    q1b = Q1B_BANKS[set_num]
    q2  = Q2_BANKS[set_num]
    q3  = Q3_BANKS[set_num]
    q4  = Q4_BANKS[set_num]

    # ── Row 0  Q1 header ──
    qt.cell(0, 0).text = ""
    _cell_run(qt.cell(0, 0), "Q1", bold=True)
    qt.cell(0, 1).text = ""
    qt.cell(0, 2).text = ""
    _cell_run(qt.cell(0, 2), "Answer briefly:", size=Pt(12))
    qt.cell(0, 3).text = ""

    # ── Row 1  Q1(a) ──
    c0 = qt.cell(1, 0); c0.text = ""
    _cell_run(c0, "CO- 1,2;", size=Pt(9))
    _cell_run(c0, "SO- 1,2;", size=Pt(9), first=False)
    _cell_run(c0, "BL- 1,2", size=Pt(9), first=False)

    qt.cell(1, 1).text = ""
    _cell_run(qt.cell(1, 1), "a.")

    c2 = qt.cell(1, 2); c2.text = ""
    for i, q in enumerate(q1a["questions"]):
        roman = ["(i)", "(ii)", "(iii)", "(iv)"][i]
        _cell_run(c2, f"{roman} {q}", first=(i == 0))

    qt.cell(1, 3).text = ""
    _cell_run(qt.cell(1, 3), "[2]")

    # ── Row 2  Q1(b) ──
    c0 = qt.cell(2, 0); c0.text = ""
    _cell_run(c0, "CO- 1,2;", size=Pt(9))
    _cell_run(c0, "SO- 1,2;", size=Pt(9), first=False)
    _cell_run(c0, "BL- 2,3", size=Pt(9), first=False)

    qt.cell(2, 1).text = ""
    _cell_run(qt.cell(2, 1), "b.")

    c2 = qt.cell(2, 2); c2.text = ""
    for i, q in enumerate(q1b["questions"]):
        roman = ["(i)", "(ii)", "(iii)", "(iv)"][i]
        _cell_run(c2, f"{roman} {q}", first=(i == 0))

    qt.cell(2, 3).text = ""
    _cell_run(qt.cell(2, 3), "[2]")

    # ── Rows 3–5  Q2, Q3, Q4 ──
    q_meta = [
        (3, q2, "Q2", "CO- 2;",   "SO- 2;",   "BL- 3"),
        (4, q3, "Q3", "CO- 2;",   "SO- 2;",   "BL- 3,4"),
        (5, q4, "Q4", "CO- 2,3;", "SO- 2,3;", "BL- 4,5"),
    ]
    for ri, qb, ql, co_t, so_t, bl_t in q_meta:
        # C0: Q# (bold) + CO/SO/BL
        c0 = qt.cell(ri, 0); c0.text = ""
        _cell_run(c0, ql, bold=True)
        _cell_run(c0, co_t, size=Pt(9), first=False)
        _cell_run(c0, so_t, size=Pt(9), first=False)
        _cell_run(c0, bl_t, size=Pt(9), first=False)

        qt.cell(ri, 1).text = ""   # C1 empty for Q2-Q4

        # C2: stem + sub-parts
        c2 = qt.cell(ri, 2); c2.text = ""
        _cell_run(c2, qb["stem"])
        for j, sub in enumerate(qb["subs"]):
            letter = chr(ord("a") + j)
            _cell_run(c2, f"({letter}) {sub}", first=False)

        # C3: marks
        qt.cell(ri, 3).text = ""
        _cell_run(qt.cell(ri, 3), "[3]")

    return doc


# ─── Solution Builder (original paragraph format) ───────────

def _build_solutions_docx(set_num):
    """Build solution .docx using the original paragraph-based layout."""
    doc = Document()
    _add_header(doc, set_num, is_solution=True)

    q1a = Q1A_BANKS[set_num]
    q1b = Q1B_BANKS[set_num]
    q2  = Q2_BANKS[set_num]
    q3  = Q3_BANKS[set_num]
    q4  = Q4_BANKS[set_num]

    q2_sols = solve_q2(set_num)
    q3_sols = solve_q3(set_num)
    q4_sols = solve_q4(set_num)

    _add_q_block(doc, "Q1(a)", "[2 Marks]",
                 "BL 1–2 (Remember / Understand)", "CO1, CO2",
                 "Answer the following:", q1a["questions"], q1a["solutions"])

    _add_q_block(doc, "Q1(b)", "[2 Marks]",
                 "BL 2–3 (Understand / Apply)", "CO1, CO2",
                 None, q1b["questions"], q1b["solutions"])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Solve any TWO from Q2, Q3, Q4  [3 Marks each = 6 Marks]")
    run.bold = True
    run.font.size = Pt(9.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_q_block(doc, "Q2.", "[3 Marks]", q2["bloom"], q2["co"],
                 q2["stem"], q2["subs"], q2_sols)
    _add_q_block(doc, "Q3.", "[3 Marks]", q3["bloom"], q3["co"],
                 q3["stem"], q3["subs"], q3_sols)
    _add_q_block(doc, "Q4.", "[3 Marks]", q4["bloom"], q4["co"],
                 q4["stem"], q4["subs"], q4_sols)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run("— End of Model Answers —")
    run.bold = True
    run.font.size = Pt(9)

    return doc


def build_docx(set_num, include_solutions=False):
    """Build one Word file — template-format QP or paragraph-format solutions."""
    if include_solutions:
        return _build_solutions_docx(set_num)
    return _build_qp_docx(set_num)


# ═══════════════════════ MAIN ════════════════════════════════

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print("=" * 65)
    print("DYNAMIC SYSTEMS MODELING — M2 QUESTION PAPER GENERATOR (v2)")
    print("=" * 65)
    print(f"  Exam   : {EXAM_TITLE}")
    print(f"  Date   : {EXAM_DATE}")
    print(f"  Sets   : 4")
    print(f"  Format : Word (.docx) — SVKM'S NMIMS template")
    print()

    # Generate QPs (new template format)
    print("  Question Papers (template format):")
    for s in range(1, 5):
        doc = build_docx(s, include_solutions=False)
        fname = f"DSM_M2_Set_{s:02d}.docx"
        doc.save(os.path.join(OUTPUT_DIR, fname))
        print(f"    ✓ {fname}")

    # Generate Solutions (paragraph format)
    print()
    print("  Model Answers:")
    for s in range(1, 5):
        doc = build_docx(s, include_solutions=True)
        fname = f"DSM_M2_Solutions_Set_{s:02d}.docx"
        doc.save(os.path.join(OUTPUT_DIR, fname))
        print(f"    ✓ {fname}")

    print()
    print(f"  All files saved to: {OUTPUT_DIR}")
    print("=" * 65)


if __name__ == "__main__":
    main()
