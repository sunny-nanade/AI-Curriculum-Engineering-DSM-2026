"""
Question Paper Template Generator with Detailed Mapping Tables
This shows two formats for the QP with mapping tables
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

def create_qp_template_format1():
    """
    FORMAT 1: Individual table after each question
    """
    doc = Document()
    
    # Header
    heading = doc.add_heading('SVKM\'s Narsee Monjee Institute of Management Studies', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subheading = doc.add_heading('Mukesh Patel School of Technology Management and Engineering', level=2)
    subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("")
    
    # Course details
    course_para = doc.add_paragraph()
    course_para.add_run("Program: B Tech (Mechatronics Engineering)\n").bold = True
    course_para.add_run("Course: Dynamic Systems Modeling\n").bold = True
    course_para.add_run("Course Code: 702MH0C025\n").bold = True
    course_para.add_run("Semester: IV\n").bold = True
    course_para.add_run("Question Paper Set: 1\n").bold = True
    course_para.add_run("Max Marks: 100 (Paper of 140 marks)\n").bold = True
    course_para.add_run("Duration: 3 Hours\n").bold = True
    
    doc.add_paragraph("")
    
    # Instructions
    inst_heading = doc.add_heading("Instructions:", level=2)
    doc.add_paragraph("1. Total Marks: 100 (Question paper contains questions worth 140 marks)")
    doc.add_paragraph("2. Q1: Solve all 4 questions (5 marks each = 20 marks)")
    doc.add_paragraph("3. Q2-Q7: Solve any 4 out of 6 questions (20 marks each = 80 marks)")
    doc.add_paragraph("4. All questions carry equal marks within their section")
    
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    # Q1 Section
    q1_heading = doc.add_heading("Q1. Answer ALL questions (5 marks each):", level=2)
    
    # Q1(a)
    doc.add_paragraph("")
    q1a = doc.add_paragraph()
    q1a.add_run("(a) Define kinematics and kinetics in the context of dynamic systems. Provide suitable examples.")
    
    # Table for Q1(a)
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    headers = ['UNIT', 'COURSE OUTCOME', 'BLOOM\'S TAXONOMY', 'DIFFICULTY LEVEL', 'MARKS']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Data row
    data = ['1', 'CO1', '2-UNDERSTANDING', 'Low', '5']
    for i, value in enumerate(data):
        table.rows[1].cells[i].text = value
    
    doc.add_paragraph("")
    
    # Q1(b)
    q1b = doc.add_paragraph()
    q1b.add_run("(b) Explain the concept of degrees of freedom in rigid body kinematics.")
    
    # Table for Q1(b)
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Data row
    data = ['3', 'CO1', '2-UNDERSTANDING', 'Low', '5']
    for i, value in enumerate(data):
        table.rows[1].cells[i].text = value
    
    doc.add_paragraph("")
    
    # Q1(c)
    q1c = doc.add_paragraph()
    q1c.add_run("(c) Calculate the work done by a constant force of 10N moving a particle through 5m.")
    
    # Table for Q1(c)
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Data row
    data = ['6', 'CO5', '3-APPLYING', 'Low', '5']
    for i, value in enumerate(data):
        table.rows[1].cells[i].text = value
    
    doc.add_paragraph("")
    
    # Q1(d)
    q1d = doc.add_paragraph()
    q1d.add_run("(d) List the applications of Lagrangian mechanics in modern mechatronics systems.")
    
    # Table for Q1(d)
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Data row
    data = ['2', 'CO4', '1-REMEMBERING', 'Low', '5']
    for i, value in enumerate(data):
        table.rows[1].cells[i].text = value
    
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    # Q2-Q7 Section
    doc.add_page_break()
    q2_heading = doc.add_heading("Q2-Q7. Answer any FOUR out of SIX questions (20 marks each):", level=2)
    
    # Q2
    doc.add_paragraph("")
    q2 = doc.add_paragraph()
    q2.add_run("Q2. Derive the equations of motion for a particle moving in a rotating reference frame. Apply the derived equations to solve a practical problem of your choice with appropriate assumptions.")
    
    # Table for Q2
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Data row
    data = ['3, 4', 'CO2, CO4', '4-ANALYZING', 'High', '20']
    for i, value in enumerate(data):
        table.rows[1].cells[i].text = value
    
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    # Q3
    q3 = doc.add_paragraph()
    q3.add_run("Q3. Design and simulate a spring-mass-damper system using computational tools. Evaluate the system response for different damping coefficients (underdamped, critically damped, overdamped). Include code snippets and result plots.")
    
    # Table for Q3
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Data row
    data = ['7', 'CO4, CO5', '6-CREATING', 'Medium', '20']
    for i, value in enumerate(data):
        table.rows[1].cells[i].text = value
    
    doc.add_paragraph("")
    
    # Save
    doc.save("QP_TEMPLATE_FORMAT1_Individual_Tables.docx")
    print("Format 1 saved: QP_TEMPLATE_FORMAT1_Individual_Tables.docx")
    print("(Individual table after each question)")


def create_qp_template_format2():
    """
    FORMAT 2: Combined mapping table at the end
    """
    doc = Document()
    
    # Header
    heading = doc.add_heading('SVKM\'s Narsee Monjee Institute of Management Studies', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subheading = doc.add_heading('Mukesh Patel School of Technology Management and Engineering', level=2)
    subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("")
    
    # Course details
    course_para = doc.add_paragraph()
    course_para.add_run("Program: B Tech (Mechatronics Engineering)\n").bold = True
    course_para.add_run("Course: Dynamic Systems Modeling\n").bold = True
    course_para.add_run("Course Code: 702MH0C025\n").bold = True
    course_para.add_run("Semester: IV\n").bold = True
    course_para.add_run("Question Paper Set: 1\n").bold = True
    course_para.add_run("Max Marks: 100 (Paper of 140 marks)\n").bold = True
    course_para.add_run("Duration: 3 Hours\n").bold = True
    
    doc.add_paragraph("")
    
    # Instructions
    inst_heading = doc.add_heading("Instructions:", level=2)
    doc.add_paragraph("1. Total Marks: 100 (Question paper contains questions worth 140 marks)")
    doc.add_paragraph("2. Q1: Solve all 4 questions (5 marks each = 20 marks)")
    doc.add_paragraph("3. Q2-Q7: Solve any 4 out of 6 questions (20 marks each = 80 marks)")
    doc.add_paragraph("4. All questions carry equal marks within their section")
    
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    # Q1 Section
    q1_heading = doc.add_heading("Q1. Answer ALL questions (5 marks each):", level=2)
    
    doc.add_paragraph("(a) Define kinematics and kinetics in the context of dynamic systems. Provide suitable examples.")
    doc.add_paragraph("")
    
    doc.add_paragraph("(b) Explain the concept of degrees of freedom in rigid body kinematics.")
    doc.add_paragraph("")
    
    doc.add_paragraph("(c) Calculate the work done by a constant force of 10N moving a particle through 5m.")
    doc.add_paragraph("")
    
    doc.add_paragraph("(d) List the applications of Lagrangian mechanics in modern mechatronics systems.")
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    # Q2-Q7 Section
    q2_heading = doc.add_heading("Q2-Q7. Answer any FOUR out of SIX questions (20 marks each):", level=2)
    
    doc.add_paragraph("")
    doc.add_paragraph("Q2. Derive the equations of motion for a particle moving in a rotating reference frame. Apply the derived equations to solve a practical problem of your choice with appropriate assumptions.")
    doc.add_paragraph("")
    
    doc.add_paragraph("Q3. Design and simulate a spring-mass-damper system using computational tools. Evaluate the system response for different damping coefficients (underdamped, critically damped, overdamped). Include code snippets and result plots.")
    doc.add_paragraph("")
    
    doc.add_paragraph("Q4. Analyze the motion of a pulley system with three masses (m₁=2kg, m₂=3kg, m₃=5kg) connected by inextensible strings. Draw free body diagrams and determine the accelerations and string tensions.")
    doc.add_paragraph("")
    
    doc.add_paragraph("Q5. Apply the work-energy principle to determine the velocity of a slider-crank mechanism at different crank positions. Given: crank length = 50mm, connecting rod length = 150mm, crank speed = 1200 rpm. Derive the general expression and calculate for θ = 45°.")
    doc.add_paragraph("")
    
    doc.add_paragraph("Q6. A rigid body rotates about a fixed axis with angular acceleration α = 2t rad/s². Starting from rest, analyze the angular velocity, angular displacement, and tangential acceleration after 5 seconds. Evaluate the total kinetic energy if the moment of inertia is 10 kg⋅m².")
    doc.add_paragraph("")
    
    doc.add_paragraph("Q7. Develop kinematic equations for a robotic arm operating in 3D space using homogeneous transformation matrices. Compare the results obtained using different reference frames (fixed vs. moving). Justify your choice of reference frame for a specific application.")
    doc.add_paragraph("")
    
    # Combined Mapping Table
    doc.add_page_break()
    mapping_heading = doc.add_heading("QUESTION MAPPING TABLE", level=1)
    mapping_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("")
    
    # Create comprehensive table
    table = doc.add_table(rows=11, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    headers = ['QUESTION', 'UNIT', 'COURSE OUTCOME', 'BLOOM\'S TAXONOMY', 'DIFFICULTY LEVEL', 'MARKS']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Data rows
    questions_data = [
        ['Q1(a)', '1', 'CO1', '2-UNDERSTANDING', 'Low', '5'],
        ['Q1(b)', '3', 'CO1', '2-UNDERSTANDING', 'Low', '5'],
        ['Q1(c)', '6', 'CO5', '3-APPLYING', 'Low', '5'],
        ['Q1(d)', '2', 'CO4', '1-REMEMBERING', 'Low', '5'],
        ['Q2', '3, 4', 'CO2, CO4', '4-ANALYZING', 'High', '20'],
        ['Q3', '7', 'CO4, CO5', '6-CREATING', 'Medium', '20'],
        ['Q4', '4', 'CO2, CO3', '4-ANALYZING', 'Medium', '20'],
        ['Q5', '5, 6', 'CO4, CO5', '3-APPLYING', 'Medium', '20'],
        ['Q6', '3, 5', 'CO2, CO4', '5-EVALUATING', 'High', '20'],
        ['Q7', '3', 'CO1, CO2', '6-CREATING', 'High', '20'],
    ]
    
    for i, data in enumerate(questions_data, 1):
        for j, value in enumerate(data):
            table.rows[i].cells[j].text = value
    
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    # Summary statistics
    summary_heading = doc.add_heading("DISTRIBUTION SUMMARY", level=2)
    
    doc.add_paragraph("Course Outcome Coverage:")
    co_para = doc.add_paragraph()
    co_para.add_run("• CO1: 3 questions (25 marks)\n")
    co_para.add_run("• CO2: 4 questions (60 marks)\n")
    co_para.add_run("• CO3: 1 question (20 marks)\n")
    co_para.add_run("• CO4: 6 questions (85 marks)\n")
    co_para.add_run("• CO5: 3 questions (45 marks)\n")
    
    doc.add_paragraph("")
    doc.add_paragraph("Bloom's Taxonomy Distribution:")
    bloom_para = doc.add_paragraph()
    bloom_para.add_run("• L1-Remembering: 1 question (5 marks)\n")
    bloom_para.add_run("• L2-Understanding: 2 questions (10 marks)\n")
    bloom_para.add_run("• L3-Applying: 2 questions (25 marks)\n")
    bloom_para.add_run("• L4-Analyzing: 2 questions (40 marks)\n")
    bloom_para.add_run("• L5-Evaluating: 1 question (20 marks)\n")
    bloom_para.add_run("• L6-Creating: 2 questions (40 marks)\n")
    
    doc.add_paragraph("")
    doc.add_paragraph("Difficulty Level Distribution:")
    diff_para = doc.add_paragraph()
    diff_para.add_run("• Low: 4 questions (20 marks)\n")
    diff_para.add_run("• Medium: 3 questions (60 marks)\n")
    diff_para.add_run("• High: 3 questions (60 marks)\n")
    
    # Save
    doc.save("QP_TEMPLATE_FORMAT2_Combined_Table.docx")
    print("Format 2 saved: QP_TEMPLATE_FORMAT2_Combined_Table.docx")
    print("(Combined mapping table at the end with summary)")


if __name__ == "__main__":
    print("=" * 80)
    print("Creating Question Paper Templates")
    print("=" * 80)
    print()
    
    create_qp_template_format1()
    print()
    create_qp_template_format2()
    
    print()
    print("=" * 80)
    print("Both template formats created successfully!")
    print("=" * 80)
    print()
    print("FORMAT 1: Individual table after each question")
    print("- Good for: Immediate reference while reading")
    print("- Takes more space")
    print()
    print("FORMAT 2: Combined table at the end")
    print("- Good for: Clean question paper, systematic review")
    print("- Includes distribution summary")
    print("- Recommended for software import")
    print()
    print("Please review both formats and let us know which one you prefer!")
