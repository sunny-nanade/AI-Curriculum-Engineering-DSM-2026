"""
Question Paper Generation Framework
This script will generate question papers based on the template and structure provided by the user.

QP Structure (as per user requirement):
- Total marks: 140 (students solve for 100 marks)
- Q1: 4 questions × 5 marks each = 20 marks (solve all 4)
- Q2-Q7: 6 questions × 20 marks each = 120 marks (solve any 4 out of 6 = 80 marks)
- Total: 20 + 80 = 100 marks (out of 140 marks paper)

Requirements:
1. Map Course Outcomes (CO) to each question
2. Fair distribution across entire syllabus
3. Map COs with Bloom's Taxonomy levels
4. Generate 2 sets per course (total 6 QPs)
5. Output in simple Word format for easy copy-paste
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from course_syllabus_summary import ALL_COURSES
from blooms_taxonomy_verbs import BLOOMS_TAXONOMY, get_bloom_level_from_verb
import json
import random

class QuestionPaper:
    """Class to manage Question Paper generation"""
    
    def __init__(self, course_key, set_number=1):
        self.course_key = course_key
        self.set_number = set_number
        self.course_info = ALL_COURSES[course_key]
        self.questions = {
            "Q1": [],  # 4 questions of 5 marks each
            "Q2_to_Q7": []  # 6 questions of 20 marks each
        }
        
    def add_question(self, question_type, question_text, marks, unit, co, bloom_level):
        """
        Add a question to the question paper
        
        Args:
            question_type: "Q1" or "Q2_to_Q7"
            question_text: The question text
            marks: Marks for the question
            unit: Unit number(s) covered
            co: Course Outcome(s) mapped
            bloom_level: Bloom's Taxonomy level
        """
        question = {
            "text": question_text,
            "marks": marks,
            "unit": unit,
            "co": co,
            "bloom_level": bloom_level
        }
        self.questions[question_type].append(question)
        
    def validate_structure(self):
        """Validate QP structure meets requirements"""
        validations = []
        
        # Check Q1 has exactly 4 questions
        if len(self.questions["Q1"]) != 4:
            validations.append(f"Q1 should have 4 questions, found {len(self.questions['Q1'])}")
            
        # Check Q2-Q7 has exactly 6 questions
        if len(self.questions["Q2_to_Q7"]) != 6:
            validations.append(f"Q2-Q7 should have 6 questions, found {len(self.questions['Q2_to_Q7'])}")
            
        # Check Q1 marks
        q1_total = sum([q["marks"] for q in self.questions["Q1"]])
        if q1_total != 20:
            validations.append(f"Q1 total marks should be 20, found {q1_total}")
            
        # Check Q2-Q7 marks
        for q in self.questions["Q2_to_Q7"]:
            if q["marks"] != 20:
                validations.append(f"Each Q2-Q7 question should be 20 marks, found {q['marks']}")
                
        return validations
        
    def get_co_distribution(self):
        """Get distribution of Course Outcomes across questions"""
        co_count = {}
        all_questions = self.questions["Q1"] + self.questions["Q2_to_Q7"]
        
        for q in all_questions:
            cos = q["co"] if isinstance(q["co"], list) else [q["co"]]
            for co in cos:
                co_count[co] = co_count.get(co, 0) + 1
                
        return co_count
        
    def get_unit_distribution(self):
        """Get distribution of units across questions"""
        unit_count = {}
        all_questions = self.questions["Q1"] + self.questions["Q2_to_Q7"]
        
        for q in all_questions:
            units = q["unit"] if isinstance(q["unit"], list) else [q["unit"]]
            for unit in units:
                unit_count[unit] = unit_count.get(unit, 0) + 1
                
        return unit_count
        
    def generate_word_document(self, output_path):
        """
        Generate Word document for the question paper
        This will be a simple format for easy copy-paste
        """
        doc = Document()
        
        # Header
        heading = doc.add_heading(self.course_info["course_name"], 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Course details
        doc.add_paragraph(f"Course Code: {self.course_info['course_code']}")
        doc.add_paragraph(f"Semester: {self.course_info['semester']}")
        doc.add_paragraph(f"Question Paper Set: {self.set_number}")
        doc.add_paragraph("")
        
        # Instructions
        doc.add_heading("Instructions:", level=2)
        doc.add_paragraph("1. Total Marks: 100 (Paper contains 140 marks)")
        doc.add_paragraph("2. Q1: Solve all 4 questions (5 marks each = 20 marks)")
        doc.add_paragraph("3. Q2-Q7: Solve any 4 out of 6 questions (20 marks each = 80 marks)")
        doc.add_paragraph("")
        
        # Q1 Section
        doc.add_heading("Q1. Answer all questions (5 marks each):", level=2)
        for i, q in enumerate(self.questions["Q1"], 1):
            para = doc.add_paragraph()
            para.add_run(f"({chr(96+i)}) {q['text']}").bold = True
            para.add_run(f" [CO: {q['co']}, BL: {q['bloom_level']}, Unit: {q['unit']}]").italic = True
        doc.add_paragraph("")
        
        # Q2-Q7 Section
        doc.add_heading("Q2-Q7. Answer any 4 out of 6 questions (20 marks each):", level=2)
        for i, q in enumerate(self.questions["Q2_to_Q7"], 2):
            para = doc.add_paragraph()
            para.add_run(f"Q{i}. {q['text']}").bold = True
            para.add_run(f" [CO: {q['co']}, BL: {q['bloom_level']}, Unit: {q['unit']}]").italic = True
            doc.add_paragraph("")
        
        # Save document
        doc.save(output_path)
        print(f"Question paper saved to: {output_path}")
        
    def display_summary(self):
        """Display summary of the question paper"""
        print("=" * 80)
        print(f"Question Paper Summary - {self.course_info['course_name']} (Set {self.set_number})")
        print("=" * 80)
        
        # Validation
        validations = self.validate_structure()
        if validations:
            print("\n⚠ VALIDATION ISSUES:")
            for v in validations:
                print(f"  - {v}")
        else:
            print("\n✓ Structure validated successfully")
            
        # CO Distribution
        print("\nCourse Outcome Distribution:")
        co_dist = self.get_co_distribution()
        for co, count in sorted(co_dist.items()):
            print(f"  {co}: {count} question(s)")
            
        # Unit Distribution
        print("\nUnit Distribution:")
        unit_dist = self.get_unit_distribution()
        for unit, count in sorted(unit_dist.items()):
            print(f"  Unit {unit}: {count} question(s)")
            
        print("=" * 80)


# Example template (to be filled based on user's template)
def create_sample_qp():
    """Create a sample question paper for demonstration"""
    qp = QuestionPaper("DSM", set_number=1)
    
    # Q1 questions (5 marks each) - These are placeholders
    qp.add_question("Q1", "Define kinematics and kinetics in the context of dynamic systems.", 
                   5, 1, "CO1", "L1")
    qp.add_question("Q1", "Explain Newton's laws of motion with suitable examples.", 
                   5, 2, "CO3", "L2")
    qp.add_question("Q1", "List the types of coordinate frames used in motion analysis.", 
                   5, 1, "CO1", "L1")
    qp.add_question("Q1", "Describe the concept of degrees of freedom in rigid body kinematics.", 
                   5, 3, "CO1", "L2")
    
    # Q2-Q7 questions (20 marks each) - These are placeholders
    qp.add_question("Q2_to_Q7", "Derive the equations of motion for a particle moving in a rotating reference frame.", 
                   20, [3, 4], ["CO2", "CO4"], "L4")
    qp.add_question("Q2_to_Q7", "Apply the work-energy principle to analyze a two degree of freedom system.", 
                   20, [5, 6], ["CO4", "CO5"], "L3")
    qp.add_question("Q2_to_Q7", "Analyze the motion of a 3D printer head using particle kinetics principles.", 
                   20, 4, ["CO2", "CO3"], "L4")
    qp.add_question("Q2_to_Q7", "Design a simulation using ODE solvers for a simple pendulum system.", 
                   20, 7, ["CO4", "CO5"], "L6")
    qp.add_question("Q2_to_Q7", "Evaluate the energy dissipated in a damped oscillator using Lagrangian mechanics.", 
                   20, [2, 6], ["CO4", "CO5"], "L5")
    qp.add_question("Q2_to_Q7", "Develop the free body diagram and equations of motion for a pulley system.", 
                   20, [2, 4], ["CO2", "CO3"], "L3")
    
    return qp


if __name__ == "__main__":
    # Create sample QP
    sample_qp = create_sample_qp()
    sample_qp.display_summary()
    
    # Generate Word document
    output_file = f"DSM/DSM_QP_Set1_SAMPLE.docx"
    sample_qp.generate_word_document(output_file)
