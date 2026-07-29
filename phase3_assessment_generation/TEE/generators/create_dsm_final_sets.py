"""
DSM Question Papers — Sets 1, 2, 3
====================================
Dynamic Systems Modeling (702MH0C025)
B.Tech Mechatronics, Semester IV

Structure per set:
  Q1 (compulsory): 4 × 5 = 20 marks  (3 numericals + 1 theory)
  Q2–Q7: 20 marks each, solve any 4 = 80 marks
    Q2, Q3 — Pure Numericals (20 marks each)
    Q4, Q5 — Theory 10M + Numerical 10M
    Q6, Q7 — FBD 10M + Difficult Numerical 10M
  Total written: 140   Evaluated: 100

Identical synoptic table stamped on all 3 sets.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
HANDBOOK_IMG = os.path.join(SCRIPT_DIR, "Numerical_Solutions_Handbook", "media", "handbook")
FBD_IMG = os.path.join(SCRIPT_DIR, "DSM_FBD_Notebooks", "DSM_FBD_Notebooks", "media")

# ── Course Outcomes ──
COURSE_OUTCOMES = {
    "CO1": "Understand and apply theoretical concepts in kinematics",
    "CO2": "Analyze the given system using the Free Body Diagram",
    "CO3": "Understand and calculate the forces acting on dynamic systems",
    "CO4": "Apply the constitutive laws to analyze dynamic systems of particles and rigid bodies",
    "CO5": "Understand and calculate the energy of a dynamic system and power dissipated",
}

# ── Synoptic Table (identical across all 3 sets) ──
MASTER_ROWS = [
    # (Q.No, Topic Area, Unit, CO, BL, Marks, Difficulty, Type)
    ("Q1 (a)", "Kinematics / Rigid Body Theory",         "1,3", "CO1",      "L2", "5",  "Easy",   "Theory"),
    ("Q1 (b)", "Position, Velocity, Acceleration",        "1",   "CO1",      "L3", "5",  "Easy",   "Numerical"),
    ("Q1 (c)", "Constrained Motion / Pulleys",            "4",   "CO3",      "L3", "5",  "Medium", "Numerical"),
    ("Q1 (d)", "Energy / Rigid Body Calculation",         "5,6", "CO5",      "L3", "5",  "Medium", "Numerical"),
    ("Q2",     "Particle Kinetics — Pure Numerical",      "4",   "CO3+CO4",  "L3", "20", "Medium", "Numerical"),
    ("Q3",     "Rigid Body Dynamics / Vibrations",        "5",   "CO4+CO5",  "L3", "20", "Medium", "Numerical"),
    ("Q4",     "Rotating Frames / Kinematics",            "3",   "CO1+CO3",  "L4", "20", "Medium", "Theory+Num"),
    ("Q5",     "Constitutive Laws / Energy Methods",      "2,6", "CO4+CO5",  "L4", "20", "Hard",   "Theory+Num"),
    ("Q6",     "FBD Analysis + Difficult Numerical",      "4,5", "CO2+CO3",  "L5", "20", "Hard",   "FBD+Num"),
    ("Q7",     "FBD Analysis + Difficult Numerical",      "5,6", "CO2+CO5",  "L5", "20", "Hard",   "FBD+Num"),
]

# ═══════════════ HELPER FUNCTIONS ═══════════════

def new_doc(set_number):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)

    # page borders
    for s in doc.sections:
        pgB = OxmlElement('w:pgBorders')
        pgB.set(qn('w:offsetFrom'), 'page')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '18')
            b.set(qn('w:space'), '24'); b.set(qn('w:color'), '000000')
            pgB.append(b)
        s._sectPr.append(pgB)

    def ctr(txt, sz=12, bold=True, color=None):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(0)
        r = p.add_run(txt); r.bold = bold; r.font.size = Pt(sz); r.font.name = 'Times New Roman'
        if color: r.font.color.rgb = color

    ctr("SVKM's NMIMS (Deemed-to-be University)", 14)
    ctr("School of Technology Management & Engineering (STME), Indore", 10)
    ctr("TERM END EXAMINATION — MARCH 2026", 11)
    ctr(f"Dynamic Systems Modeling (702MH0C025) — SET {set_number}", 13)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    for label, val in [("Program: ","B.Tech Mechatronics"), ("  Semester: ","IV"),
                       ("  Duration: ","3 Hours"), ("  Max Marks: ","100")]:
        r = p.add_run(label); r.bold = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'
        r = p.add_run(val); r.font.size = Pt(9); r.font.name = 'Times New Roman'

    line(doc)

    inst = doc.add_paragraph()
    inst.paragraph_format.space_after = Pt(4)
    r = inst.add_run("Instructions: "); r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'
    r = inst.add_run("Q1 is compulsory (20 marks). Attempt any FOUR from Q2–Q7 (4 × 20 = 80 marks). "
                     "Draw neat Free Body Diagrams wherever applicable. No programmable calculators. "
                     "All numerical answers must include proper units. "
                     "Assume g = 9.81 m/s² unless stated otherwise.")
    r.font.size = Pt(9); r.font.name = 'Times New Roman'

    line(doc)
    return doc


def line(doc):
    p = doc.add_paragraph("─" * 95)
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    for r in p.runs:
        r.font.size = Pt(6); r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0x99,0x99,0x99)


def q_head(doc, number, marks, text=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"Q{number}  [{marks} marks]"); r.bold = True; r.font.size = Pt(11); r.font.name = 'Times New Roman'
    if text:
        r = p.add_run(f"  —  {text}"); r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0x33,0x33,0x33)


def sub_q(doc, label, text, marks=5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(f"({label})  "); r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'
    r = p.add_run(text); r.font.size = Pt(10); r.font.name = 'Times New Roman'
    r = p.add_run(f"  [{marks}]"); r.bold = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0x00,0x00,0x99)


def big_q(doc, number, q_type, intro, subs):
    """Q2–Q7: 20 marks, with 2 or 4 sub-questions depending on type."""
    q_head(doc, number, 20, q_type)
    if intro:
        p = doc.add_paragraph(intro)
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(10); r.font.name = 'Times New Roman'; r.italic = True
    for label, text, marks in subs:
        sub_q(doc, label, text, marks)
    line(doc)


def try_image(doc, folder, filename, width=Inches(3.5)):
    path = os.path.join(folder, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)


def add_synoptic_table(doc):
    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SYNOPTIC TABLE — QUESTION PAPER MAPPING"); r.bold = True
    r.font.size = Pt(12); r.font.name = 'Times New Roman'

    headers = ["Q. No.", "Topic Area", "Unit", "CO", "BL", "Marks", "Difficulty", "Type"]
    tbl = doc.add_table(rows=1+len(MASTER_ROWS), cols=len(headers))
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]; cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(8); r.font.name = 'Times New Roman'
    for ri, row_data in enumerate(MASTER_ROWS, 1):
        for ci, val in enumerate(row_data):
            cell = tbl.rows[ri].cells[ci]; cell.text = val
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(8); r.font.name = 'Times New Roman'

    # Coverage summary
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Coverage Summary\n"); r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'
    r = p.add_run(
        "Units: 1, 2, 3, 4, 5, 6 (all theory units covered; Unit 7 has no numerical)\n"
        "COs: CO1–CO5 (all Course Outcomes assessed)\n"
        "Bloom's Levels: L2–L5 (Understand → Evaluate)\n"
        "Difficulty: Easy → Medium → Hard (50 marks easy, 80+ difficult)\n"
        "Q2–Q7: 2 Pure Numerical + 2 Theory+Numerical + 2 FBD+Difficult Numerical"
    )
    r.font.size = Pt(9); r.font.name = 'Times New Roman'

    # CO table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Course Outcomes:\n"); r.bold = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'
    for co, desc in COURSE_OUTCOMES.items():
        r = p.add_run(f"  {co}: {desc}\n"); r.font.size = Pt(8); r.font.name = 'Times New Roman'


# ═══════════════ SET 1 ═══════════════

def write_set1(doc):
    # ── Q1 (Compulsory) ──
    q_head(doc, 1, 20, "Compulsory — Answer ALL parts")

    sub_q(doc, "a",
          "Define the following with respect to rigid body motion: "
          "(i) Degrees of Freedom for a rigid body in 2D and 3D, "
          "(ii) Rolling without slipping condition, and "
          "(iii) Instantaneous center of rotation. "
          "Give one real-world example for each.", 5)

    sub_q(doc, "b",
          "A projectile is launched from the ground with initial velocity components "
          "v_x0 = 12 m/s (horizontal) and v_y0 = 16 m/s (vertical). "
          "Calculate: (i) the maximum height reached, "
          "(ii) the total time of flight, "
          "(iii) the horizontal range, and "
          "(iv) the velocity vector and speed at t = 1 s. "
          "Take g = 9.81 m/s².", 5)

    sub_q(doc, "c",
          "In an Atwood machine, two masses m₁ = 5 kg and m₂ = 3 kg are connected "
          "by a light inextensible string passing over a frictionless, massless pulley. "
          "The system starts from rest. Find: (i) the acceleration of the system, "
          "(ii) the tension in the string, and "
          "(iii) the velocity of each mass after 1.5 seconds. "
          "Verify that the tension lies between the two weights.", 5)

    sub_q(doc, "d",
          "A block of mass 3 kg slides down a rough inclined plane making an angle of "
          "35° with the horizontal. The coefficient of kinetic friction is μ_k = 0.25. "
          "The block starts from rest. Using the work-energy theorem, find the speed "
          "of the block after it has traveled 4 m along the incline.", 5)

    line(doc)

    # ── Q2 — Pure Numerical: Particle Kinetics ──
    big_q(doc, 2, "Pure Numerical — Particle Kinetics",
          "Attempt all four parts.",
          [("a",
            "A 5 kg block is pushed up a 30° frictionless inclined plane with an initial "
            "velocity of 10 m/s. Find: (i) the acceleration of the block, "
            "(ii) the distance traveled along the incline before the block stops, and "
            "(iii) the time taken to stop.",
            5),
           ("b",
            "A block-and-tackle system consists of 4 pulleys, each with an efficiency of 95%. "
            "A load of 200 kg needs to be lifted through a height of 2 m. "
            "Calculate: (i) the ideal mechanical advantage, "
            "(ii) the overall system efficiency, "
            "(iii) the actual input force required, and "
            "(iv) the total energy wasted as heat.",
            5),
           ("c",
            "A particle of mass 1 kg moves along a vertical circular loop of radius 2 m. "
            "Determine: (i) the minimum speed at the topmost point to maintain contact, "
            "(ii) the minimum entry speed at the bottom of the loop, and "
            "(iii) the normal force at the bottom when the entry speed is 12 m/s.",
            5),
           ("d",
            "A movable pulley system is used to lift a 50 kg load. "
            "If a force of 300 N is applied to the free end of the rope, "
            "find the upward acceleration of the load. "
            "Also determine the force required for the load to move at constant velocity.",
            5)])

    # ── Q3 — Pure Numerical: Rigid Body & Vibrations ──
    big_q(doc, 3, "Pure Numerical — Rigid Body Dynamics & Vibrations",
          "Attempt all four parts.",
          [("a",
            "Three objects — a solid sphere, a solid disk, and a hollow cylinder — "
            "all having the same mass and radius, are released from rest at the top of "
            "a 30° incline of length 5 m. Each rolls without slipping. "
            "Calculate the linear acceleration, time to reach the bottom, and final velocity "
            "for each object. Which object arrives first and why?",
            5),
           ("b",
            "A physical pendulum consists of a uniform rod of mass 2 kg and length 1 m, "
            "pivoted at one end. Find: (i) the moment of inertia about the pivot, "
            "(ii) the natural frequency of small oscillations in rad/s, and "
            "(iii) the period. Compare this period with that of a simple pendulum "
            "of the same length 1 m.",
            5),
           ("c",
            "A spring-mass system consists of a block of mass 2 kg attached to a spring "
            "of stiffness 200 N/m. The block is displaced 0.1 m from equilibrium and "
            "released from rest (no damping). Find: (i) the natural frequency, "
            "(ii) the period, (iii) the maximum velocity, and (iv) the total mechanical energy.",
            5),
           ("d",
            "A spring-mass-damper system has m = 2 kg, k = 500 N/m, and c = 20 N·s/m. "
            "Determine: (i) the natural frequency, "
            "(ii) the critical damping coefficient, "
            "(iii) the damping ratio, "
            "(iv) the damped frequency, and "
            "(v) classify the system. "
            "Find the time for the oscillation amplitude to decay to 10% of the initial value.",
            5)])

    # ── Q4 — Theory + Numerical: Rotating Frames ──
    big_q(doc, 4, "Theory 10M + Numerical 10M — Rotating Frames",
          None,
          [("a",
            "State and derive the Transport Theorem for a vector quantity observed from "
            "both a fixed and a rotating reference frame. Explain each term in the velocity "
            "relation between the two frames and state its physical significance.",
            5),
           ("b",
            "Starting from the Transport Theorem, derive the complete acceleration equation "
            "in a rotating reference frame. Identify and explain the physical origin of the "
            "Coriolis, centripetal, and Euler acceleration terms.",
            5),
           ("c",
            "An object on a rotating platform has the following parameters: "
            "angular velocity Ω = 3.0 rad/s (CCW about the z-axis), "
            "position vector r = 1.0 m (along the positive x-axis), "
            "relative velocity v_rel = (0.5 î + 0.8 ĵ) m/s, "
            "relative acceleration a_rel = (0.2 î − 0.1 ĵ) m/s², and angular acceleration α = 0. "
            "Calculate each component of the acceleration (Coriolis, centripetal, relative) "
            "and determine the total acceleration vector in the fixed frame.",
            5),
           ("d",
            "A projectile is fired due north at a latitude of 45° N with a speed of 500 m/s. "
            "Given the Earth's angular velocity Ω = 7.3 × 10⁻⁵ rad/s, "
            "calculate: (i) the Coriolis acceleration magnitude, and "
            "(ii) the lateral deflection after 10 seconds of flight. "
            "State and justify the direction of deflection.",
            5)])

    # ── Q5 — Theory + Numerical: Constitutive Laws & Energy ──
    big_q(doc, 5, "Theory 10M + Numerical 10M — Constitutive Laws & Energy",
          None,
          [("a",
            "Derive the Euler-Lagrange equation starting from Hamilton's principle of "
            "least action. Explain the significance of generalized coordinates and "
            "how Lagrangian mechanics automatically eliminates constraint forces.",
            5),
           ("b",
            "Compare Newton's force-based formulation with Lagrange's energy-based formulation "
            "for deriving equations of motion. State the advantages and limitations of each. "
            "Illustrate with the example of a simple pendulum.",
            5),
           ("c",
            "A block of mass 4 kg is released from rest at the top of a frictionless curved "
            "ramp of height 3 m. At the bottom, it enters a rough horizontal surface with "
            "coefficient of kinetic friction μ_k = 0.3. Using conservation of energy "
            "and the work-energy theorem, find: (i) the speed at the bottom of the ramp, and "
            "(ii) the distance the block slides on the rough surface before stopping.",
            5),
           ("d",
            "A ball of mass 0.15 kg traveling at 40 m/s is struck by a bat. "
            "The ball leaves at 60 m/s in the opposite direction. "
            "The contact time is 0.002 s. Find: "
            "(i) the impulse delivered to the ball, "
            "(ii) the average force exerted by the bat, and "
            "(iii) the kinetic energy before and after the hit. "
            "Is kinetic energy conserved? Explain.",
            5)])

    # ── Q6 — FBD + Difficult Numerical ──
    big_q(doc, 6, "FBD 10M + Difficult Numerical 10M",
          None,
          [("a",
            "A 2 kg package is transferred from a conveyor belt moving at 1.5 m/s to another "
            "belt moving at 3.0 m/s. The coefficients of friction are μ_s = 0.35 and "
            "μ_k = 0.25. Draw a detailed Free Body Diagram showing all forces acting on the "
            "package while it is on the faster belt. "
            "Determine: (i) the friction force and acceleration, "
            "(ii) the time for the package to match the speed of the second belt, "
            "(iii) the distance traveled by the package and by the belt surface during this time, and "
            "(iv) the energy dissipated as heat during the transfer.",
            10),
           ("b",
            "A slider-crank mechanism has a crank radius r = 50 mm, connecting rod length "
            "L = 150 mm, and operates at 3000 RPM. At a crank angle θ = 30° from TDC, "
            "calculate: (i) the angular velocity of the crank in rad/s, "
            "(ii) the connecting rod angle φ, "
            "(iii) the piston position measured from the crank center, and "
            "(iv) the piston velocity. State the direction of piston motion.",
            10)])

    # ── Q7 — FBD + Difficult Numerical ──
    big_q(doc, 7, "FBD 10M + Difficult Numerical 10M",
          None,
          [("a",
            "A yo-yo of mass 50 g has an outer radius R = 30 mm and an inner spool radius "
            "r = 10 mm. Model it as a solid disk (I = ½mR²). "
            "Draw the complete Free Body Diagram showing all forces and the rolling constraint. "
            "Derive the expression for the downward acceleration a = g/(1 + β), where β = I/(mr²). "
            "Calculate: (i) the acceleration, (ii) the string tension, "
            "(iii) the angular acceleration, and "
            "(iv) the fraction of total kinetic energy that is rotational.",
            10),
           ("b",
            "A 2-DOF mass-spring system has masses m₁ = 1.0 kg and m₂ = 1.5 kg, with three "
            "springs arranged as: Wall — k₁(100 N/m) — m₁ — k₂(150 N/m) — m₂ — k₃(100 N/m) — Wall. "
            "Write the mass matrix [M] and stiffness matrix [K]. "
            "Set up the eigenvalue equation det([K] − ω²[M]) = 0, solve for the two natural "
            "frequencies, and describe the two mode shapes physically (in-phase vs out-of-phase).",
            10)])


# ═══════════════ SET 2 ═══════════════

def write_set2(doc):
    # ── Q1 (Compulsory) ──
    q_head(doc, 1, 20, "Compulsory — Answer ALL parts")

    sub_q(doc, "a",
          "Explain the difference between holonomic and non-holonomic constraints "
          "with one example each. How does the number of constraints affect the degrees "
          "of freedom of a mechanical system? State the general formula for DOF.", 5)

    sub_q(doc, "b",
          "A particle moves in a plane such that its position is given by x(t) = 4t³ − 2t "
          "and y(t) = 5t² + 3, where x and y are in meters and t in seconds. "
          "Find: (i) the position vector and distance from the origin at t = 2 s, "
          "(ii) the velocity vector and speed at t = 2 s, and "
          "(iii) the acceleration vector and its magnitude at t = 2 s.", 5)

    sub_q(doc, "c",
          "A 2 kg block is placed on an inclined plane making an angle of 40° with the "
          "horizontal. The coefficient of kinetic friction between the block and the surface "
          "is μ_k = 0.3. The block is released from rest. "
          "Find: (i) the normal force, (ii) the friction force, and "
          "(iii) the acceleration down the incline.", 5)

    sub_q(doc, "d",
          "A solid disk of mass 5 kg and radius 0.2 m is subjected to a constant angular "
          "acceleration of α = 3.75 rad/s² starting from rest. "
          "Find: (i) the moment of inertia about the center, "
          "(ii) the angular velocity after 2 s, and "
          "(iii) the tangential and centripetal acceleration components at a point on the "
          "rim at t = 2 s. Also find the total acceleration magnitude at that point.", 5)

    line(doc)

    # ── Q2 — Pure Numerical: Constrained Motion & Pulleys ──
    big_q(doc, 2, "Pure Numerical — Constrained Motion & Pulleys",
          "Attempt all four parts.",
          [("a",
            "A particle enters the bottom of a vertical circular loop of radius 8 m. "
            "Find: (i) the minimum entry speed to complete the loop, and "
            "(ii) the normal force at the topmost point when the entry speed is 22 m/s. "
            "The particle has mass 2 kg.",
            5),
           ("b",
            "A 6-pulley block-and-tackle system is used to lift a 2000 kg load "
            "through a height of 3 m. Each pulley has an efficiency of 92%. "
            "Calculate: (i) the ideal mechanical advantage, "
            "(ii) the overall system efficiency, "
            "(iii) the required input force, and "
            "(iv) the total energy wasted as heat.",
            5),
           ("c",
            "An Atwood machine has masses m₁ = 7 kg and m₂ = 5 kg connected by a "
            "light inextensible string over a frictionless, massless pulley. "
            "Starting from rest, find: (i) the acceleration, "
            "(ii) the tension in the string, and "
            "(iii) the velocity of each mass after 2 seconds. "
            "Verify that the tension lies between the two weights.",
            5),
           ("d",
            "A CNC machine tool must execute a straight-line cut from point "
            "P₀ = (10, 20, 5) mm to P₁ = (110, 70, 35) mm at a feedrate of 50 mm/s. "
            "Calculate: (i) the total travel distance, "
            "(ii) the direction unit vector, "
            "(iii) the required speed for each axis motor, and "
            "(iv) the total move time.",
            5)])

    # ── Q3 — Pure Numerical: Vibrations & Rolling ──
    big_q(doc, 3, "Pure Numerical — Vibrations & Rolling Dynamics",
          "Attempt all four parts.",
          [("a",
            "In a free vibration test, successive peak amplitudes are measured as 20 mm and "
            "15 mm, with a damped period of 0.35 s. "
            "Find: (i) the logarithmic decrement, "
            "(ii) the damping ratio, "
            "(iii) the undamped natural frequency (rad/s and Hz), and "
            "(iv) the spring stiffness if the vibrating mass is 2 kg.",
            5),
           ("b",
            "A wheel of radius 0.4 m rolls without slipping along a straight path. "
            "The center of the wheel moves at 3.0 m/s with an acceleration of 1.5 m/s². "
            "Calculate the angular velocity ω and angular acceleration α. "
            "Then find the magnitude and direction of the total acceleration at the "
            "topmost and bottommost points of the wheel.",
            5),
           ("c",
            "A solid disk of mass 5 kg and radius 0.3 m is driven by a constant torque "
            "of 2 N·m against viscous friction with damping coefficient c = 0.5 N·m·s/rad. "
            "Calculate: (i) the moment of inertia, "
            "(ii) the time constant, and "
            "(iii) the steady-state angular velocity.",
            5),
           ("d",
            "A two-stage gear train has Stage 1 (15-tooth driver meshing with 45-tooth driven) "
            "and Stage 2 (12-tooth driver meshing with 60-tooth driven). "
            "The motor runs at 1500 RPM with an input torque of 0.5 N·m. "
            "Find: (i) the overall gear ratio, "
            "(ii) the output speed in RPM, "
            "(iii) the output torque, and "
            "(iv) verify that power is conserved across the gear train.",
            5)])

    # ── Q4 — Theory + Numerical: Kinematics & Coordinate Systems ──
    big_q(doc, 4, "Theory 10M + Numerical 10M — Rigid Body Kinematics",
          None,
          [("a",
            "Derive the velocity formula for a point P on a rigid body: "
            "v_P = v_center + ω × r. "
            "Explain the physical significance of each term. "
            "What is the velocity of the contact point during rolling without slipping?",
            5),
           ("b",
            "Define the three basic rotation matrices R_x(α), R_y(β), and R_z(γ). "
            "What is the gimbal lock problem in Euler angle representation? "
            "How do quaternions address this issue?",
            5),
           ("c",
            "A wheel of radius R = 0.5 m rolls without slipping with its center moving "
            "at v_center = 2.0 m/s to the right. "
            "Find: (i) the angular velocity, and the velocity (magnitude and direction) at "
            "(ii) the topmost point, (iii) the bottommost (contact) point, and "
            "(iv) the rightmost point of the wheel.",
            5),
           ("d",
            "A rod of length 1.0 m rotates about a fixed pivot at one end with "
            "angular velocity ω = 5 rad/s and angular acceleration α = 2 rad/s². "
            "Find the tangential acceleration, normal acceleration, total acceleration "
            "magnitude, and velocity at: (i) the free end, (ii) the midpoint, and "
            "(iii) a point 0.3 m from the pivot.",
            5)])

    # ── Q5 — Theory + Numerical: Energy & Momentum ──
    big_q(doc, 5, "Theory 10M + Numerical 10M — Energy & Momentum",
          None,
          [("a",
            "State and derive the work-energy theorem for a particle. "
            "Distinguish between conservative and non-conservative forces with examples. "
            "How is the theorem modified when non-conservative forces are present?",
            5),
           ("b",
            "Define impulse, linear momentum, and coefficient of restitution. "
            "Derive the impulse-momentum theorem. "
            "Explain the conservation of linear momentum for an isolated system.",
            5),
           ("c",
            "A car of mass 1200 kg accelerates from 20 m/s to 40 m/s in 8 s on a "
            "level road against a constant drag force of 400 N. "
            "Find: (i) the net force from Newton's second law, "
            "(ii) the engine force, "
            "(iii) the average power delivered by the engine, and "
            "(iv) the instantaneous power at 40 m/s.",
            5),
           ("d",
            "Two identical balls (mass 2 kg each) undergo a head-on collision. Ball A "
            "travels at 6 m/s and Ball B is at rest. The coefficient of restitution is "
            "e = 0.5. Find: (i) the velocities of both balls after collision, "
            "(ii) the impulse during impact, and "
            "(iii) the kinetic energy lost in the collision.",
            5)])

    # ── Q6 — FBD + Difficult Numerical ──
    big_q(doc, 6, "FBD 10M + Difficult Numerical 10M",
          None,
          [("a",
            "A mobile manipulator on an AGV moves at v_AGV = 0.5 m/s (horizontal). "
            "The robot arm of length L = 1.2 m is at an angle θ = 30° from the vertical, "
            "rotating at angular velocity dθ/dt = −0.2 rad/s. A payload of total mass 2.5 kg "
            "(end-effector + object) is at the tip. "
            "Draw the complete Free Body Diagram of the tip mass showing weight, arm reaction, "
            "and any other relevant forces. "
            "Using the transport theorem, calculate the absolute velocity of the tip. "
            "If a vibration isolator (k = 1000 N/m, c = 40 N·s/m) is attached to the "
            "end-effector, determine the natural frequency and damping ratio.",
            10),
           ("b",
            "A CNC machine executes a move of 200 mm using a trapezoidal velocity profile "
            "with v_max = 100 mm/s and a_max = 500 mm/s². "
            "Calculate: (i) the acceleration time and distance, "
            "(ii) the cruise distance and time, "
            "(iii) the total move time. "
            "Then, repeat the calculation for a short move of 15 mm and determine whether "
            "the profile is trapezoidal or triangular. If triangular, find the peak velocity "
            "and total time.",
            10)])

    # ── Q7 — FBD + Difficult Numerical ──
    big_q(doc, 7, "FBD 10M + Difficult Numerical 10M",
          None,
          [("a",
            "A spring-mass-damper system has m = 2 kg, k = 500 N/m, c = 20 N·s/m, "
            "with initial displacement x₀ = 0.1 m and zero initial velocity. "
            "Draw the complete Free Body Diagram of the mass in the displaced position, "
            "clearly labeling the spring force, damping force, and weight (about equilibrium). "
            "Write the equation of motion in standard form. "
            "Determine: (i) the natural frequency, (ii) the critical damping coefficient, "
            "(iii) the damping ratio, (iv) the damped frequency, and "
            "(v) classify the system response. "
            "Write the expression for the free response x(t).",
            10),
           ("b",
            "A double pendulum consists of two masses m₁ = m₂ = 1 kg at the ends of "
            "two massless rods of lengths L₁ = L₂ = 1 m. "
            "Write the kinetic energy T and potential energy V for the system. "
            "Using Euler-Lagrange equations, derive the two coupled equations of motion. "
            "For small angles, write the linearized form. "
            "Discuss the significance of this system in the study of deterministic chaos.",
            10)])


# ═══════════════ SET 3 ═══════════════

def write_set3(doc):
    # ── Q1 (Compulsory) ──
    q_head(doc, 1, 20, "Compulsory — Answer ALL parts")

    sub_q(doc, "a",
          "State the three damping regimes in a single-DOF vibrating system: underdamped, "
          "critically damped, and overdamped. For each, state the condition on the damping "
          "ratio ζ and describe the physical nature of the response. "
          "Which regime is preferred for automotive suspension design and why?", 5)

    sub_q(doc, "b",
          "A rocket is launched vertically with its height given by y(t) = 60t − 10t² meters "
          "(t in seconds). Find: (i) the initial velocity, "
          "(ii) the constant acceleration, "
          "(iii) the time at which the rocket reaches maximum height, and "
          "(iv) the maximum height reached.", 5)

    sub_q(doc, "c",
          "An Atwood machine has masses m₁ = 6 kg and m₂ = 4 kg connected by a light "
          "inextensible string over a frictionless, massless pulley. Starting from rest, "
          "find: (i) the acceleration of the system, "
          "(ii) the tension in the string, and "
          "(iii) the velocity of each mass after 1 second.", 5)

    sub_q(doc, "d",
          "Convert the Cartesian coordinates (6, 8) to polar coordinates (r, θ). "
          "Then convert the cylindrical coordinates (ρ = 3 m, φ = 45°, z = 4 m) to "
          "Cartesian coordinates (x, y, z). "
          "Verify both conversions by computing the distance from the origin.", 5)

    line(doc)

    # ── Q2 — Pure Numerical: Incline & Pulley Systems ──
    big_q(doc, 2, "Pure Numerical — Incline & Pulley Systems",
          "Attempt all four parts.",
          [("a",
            "A 5 kg block on a 30° incline with coefficient of kinetic friction μ_k = 0.2 "
            "is given an initial velocity of 10 m/s directed up the incline. "
            "Find: (i) the deceleration while the block moves up, "
            "(ii) the distance traveled before the block stops, "
            "(iii) the time to stop, and "
            "(iv) whether the block slides back down (compare gravitational component along "
            "the incline with the maximum static friction force, given μ_s = 0.3).",
            5),
           ("b",
            "An elevator car of mass 1200 kg is counterbalanced by a 1000 kg counterweight "
            "connected via a cable over a pulley. Four passengers (each weighing 75 kg) are "
            "inside the car. When the brake is released, find: "
            "(i) the total mass on the car side, "
            "(ii) the acceleration of the system, and "
            "(iii) the tension in the cable.",
            5),
           ("c",
            "A 1 kg particle on a frictionless inclined plane at 30° is released from rest. "
            "Find: (i) the normal force, "
            "(ii) the acceleration down the incline, "
            "(iii) the velocity after 3 seconds, and "
            "(iv) the distance traveled in 3 seconds.",
            5),
           ("d",
            "A CNC machine must execute a 250 mm move with v_max = 150 mm/s "
            "and a_max = 1000 mm/s². Calculate the complete trapezoidal velocity profile: "
            "acceleration time, cruise time, deceleration time, and total time. "
            "If the move distance were only 10 mm, determine whether the profile would be "
            "trapezoidal or triangular, and find the peak velocity and total time.",
            5)])

    # ── Q3 — Pure Numerical: Rigid Body & Energy ──
    big_q(doc, 3, "Pure Numerical — Rigid Body Dynamics & Energy",
          "Attempt all four parts.",
          [("a",
            "A solid disk of mass 10 kg and radius 0.15 m rolls without slipping down a 20° "
            "incline starting from rest. Find: (i) the linear acceleration of the center, "
            "(ii) the velocity of the center after rolling 3 m, and "
            "(iii) the time to travel 3 m. "
            "Compare the acceleration with that of a frictionless sliding block on the same incline.",
            5),
           ("b",
            "A pendulum consists of a 2 kg mass on a string of length 1.5 m. "
            "The pendulum is released from rest at an angle of 60° from the vertical. "
            "Using energy conservation, find: (i) the speed at the lowest point, and "
            "(ii) the tension in the string at the lowest point.",
            5),
           ("c",
            "A car of mass 1500 kg is traveling at 100 km/h when the brakes are applied, "
            "exerting a constant friction force of 8000 N. Using the work-energy theorem, "
            "find: (i) the braking distance, and "
            "(ii) the total energy dissipated as heat.",
            5),
           ("d",
            "A vibrating system has m = 2 kg, k = 800 N/m, and damping ratio ζ = 0.2. "
            "Find: (i) the natural frequency in rad/s, "
            "(ii) the damped frequency, "
            "(iii) the critical damping coefficient, "
            "(iv) the actual damping coefficient, and "
            "(v) the number of complete oscillations for the amplitude to decay to half "
            "its initial value.",
            5)])

    # ── Q4 — Theory + Numerical: Rigid Body Kinematics ──
    big_q(doc, 4, "Theory 10M + Numerical 10M — Rigid Body Properties",
          None,
          [("a",
            "Define moment of inertia and state the Parallel Axis Theorem with proof. "
            "Starting from the definition, derive the moment of inertia of a uniform rod "
            "of mass m and length L about its center. Then use the Parallel Axis Theorem "
            "to find the MOI about one end.",
            5),
           ("b",
            "Explain the slider-crank mechanism with a neat sketch. "
            "State the geometric constraint equation relating the crank angle θ "
            "to the connecting rod angle φ. "
            "Discuss what causes asymmetric piston motion in an IC engine and the "
            "effect of increasing the L/r ratio (connecting rod to crank) on piston kinematics.",
            5),
           ("c",
            "Calculate the moment of inertia about the center of mass for: "
            "(i) a solid disk of mass 5 kg and radius 0.2 m, "
            "(ii) a uniform rod of mass 5 kg and length 1 m about its center, "
            "(iii) the same rod about one end, and "
            "(iv) a solid sphere of mass 5 kg and radius 0.2 m. "
            "For case (iii), verify the result using the Parallel Axis Theorem from case (ii).",
            5),
           ("d",
            "A slider-crank mechanism has crank radius r = 50 mm, connecting rod length "
            "L = 150 mm, and operates at 6000 RPM. At crank angle θ = 90°, determine: "
            "(i) the angular velocity of the crank in rad/s, "
            "(ii) the connecting rod angle φ, "
            "(iii) the piston velocity, and "
            "(iv) state whether the piston acceleration at θ = 0° (TDC) at 6000 RPM "
            "exceeds 500 g. Justify your answer.",
            5)])

    # ── Q5 — Theory + Numerical: Lagrangian & Advanced ──
    big_q(doc, 5, "Theory 10M + Numerical 10M — Lagrangian & Advanced Topics",
          None,
          [("a",
            "Derive the equation of motion for a simple pendulum (mass m, length L) using "
            "the Lagrangian method. Identify the kinetic energy T, potential energy V, "
            "and generalized coordinate. Apply the Euler-Lagrange equation and show that it "
            "yields the same EOM as Newton's force method.",
            5),
           ("b",
            "Explain the concept of normal modes and modal analysis for a multi-DOF system. "
            "What is the eigenvalue problem in the context of vibrations? "
            "How are natural frequencies and mode shapes determined from the mass and "
            "stiffness matrices?",
            5),
           ("c",
            "A simple pendulum has mass m = 2 kg and length L = 1.5 m. "
            "Find: (i) the Lagrangian, (ii) the equation of motion, and "
            "(iii) the natural frequency and period for small oscillations. "
            "Compare the period with a physical pendulum consisting of a uniform rod "
            "of mass 2 kg and length 1.5 m pivoted at one end.",
            5),
           ("d",
            "Briefly explain three types of numerical ODE solvers: Forward Euler, "
            "Runge-Kutta 2nd order (RK2), and Runge-Kutta 4th order (RK4). "
            "For each, state the order of accuracy and comment on stability. "
            "Why is RK4 generally preferred for engineering simulations?",
            5)])

    # ── Q6 — FBD + Difficult Numerical ──
    big_q(doc, 6, "FBD 10M + Difficult Numerical 10M",
          None,
          [("a",
            "A rotating sensor array has angular velocity ω = 2.0 rad/s about a vertical axis. "
            "A sensor of mass 0.2 kg is at radial position r = 0.5 m and moves radially "
            "outward at ṙ = 0.3 m/s. "
            "Draw the complete Free Body Diagram in the rotating frame, showing all real "
            "and fictitious forces (centrifugal, Coriolis). "
            "Calculate: (i) the centripetal acceleration, "
            "(ii) the Coriolis acceleration, "
            "(iii) the total acceleration magnitude and direction, and "
            "(iv) the constraint forces required to maintain this motion.",
            10),
           ("b",
            "A block of mass 3 kg starts from rest and slides 4 m down a rough incline at 35° "
            "(μ_k = 0.25). At the bottom, it transitions to a rough horizontal surface "
            "(μ_k = 0.3). Using energy methods throughout, find: "
            "(i) the speed at the bottom of the incline, "
            "(ii) the distance traveled on the horizontal surface before stopping, and "
            "(iii) the total energy dissipated by friction (on both surfaces combined).",
            10)])

    # ── Q7 — FBD + Difficult Numerical ──
    big_q(doc, 7, "FBD 10M + Difficult Numerical 10M",
          None,
          [("a",
            "A 2 kg package is transferred from Conveyor Belt 1 (moving at 1.5 m/s) "
            "to Conveyor Belt 2 (moving at 3.0 m/s). The coefficients of friction are "
            "μ_s = 0.35 and μ_k = 0.25. "
            "Draw the complete Free Body Diagram showing all forces during the sliding phase. "
            "Calculate: (i) the kinetic friction force and the resulting acceleration, "
            "(ii) the time to match Belt 2's speed, "
            "(iii) the change in kinetic energy of the package, "
            "(iv) the total work done by friction, and "
            "(v) the energy efficiency of the transfer (ratio of KE gained to total friction work).",
            10),
           ("b",
            "A compound gear train has two stages: Stage 1 (20-tooth driver meshing with "
            "80-tooth driven) and Stage 2 (15-tooth driver meshing with 60-tooth driven). "
            "The input shaft runs at 2000 RPM with torque 2.0 N·m. Each gear stage has "
            "an efficiency of 96%. "
            "Find: (i) the overall gear ratio, "
            "(ii) the output speed, "
            "(iii) the output torque accounting for efficiency losses, and "
            "(iv) the total power loss. "
            "If the load requires 25 N·m at the output speed, determine if the motor is sufficient.",
            10)])


# ═══════════════ MAIN ═══════════════

if __name__ == "__main__":
    os.makedirs("Final Papers", exist_ok=True)

    sets = {1: write_set1, 2: write_set2, 3: write_set3}

    for num, writer in sets.items():
        doc = new_doc(num)
        writer(doc)
        add_synoptic_table(doc)
        path = f"Final Papers/DSM_QP_Set{num}_FINAL.docx"
        doc.save(path)
        print(f"  [OK] {path}")

    print("\nDone — 3 DSM question papers with identical synoptic table.")
