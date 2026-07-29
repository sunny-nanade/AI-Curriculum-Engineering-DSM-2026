"""
DSM Answer Booklets — Sets 1, 2, 3
====================================
Dynamic Systems Modeling (702MH0C025)
B.Tech Mechatronics, Semester IV

Comprehensive solutions with step-by-step working,
mark allocation, and key notes.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# ═══════════════ HELPER FUNCTIONS ═══════════════

def new_doc(set_number):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.5); s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)
    for s in doc.sections:
        pgB = OxmlElement('w:pgBorders')
        pgB.set(qn('w:offsetFrom'), 'page')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '18')
            b.set(qn('w:space'), '24'); b.set(qn('w:color'), '000000')
            pgB.append(b)
        s._sectPr.append(pgB)

    def ctr(txt, sz=12, bold=True, color=None):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(0)
        r = p.add_run(txt); r.bold = bold; r.font.size = Pt(sz); r.font.name = 'Times New Roman'
        if color: r.font.color.rgb = color

    ctr("SVKM's NMIMS — STME Indore", 13)
    ctr(f"DSM (702MH0C025) — ANSWER BOOKLET — SET {set_number}", 12)
    ctr("Term End Examination — March 2026", 10, bold=False)
    sep(doc)
    return doc


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'
    r.font.size = Pt(12 if level == 1 else 10)
    r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)


def point(doc, text, marks=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(9); r.font.name = 'Times New Roman'
    if marks:
        r = p.add_run(f"  [{marks}]"); r.bold = True; r.font.size = Pt(8)
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(8); r.font.name = 'Times New Roman'
    r.italic = True; r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def marks_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0x00, 0x00, 0xCC)


def sep(doc):
    p = doc.add_paragraph("─" * 95)
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    for r in p.runs:
        r.font.size = Pt(6); r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0x99,0x99,0x99)


# ═══════════════ SET 1 ANSWERS ═══════════════

def answers_set1(doc):
    # ---- Q1(a) ----
    heading(doc, "Q1 (a)  —  Definitions [5 marks]")
    point(doc, "(i) DOF: In 2D, a rigid body has 3 DOF (x, y, θ); in 3D, it has 6 DOF "
          "(x, y, z + three rotational angles). DOF = total coordinates − number of constraints.", "1.5")
    point(doc, "(ii) Rolling without slipping: v_cm = ωR. The contact point has zero velocity "
          "relative to the surface. Example: car wheel on road.", "2")
    point(doc, "(iii) Instantaneous Center of Rotation (ICR): A point on or off the body about "
          "which all points appear to rotate at that instant. Example: contact point of a rolling wheel.", "1.5")
    sep(doc)

    # ---- Q1(b) ----
    heading(doc, "Q1 (b)  —  Projectile Motion [5 marks]")
    point(doc, "Given: v_x0 = 12 m/s, v_y0 = 16 m/s, g = 9.81 m/s²")
    point(doc, "h_max = v_y0²/(2g) = 256 / 19.62 = 13.05 m", "1")
    point(doc, "t_peak = v_y0/g = 16/9.81 = 1.631 s; Total T = 2 × 1.631 = 3.262 s", "1")
    point(doc, "Range R = v_x0 × T = 12 × 3.262 = 39.14 m", "1")
    point(doc, "At t = 1 s: v_x = 12 m/s, v_y = 16 − 9.81(1) = 6.19 m/s", "1")
    point(doc, "Speed = √(12² + 6.19²) = √(144 + 38.32) = √182.32 = 13.50 m/s", "1")
    sep(doc)

    # ---- Q1(c) ----
    heading(doc, "Q1 (c)  —  Atwood Machine [5 marks]")
    point(doc, "Given: m₁ = 5 kg, m₂ = 3 kg, g = 9.81 m/s²")
    point(doc, "a = (m₁ − m₂)g / (m₁ + m₂) = (2)(9.81)/8 = 2.453 m/s²", "1.5")
    point(doc, "T = 2m₁m₂g / (m₁ + m₂) = 2(5)(3)(9.81)/8 = 36.79 N", "1.5")
    point(doc, "v(1.5 s) = a × t = 2.453 × 1.5 = 3.68 m/s", "1")
    point(doc, "Check: m₂g = 29.43 N < T = 36.79 N < m₁g = 49.05 N ✓", "1")
    sep(doc)

    # ---- Q1(d) ----
    heading(doc, "Q1 (d)  —  Work-Energy on Incline [5 marks]")
    point(doc, "Given: m = 3 kg, θ = 35°, μ_k = 0.25, s = 4 m, v₀ = 0")
    point(doc, "N = mg cos 35° = 3(9.81)(0.8192) = 24.11 N", "1")
    point(doc, "W_gravity = mgs sin 35° = 3(9.81)(4)(0.5736) = 67.51 J", "1")
    point(doc, "W_friction = −μ_k N s = −0.25(24.11)(4) = −24.11 J", "1")
    point(doc, "W_net = 67.51 − 24.11 = 43.40 J", "1")
    point(doc, "½mv² = W_net → v = √(2 × 43.40 / 3) = √28.93 = 5.38 m/s", "1")
    sep(doc)

    # ---- Q2 ----
    heading(doc, "Q2  —  Particle Kinetics [20 marks]")

    heading(doc, "Q2 (a)  —  Frictionless Incline [5 marks]", 2)
    point(doc, "Given: m = 5 kg, θ = 30°, v₀ = 10 m/s up, frictionless")
    point(doc, "a = −g sin 30° = −9.81 × 0.5 = −4.905 m/s² (decelerating)", "1")
    point(doc, "d = v₀² / (2 × 4.905) = 100 / 9.81 = 10.19 m", "2")
    point(doc, "t = v₀ / 4.905 = 10 / 4.905 = 2.039 s", "2")

    heading(doc, "Q2 (b)  —  Block and Tackle [5 marks]", 2)
    point(doc, "Given: 4 pulleys, η = 0.95 each, m_load = 200 kg, h = 2 m")
    point(doc, "MA_ideal = 4; η_total = 0.95⁴ = 0.8145", "1")
    point(doc, "F_ideal = 200(9.81)/4 = 490.5 N", "1")
    point(doc, "F_actual = F_ideal / η_total = 490.5 / 0.8145 = 602.2 N", "1")
    point(doc, "W_in = F_actual × (4 × 2) = 602.2 × 8 = 4817.6 J", "1")
    point(doc, "W_out = mgh = 200(9.81)(2) = 3924.0 J; Heat = 4817.6 − 3924.0 = 893.6 J", "1")

    heading(doc, "Q2 (c)  —  Vertical Circular Loop [5 marks]", 2)
    point(doc, "Given: R = 2 m, m = 1 kg, v_bottom = 12 m/s")
    point(doc, "v_top_min = √(gR) = √(9.81 × 2) = 4.43 m/s (when N_top = 0)", "1")
    point(doc, "v_bottom_min = √(5gR) = √(5 × 9.81 × 2) = √98.1 = 9.90 m/s", "1")
    point(doc, "Energy conservation: ½mv²_bottom = ½mv²_top + 2mgR")
    point(doc, "v²_top = 144 − 4(9.81) = 144 − 39.24 = 104.76; v_top = 10.24 m/s", "1")
    point(doc, "N_bottom = mv²/R + mg = 1(144)/2 + 9.81 = 72 + 9.81 = 81.81 N", "2")
    note(doc, "N_top = mv²_top/R − mg = 1(104.76)/2 − 9.81 = 52.38 − 9.81 = 42.57 N")

    heading(doc, "Q2 (d)  —  Movable Pulley [5 marks]", 2)
    point(doc, "Given: m_load = 50 kg, F = 300 N applied")
    point(doc, "For constant velocity: F = W/2 = 50(9.81)/2 = 245.25 N", "2")
    point(doc, "For F = 300 N: Net force = 2F − mg = 600 − 490.5 = 109.5 N", "1.5")
    point(doc, "a = 109.5 / 50 = 2.19 m/s²", "1.5")
    sep(doc)

    # ---- Q3 ----
    heading(doc, "Q3  —  Rigid Body & Vibrations [20 marks]")

    heading(doc, "Q3 (a)  —  Rolling Race [5 marks]", 2)
    point(doc, "For rolling objects: a = g sin θ / (1 + β), where β = I/(mR²)")
    point(doc, "Sphere: β = 2/5 → a = 9.81 sin 30° / 1.4 = 4.905 / 1.4 = 3.504 m/s²", "1")
    point(doc, "Disk: β = 1/2 → a = 4.905 / 1.5 = 3.270 m/s²", "1")
    point(doc, "Cylinder (hollow): β = 1 → a = 4.905 / 2 = 2.453 m/s²", "1")
    point(doc, "Time (t = √(2s/a)): Sphere: 1.689 s, Disk: 1.749 s, Cylinder: 2.019 s", "1")
    point(doc, "Velocity (v = √(2as)): Sphere: 5.92 m/s, Disk: 5.72 m/s, Cylinder: 4.95 m/s", "0.5")
    point(doc, "Winner: Solid sphere (lowest β → highest acceleration)", "0.5")

    heading(doc, "Q3 (b)  —  Physical Pendulum [5 marks]", 2)
    point(doc, "Given: rod m = 2 kg, L = 1 m, pivot at end")
    point(doc, "I_O = mL²/3 = 2(1)/3 = 0.667 kg·m²", "1")
    point(doc, "d (pivot to CG) = L/2 = 0.5 m", "0.5")
    point(doc, "ω_n = √(mgd/I_O) = √(2 × 9.81 × 0.5 / 0.667) = √(14.72) = 3.836 rad/s", "1.5")
    point(doc, "T = 2π/ω_n = 2π/3.836 = 1.638 s", "1")
    point(doc, "Simple pendulum (same L): T_s = 2π√(1/9.81) = 2.006 s → Physical is 18.4% faster", "1")

    heading(doc, "Q3 (c)  —  Undamped Spring-Mass [5 marks]", 2)
    point(doc, "Given: m = 2 kg, k = 200 N/m, x₀ = 0.1 m, v₀ = 0")
    point(doc, "ω_n = √(k/m) = √(200/2) = √100 = 10 rad/s", "1")
    point(doc, "f_n = ω_n/(2π) = 10/6.283 = 1.592 Hz; T = 1/f_n = 0.628 s", "1")
    point(doc, "v_max = ω_n × x₀ = 10 × 0.1 = 1.0 m/s", "1.5")
    point(doc, "E_total = ½kx₀² = ½(200)(0.01) = 1.0 J", "1.5")

    heading(doc, "Q3 (d)  —  Spring-Mass-Damper Classification [5 marks]", 2)
    point(doc, "Given: m = 2 kg, k = 500 N/m, c = 20 N·s/m")
    point(doc, "ω_n = √(500/2) = √250 = 15.81 rad/s", "1")
    point(doc, "c_cr = 2√(km) = 2√(1000) = 63.25 N·s/m", "1")
    point(doc, "ζ = c/c_cr = 20/63.25 = 0.316 → Underdamped (ζ < 1)", "1")
    point(doc, "ω_d = ω_n √(1 − ζ²) = 15.81 × √(1 − 0.0999) = 15.81 × 0.9487 = 15.00 rad/s", "1")
    point(doc, "Amplitude envelope: e^(−ζω_n t) = e^(−5t); For 10%: t = ln(10)/5 = 0.461 s", "1")
    sep(doc)

    # ---- Q4 ----
    heading(doc, "Q4  —  Rotating Frames [20 marks]")

    heading(doc, "Q4 (a)  —  Transport Theorem [5 marks]", 2)
    point(doc, "Statement: (dA/dt)_fixed = (dA/dt)_rotating + Ω × A for any vector A", "2")
    point(doc, "Velocity relation: v_fixed = v_rel + Ω × r + v_origin")
    point(doc, "v_rel: velocity measured in rotating frame", "1")
    point(doc, "Ω × r: velocity due to rotation of the frame itself", "1")
    point(doc, "v_origin: velocity of the rotating frame's origin", "1")

    heading(doc, "Q4 (b)  —  Acceleration in Rotating Frame [5 marks]", 2)
    point(doc, "a_fixed = a_rel + 2Ω × v_rel + Ω × (Ω × r) + α × r + a_origin", "2")
    point(doc, "Coriolis: 2Ω × v_rel — deflection of moving objects due to frame rotation", "1")
    point(doc, "Centripetal: Ω × (Ω × r) — always directed toward the rotation axis", "1")
    point(doc, "Euler: α × r — present only when angular velocity is changing", "1")

    heading(doc, "Q4 (c)  —  Rotating Platform Numerical [5 marks]", 2)
    point(doc, "Given: Ω = 3.0 k̂ rad/s, r = 1.0 î m, v_rel = (0.5î + 0.8ĵ) m/s, a_rel = (0.2î − 0.1ĵ), α = 0")
    point(doc, "Coriolis = 2Ω × v_rel = 2(3k̂) × (0.5î + 0.8ĵ) = 6(0.5ĵ − 0.8î) = (−4.8î + 3.0ĵ) m/s²", "1.5")
    note(doc, "Using k̂ × î = ĵ and k̂ × ĵ = −î")
    point(doc, "Centripetal = Ω × (Ω × r) = 3k̂ × (3k̂ × 1.0î) = 3k̂ × 3ĵ = 9(−î) = −9.0î m/s²", "1.5")
    point(doc, "Euler = α × r = 0 (since α = 0)")
    point(doc, "Total: a = (0.2 − 4.8 − 9.0)î + (−0.1 + 3.0)ĵ = (−13.6î + 2.9ĵ) m/s²", "1")
    point(doc, "|a| = √(13.6² + 2.9²) = √(184.96 + 8.41) = √193.37 = 13.91 m/s²", "1")

    heading(doc, "Q4 (d)  —  Coriolis on Earth [5 marks]", 2)
    point(doc, "Given: latitude λ = 45°N, v = 500 m/s due north, Ω = 7.3 × 10⁻⁵ rad/s, t = 10 s")
    point(doc, "Vertical component: Ω_v = Ω sin λ = 7.3 × 10⁻⁵ × sin 45° = 5.162 × 10⁻⁵ rad/s", "1")
    point(doc, "a_Coriolis = 2Ω_v × v = 2(5.162 × 10⁻⁵)(500) = 0.05162 m/s²", "2")
    point(doc, "Deflection d = ½ a_C t² = ½(0.05162)(100) = 2.58 m to the east (right)", "2")
    note(doc, "In Northern hemisphere, Coriolis deflects to the right of the velocity direction")
    sep(doc)

    # ---- Q5 ----
    heading(doc, "Q5  —  Constitutive Laws & Energy [20 marks]")

    heading(doc, "Q5 (a)  —  Euler-Lagrange Derivation [5 marks]", 2)
    point(doc, "Hamilton's Principle: δ∫(T − V)dt = 0 over t₁ to t₂", "1")
    point(doc, "Lagrangian: L = T − V, where T = kinetic energy, V = potential energy", "1")
    point(doc, "Euler-Lagrange: d/dt(∂L/∂q̇) − ∂L/∂q = 0 for each generalized coordinate q", "1")
    point(doc, "Generalized coordinates: minimum set of independent variables describing configuration", "1")
    point(doc, "Advantage: constraint forces never appear → works naturally with complex systems", "1")

    heading(doc, "Q5 (b)  —  Newton vs Lagrange Comparison [5 marks]", 2)
    point(doc, "Newton: Force-based, requires FBD, explicitly handles constraint forces", "1")
    point(doc, "Lagrange: Energy-based, uses scalar T and V, constraints handled automatically", "1")
    point(doc, "Newton preferred: when forces are needed (e.g., joint reactions)", "1")
    point(doc, "Lagrange preferred: multi-DOF, complex constraints, no need for reaction forces", "1")
    point(doc, "Simple pendulum: Newton gives mg sin θ = mLα; Lagrange gives same via T = ½mL²θ̇², V = −mgL cos θ", "1")

    heading(doc, "Q5 (c)  —  Ramp + Rough Surface [5 marks]", 2)
    point(doc, "Given: m = 4 kg, h = 3 m (frictionless), μ_k = 0.3 (horizontal)")
    point(doc, "At bottom: mgh = ½mv² → v = √(2gh) = √(2 × 9.81 × 3) = √58.86 = 7.67 m/s", "2")
    point(doc, "On horizontal: ½mv² = μ_k mg d → d = v²/(2μ_k g) = 58.86/(2 × 0.3 × 9.81) = 10.0 m", "3")

    heading(doc, "Q5 (d)  —  Impulse and Ball-Bat Problem [5 marks]", 2)
    point(doc, "Given: m = 0.15 kg, v₁ = 40 m/s (→), v₂ = 60 m/s (←), Δt = 0.002 s")
    point(doc, "Taking initial direction as positive: v₂ = −60 m/s")
    point(doc, "J = m(v₂ − v₁) = 0.15(−60 − 40) = −15 N·s; |J| = 15 N·s", "1.5")
    point(doc, "F_avg = |J|/Δt = 15/0.002 = 7500 N", "1.5")
    point(doc, "KE_before = ½(0.15)(40²) = 120 J; KE_after = ½(0.15)(60²) = 270 J", "1")
    point(doc, "KE NOT conserved — the bat does positive work on the ball (external energy input)", "1")
    sep(doc)

    # ---- Q6 ----
    heading(doc, "Q6  —  FBD + Difficult Numerical [20 marks]")

    heading(doc, "Q6 (a)  —  Conveyor Belt Transfer [10 marks]", 2)
    point(doc, "Given: m = 2 kg, v₁ = 1.5 m/s, v₂ = 3.0 m/s, μ_s = 0.35, μ_k = 0.25")
    marks_note(doc, "FBD [4 marks]:")
    point(doc, "Weight W = mg = 2(9.81) = 19.62 N (downward)")
    point(doc, "Normal N = 19.62 N (upward)")
    point(doc, "Kinetic friction f_k = μ_k N = 0.25(19.62) = 4.905 N (forward, in belt direction)", "2")
    note(doc, "FBD shows: block on belt surface; W↓, N↑, f_k→ (direction of belt motion)")
    point(doc, "Check: package initially slower than belt → belt drags package forward → kinetic friction acts in belt direction ✓", "2")
    marks_note(doc, "Calculations [6 marks]:")
    point(doc, "a = f_k/m = 4.905/2 = 2.453 m/s²", "1.5")
    point(doc, "Time to match: t = (v₂ − v₁)/a = 1.5/2.453 = 0.612 s", "1.5")
    point(doc, "Package distance: s_pkg = v₁t + ½at² = 1.5(0.612) + ½(2.453)(0.612²) = 0.918 + 0.459 = 1.377 m", "1")
    point(doc, "Belt distance: s_belt = v₂ × t = 3.0(0.612) = 1.836 m", "0.5")
    point(doc, "Relative sliding = 1.836 − 1.377 = 0.459 m", "0.5")
    point(doc, "Heat = f_k × relative sliding = 4.905 × 0.459 = 2.25 J", "1")

    heading(doc, "Q6 (b)  —  Slider-Crank Mechanism [10 marks]", 2)
    point(doc, "Given: r = 50 mm, L = 150 mm, n = 3000 RPM, θ = 30°")
    point(doc, "ω = 2π(3000)/60 = 314.16 rad/s", "2")
    point(doc, "sin φ = (r/L) sin θ = (50/150) sin 30° = 0.3333 × 0.5 = 0.1667 → φ = 9.59°", "2")
    point(doc, "cos φ = 0.9860", "0.5")
    point(doc, "Piston position from crank center: x = r cos θ + L cos φ = 50(0.866) + 150(0.986) = 43.30 + 147.90 = 191.20 mm", "2")
    point(doc, "Angular velocity of connecting rod: φ̇ = rω cos θ/(L cos φ) = 50(314.16)(0.866)/(150 × 0.986) = 13599.1/147.9 = 91.95 rad/s", "1.5")
    point(doc, "Piston velocity: v = −rω sin θ − Lφ̇ sin φ = −50(314.16)(0.5) − 150(91.95)(0.1667) "
          "= −7854 − 2299 = −10153 mm/s = −10.15 m/s", "2")
    note(doc, "Negative sign: piston moves toward crank (away from TDC)")
    sep(doc)

    # ---- Q7 ----
    heading(doc, "Q7  —  FBD + Difficult Numerical [20 marks]")

    heading(doc, "Q7 (a)  —  Yo-Yo Dynamics [10 marks]", 2)
    point(doc, "Given: m = 50 g = 0.05 kg, R = 30 mm = 0.03 m, r = 10 mm = 0.01 m (solid disk)")
    marks_note(doc, "FBD [3 marks]:")
    point(doc, "Weight W = mg = 0.4905 N (downward at center)")
    point(doc, "String tension T (upward, applied at spool radius r)")
    point(doc, "Rolling constraint: a = αr (linear and angular acceleration linked)")
    marks_note(doc, "Derivation [3 marks]:")
    point(doc, "I = ½mR² = ½(0.05)(0.03²) = 2.25 × 10⁻⁵ kg·m²", "1")
    point(doc, "β = I/(mr²) = R²/(2r²) = 0.0009/(2 × 0.0001) = 4.5", "1")
    point(doc, "From ΣF = ma: mg − T = ma; From Στ = Iα: Tr = Iα = I(a/r)")
    point(doc, "Solving: a = g/(1 + β) = 9.81/5.5 = 1.784 m/s²", "1")
    marks_note(doc, "Calculations [4 marks]:")
    point(doc, "T = m(g − a) = 0.05(9.81 − 1.784) = 0.05(8.026) = 0.401 N", "1")
    point(doc, "α = a/r = 1.784/0.01 = 178.4 rad/s²", "1")
    point(doc, "KE_rot/KE_total = β/(1 + β) = 4.5/5.5 = 0.818 = 81.8% rotational", "2")

    heading(doc, "Q7 (b)  —  2-DOF Mass-Spring System [10 marks]", 2)
    point(doc, "Given: m₁ = 1.0 kg, m₂ = 1.5 kg, k₁ = 100, k₂ = 150, k₃ = 100 N/m")
    marks_note(doc, "Matrices [3 marks]:")
    point(doc, "[M] = [[1.0, 0], [0, 1.5]]", "1")
    point(doc, "[K] = [[k₁ + k₂, −k₂], [−k₂, k₂ + k₃]] = [[250, −150], [−150, 250]]", "2")
    marks_note(doc, "Eigenvalue problem [4 marks]:")
    point(doc, "det([K] − ω²[M]) = 0")
    point(doc, "(250 − ω²)(250 − 1.5ω²) − (−150)² = 0", "1")
    point(doc, "1.5ω⁴ − 625ω² + 40000 = 0", "1")
    point(doc, "ω² = (625 ± √(625² − 4(1.5)(40000))) / (2 × 1.5) = (625 ± √150625) / 3 = (625 ± 388.1) / 3", "1")
    point(doc, "ω₁² = 236.9/3 = 78.97 → ω₁ = 8.887 rad/s (f₁ = 1.414 Hz)", "0.5")
    point(doc, "ω₂² = 1013.1/3 = 337.7 → ω₂ = 18.38 rad/s (f₂ = 2.925 Hz)", "0.5")
    marks_note(doc, "Mode shapes [3 marks]:")
    point(doc, "Mode 1 (ω₁): Both masses move in-phase (same direction)", "1.5")
    point(doc, "Mode 2 (ω₂): Masses move out-of-phase (opposite directions)", "1.5")
    sep(doc)


# ═══════════════ SET 2 ANSWERS ═══════════════

def answers_set2(doc):
    # ---- Q1(a) ----
    heading(doc, "Q1 (a)  —  Constraints & DOF [5 marks]")
    point(doc, "Holonomic: constraints expressible as f(q₁, q₂, ..., t) = 0. Example: pendulum (L = const)", "1.5")
    point(doc, "Non-holonomic: involve velocities, cannot be integrated. Example: rolling disk (v_contact = 0)", "1.5")
    point(doc, "DOF = n × f − c, where n = number of bodies, f = freedom per body, c = constraints", "1")
    point(doc, "Constraints reduce DOF: each independent constraint removes one degree of freedom", "1")
    sep(doc)

    # ---- Q1(b) ----
    heading(doc, "Q1 (b)  —  Position-Velocity-Acceleration [5 marks]")
    point(doc, "Given: x(t) = 4t³ − 2t; y(t) = 5t² + 3; at t = 2 s")
    point(doc, "r(2) = (4(8) − 4)î + (5(4) + 3)ĵ = 28î + 23ĵ m; |r| = √(784 + 529) = √1313 = 36.24 m", "1.5")
    point(doc, "v = (12t² − 2)î + 10t ĵ → v(2) = 46î + 20ĵ m/s; speed = √(2116 + 400) = √2516 = 50.16 m/s", "1.5")
    point(doc, "a = 24t î + 10 ĵ → a(2) = 48î + 10ĵ m/s²; |a| = √(2304 + 100) = √2404 = 49.03 m/s²", "2")
    sep(doc)

    # ---- Q1(c) ----
    heading(doc, "Q1 (c)  —  Block on 40° Incline [5 marks]")
    point(doc, "Given: m = 2 kg, θ = 40°, μ_k = 0.3, from rest")
    point(doc, "N = mg cos 40° = 2(9.81)(0.766) = 15.03 N", "1.5")
    point(doc, "f_k = μ_k N = 0.3(15.03) = 4.51 N", "1.5")
    point(doc, "a = g(sin 40° − μ_k cos 40°) = 9.81(0.6428 − 0.2298) = 9.81 × 0.4130 = 4.05 m/s²", "2")
    sep(doc)

    # ---- Q1(d) ----
    heading(doc, "Q1 (d)  —  Spinning Disk [5 marks]")
    point(doc, "Given: m = 5 kg, R = 0.2 m, α = 3.75 rad/s², starts from rest")
    point(doc, "I = ½mR² = ½(5)(0.04) = 0.1 kg·m²", "1")
    point(doc, "ω(2) = αt = 3.75 × 2 = 7.5 rad/s", "1")
    point(doc, "At rim: a_t = αR = 3.75 × 0.2 = 0.75 m/s²", "1")
    point(doc, "a_n = ω²R = 56.25 × 0.2 = 11.25 m/s²", "1")
    point(doc, "|a_total| = √(0.75² + 11.25²) = √(0.5625 + 126.5625) = √127.125 = 11.27 m/s²", "1")
    sep(doc)

    # ---- Q2 ----
    heading(doc, "Q2  —  Constrained Motion & Pulleys [20 marks]")

    heading(doc, "Q2 (a)  —  Vertical Loop R = 8 m [5 marks]", 2)
    point(doc, "Given: R = 8 m, m = 2 kg, v_bottom = 22 m/s")
    point(doc, "v_bottom_min = √(5gR) = √(5 × 9.81 × 8) = √392.4 = 19.81 m/s", "1.5")
    point(doc, "At top: ½mv²_bottom = ½mv²_top + 2mgR → v²_top = 484 − 4(9.81)(8) = 484 − 313.92 = 170.08", "1.5")
    point(doc, "N_top = mv²_top/R − mg = 2(170.08)/8 − 2(9.81) = 42.52 − 19.62 = 22.90 N", "2")

    heading(doc, "Q2 (b)  —  6-Pulley System [5 marks]", 2)
    point(doc, "Given: 6 pulleys, η = 0.92 each, m = 2000 kg, h = 3 m")
    point(doc, "MA = 6; η_total = 0.92⁶ = 0.6065", "1")
    point(doc, "F_ideal = 2000(9.81)/6 = 3270.0 N; F_actual = 3270.0/0.6065 = 5392.4 N", "1.5")
    point(doc, "W_out = 2000(9.81)(3) = 58860 J", "0.5")
    point(doc, "W_in = F_actual × (6 × 3) = 5392.4 × 18 = 97063.2 J", "1")
    point(doc, "Heat = 97063.2 − 58860.0 = 38203.2 J ≈ 38.2 kJ", "1")

    heading(doc, "Q2 (c)  —  Atwood m₁ = 7, m₂ = 5 [5 marks]", 2)
    point(doc, "a = (7 − 5)(9.81)/(7 + 5) = 2(9.81)/12 = 1.635 m/s²", "1.5")
    point(doc, "T = 2m₁m₂g/(m₁ + m₂) = 2(7)(5)(9.81)/12 = 57.225 N", "1.5")
    point(doc, "v(2 s) = 1.635 × 2 = 3.27 m/s", "1")
    point(doc, "Check: m₂g = 49.05 < T = 57.23 < m₁g = 68.67 ✓", "1")

    heading(doc, "Q2 (d)  —  CNC 3D Path [5 marks]", 2)
    point(doc, "Given: P₀ = (10, 20, 5), P₁ = (110, 70, 35) mm, F = 50 mm/s")
    point(doc, "ΔP = (100, 50, 30); d = √(10000 + 2500 + 900) = √13400 = 115.76 mm", "1.5")
    point(doc, "û = (100, 50, 30)/115.76 = (0.864, 0.432, 0.259)", "1")
    point(doc, "Axis speeds: v_x = 43.20, v_y = 21.60, v_z = 12.96 mm/s", "1.5")
    point(doc, "t = 115.76/50 = 2.315 s", "1")
    sep(doc)

    # ---- Q3 ----
    heading(doc, "Q3  —  Vibrations & Rolling [20 marks]")

    heading(doc, "Q3 (a)  —  Log Decrement [5 marks]", 2)
    point(doc, "Given: x₁ = 20 mm, x₂ = 15 mm, T_d = 0.35 s, m = 2 kg")
    point(doc, "δ = ln(x₁/x₂) = ln(20/15) = ln(1.333) = 0.2877", "1")
    point(doc, "ζ = δ/√(4π² + δ²) = 0.2877/√(39.478 + 0.0828) = 0.2877/6.290 = 0.0457", "1.5")
    point(doc, "ω_d = 2π/T_d = 2π/0.35 = 17.952 rad/s", "0.5")
    point(doc, "ω_n = ω_d/√(1 − ζ²) = 17.952/0.9990 = 17.971 rad/s → f_n = 2.860 Hz", "1")
    point(doc, "k = mω_n² = 2(17.971²) = 2(322.96) = 645.9 N/m", "1")

    heading(doc, "Q3 (b)  —  Rolling Wheel Acceleration [5 marks]", 2)
    point(doc, "Given: R = 0.4 m, v_center = 3.0 m/s, a_center = 1.5 m/s²")
    point(doc, "ω = v/R = 3.0/0.4 = 7.5 rad/s; α = a/R = 1.5/0.4 = 3.75 rad/s²", "1")
    point(doc, "Top point: a_tang = αR = 1.5 m/s² (backward); a_cent = ω²R = 22.5 m/s² (toward center)", "1.5")
    point(doc, "a_top = a_center + a_tang + a_cent: net = (1.5 − 1.5)î + (−22.5)ĵ → |a_top| = 22.5 m/s² (downward)", "1")
    point(doc, "Bottom point: a_tang = 1.5 m/s² (forward); a_cent = 22.5 m/s² (upward)", "0.5")
    point(doc, "a_bottom = (1.5 + 1.5)î + 22.5ĵ = 3.0î + 22.5ĵ → |a_bottom| = √(9 + 506.25) = 22.70 m/s²", "1")

    heading(doc, "Q3 (c)  —  Spinning Disk with Viscous Friction [5 marks]", 2)
    point(doc, "Given: m = 5 kg, R = 0.3 m, τ_applied = 2 N·m, c = 0.5 N·m·s/rad")
    point(doc, "I = ½mR² = ½(5)(0.09) = 0.225 kg·m²", "1.5")
    point(doc, "Time constant τ_d = I/c = 0.225/0.5 = 0.45 s", "1.5")
    point(doc, "Steady-state: Iα = 0 → τ = cω_ss → ω_ss = 2/0.5 = 4.0 rad/s", "2")

    heading(doc, "Q3 (d)  —  Gear Train [5 marks]", 2)
    point(doc, "Given: Stage 1: 15T → 45T; Stage 2: 12T → 60T; Motor: 1500 RPM, 0.5 N·m")
    point(doc, "GR₁ = 45/15 = 3; GR₂ = 60/12 = 5; Total GR = 15", "1")
    point(doc, "ω_out = 1500/15 = 100 RPM = 10.47 rad/s", "1")
    point(doc, "τ_out = 0.5 × 15 = 7.5 N·m (ideal)", "1")
    point(doc, "P_in = 0.5 × 2π(1500)/60 = 0.5 × 157.08 = 78.54 W", "1")
    point(doc, "P_out = τ_out × ω_out = 7.5 × 10.47 = 78.54 W → Power conserved ✓", "1")
    sep(doc)

    # ---- Q4 ----
    heading(doc, "Q4  —  Rigid Body Kinematics [20 marks]")

    heading(doc, "Q4 (a)  —  Rigid Body Velocity Formula [5 marks]", 2)
    point(doc, "v_P = v_center + ω × r_P/C", "1")
    point(doc, "v_center: translational velocity of the body's center of mass", "1")
    point(doc, "ω × r: rotational contribution — velocity due to angular velocity", "1")
    point(doc, "r is position vector from center of mass to point P", "0.5")
    point(doc, "For rolling: v_contact = v_center + ω × (−R ĵ) = v_center − ωR = 0 → v_center = ωR", "1.5")

    heading(doc, "Q4 (b)  —  Rotation Matrices & Gimbal Lock [5 marks]", 2)
    point(doc, "R_x(α): rotation about x-axis — [[1,0,0],[0,cosα,−sinα],[0,sinα,cosα]]", "1")
    point(doc, "R_y(β): rotation about y-axis — [[cosβ,0,sinβ],[0,1,0],[−sinβ,0,cosβ]]", "0.5")
    point(doc, "R_z(γ): rotation about z-axis — [[cosγ,−sinγ,0],[sinγ,cosγ,0],[0,0,1]]", "0.5")
    point(doc, "Gimbal lock: at β = ±90°, the first and third axes align → loss of one DOF", "1.5")
    point(doc, "Quaternions: q = (w, x, y, z) with |q| = 1; interpolation is smooth, no singularity", "1.5")

    heading(doc, "Q4 (c)  —  Rolling Wheel Velocities [5 marks]", 2)
    point(doc, "Given: R = 0.5 m, v_center = 2.0 m/s (right), rolling without slipping")
    point(doc, "ω = v/R = 2.0/0.5 = 4.0 rad/s", "1")
    point(doc, "Top: v = v_center + ωR = 2.0 + 2.0 = 4.0 m/s (rightward)", "1")
    point(doc, "Bottom (contact): v = v_center − ωR = 2.0 − 2.0 = 0 m/s", "1")
    point(doc, "Right: v = v_center(→) + ωR(↓) = 2.0î − 2.0ĵ; |v| = 2√2 = 2.83 m/s at 45° below horizontal", "2")

    heading(doc, "Q4 (d)  —  Rotating Rod [5 marks]", 2)
    point(doc, "Given: L = 1.0 m, ω = 5 rad/s, α = 2 rad/s², pivot at one end")
    point(doc, "Free end (r = 1.0): v = ωr = 5 m/s; a_t = αr = 2 m/s²; a_n = ω²r = 25 m/s²; |a| = √(4 + 625) = 25.08 m/s²", "2")
    point(doc, "Midpoint (r = 0.5): v = 2.5 m/s; a_t = 1.0; a_n = 12.5; |a| = √(1 + 156.25) = 12.54 m/s²", "1.5")
    point(doc, "r = 0.3: v = 1.5 m/s; a_t = 0.6; a_n = 7.5; |a| = √(0.36 + 56.25) = 7.52 m/s²", "1.5")
    sep(doc)

    # ---- Q5 ----
    heading(doc, "Q5  —  Energy & Momentum [20 marks]")

    heading(doc, "Q5 (a)  —  Work-Energy Theorem [5 marks]", 2)
    point(doc, "W_net = ΔKE = ½mv₂² − ½mv₁²", "1")
    point(doc, "Conservative force: path-independent work; has potential energy (gravity, spring)", "1")
    point(doc, "Non-conservative: path-dependent (friction, drag)", "1")
    point(doc, "Modified: W_conservative + W_non-conservative = ΔKE → ΔKE + ΔPE = W_nc", "2")

    heading(doc, "Q5 (b)  —  Impulse-Momentum [5 marks]", 2)
    point(doc, "Impulse: J = ∫F dt = F_avg × Δt; Momentum: p = mv", "1")
    point(doc, "Impulse-momentum theorem: J = Δp = mv₂ − mv₁", "1")
    point(doc, "Coefficient of restitution: e = (v₂B − v₂A)/(v₁A − v₁B) for collision", "1")
    point(doc, "Conservation: if ΣF_ext = 0, then Σm_i v_i = constant (isolated system)", "2")

    heading(doc, "Q5 (c)  —  Car Acceleration [5 marks]", 2)
    point(doc, "Given: m = 1200 kg, v₁ = 20 m/s, v₂ = 40 m/s, t = 8 s, F_drag = 400 N")
    point(doc, "a = (40 − 20)/8 = 2.5 m/s²; F_net = ma = 1200 × 2.5 = 3000 N", "1")
    point(doc, "F_engine = F_net + F_drag = 3000 + 400 = 3400 N", "1")
    point(doc, "W_engine = F_engine × 0.5(v₁ + v₂) × t = 3400 × 30 × 8 = 816000 J", "0.5")
    point(doc, "P_avg = W/t = 816000/8 = 102000 W = 102 kW", "1")
    point(doc, "P_inst(40 m/s) = F_engine × v = 3400 × 40 = 136000 W = 136 kW", "1.5")
    note(doc, "Alternative: P_avg = F_engine × v_avg = 3400 × 30 = 102 kW ✓")

    heading(doc, "Q5 (d)  —  Collision [5 marks]", 2)
    point(doc, "Given: m_A = m_B = 2 kg, v_A1 = 6 m/s, v_B1 = 0, e = 0.5")
    point(doc, "Momentum: 2(6) = 2v_A2 + 2v_B2 → v_A2 + v_B2 = 6", "1")
    point(doc, "Restitution: e = (v_B2 − v_A2)/(v_A1 − v_B1) → 0.5 = (v_B2 − v_A2)/6 → v_B2 − v_A2 = 3", "1")
    point(doc, "Solving: v_A2 = 1.5 m/s, v_B2 = 4.5 m/s", "1")
    point(doc, "Impulse: J = m_B(v_B2) = 2(4.5) = 9 N·s", "1")
    point(doc, "KE_loss = ½(2)(36) − ½(2)(2.25 + 20.25) = 36 − 22.5 = 13.5 J", "1")
    sep(doc)

    # ---- Q6 ----
    heading(doc, "Q6  —  FBD + Difficult Numerical [20 marks]")

    heading(doc, "Q6 (a)  —  Mobile Manipulator [10 marks]", 2)
    point(doc, "Given: v_AGV = 0.5 m/s, L = 1.2 m, θ = 30° from vertical, dθ/dt = −0.2 rad/s, m = 2.5 kg")
    marks_note(doc, "FBD [4 marks]:")
    point(doc, "Weight W = 2.5(9.81) = 24.525 N (downward)", "1")
    point(doc, "Arm reaction force at pivot (internal force balancing)", "1")
    point(doc, "Centripetal and Coriolis forces in rotating frame", "2")
    marks_note(doc, "Velocity Calculation [3 marks]:")
    point(doc, "Tip position relative to AGV: x_tip = L sin θ = 1.2 sin 30° = 0.6 m, y_tip = L cos θ = 1.2 cos 30° = 1.039 m", "1")
    point(doc, "v_tip_x = v_AGV + L cos θ × dθ/dt = 0.5 + 1.2(0.866)(−0.2) = 0.5 − 0.2078 = 0.292 m/s", "1")
    point(doc, "v_tip_y = −L sin θ × dθ/dt = −1.2(0.5)(−0.2) = +0.12 m/s", "0.5")
    point(doc, "|v_tip| = √(0.292² + 0.12²) = √(0.0853 + 0.0144) = √0.0997 = 0.316 m/s", "0.5")
    marks_note(doc, "Vibration Isolator [3 marks]:")
    point(doc, "ω_n = √(k/m) = √(1000/2.5) = √400 = 20 rad/s → f_n = 3.18 Hz", "1.5")
    point(doc, "c_cr = 2√(km) = 2√(1000 × 2.5) = 2√2500 = 100 N·s/m", "0.5")
    point(doc, "ζ = c/c_cr = 40/100 = 0.4 → underdamped", "1")

    heading(doc, "Q6 (b)  —  CNC Trapezoidal Profile [10 marks]", 2)
    point(doc, "Given: d = 200 mm, v_max = 100 mm/s, a = 500 mm/s²")
    marks_note(doc, "200 mm move (trapezoidal) [5 marks]:")
    point(doc, "t_acc = v_max/a = 100/500 = 0.2 s; d_acc = ½ × a × t² = ½(500)(0.04) = 10 mm", "1.5")
    point(doc, "d_cruise = 200 − 2(10) = 180 mm; t_cruise = 180/100 = 1.8 s", "1.5")
    point(doc, "t_total = 0.2 + 1.8 + 0.2 = 2.2 s", "2")
    marks_note(doc, "15 mm move (triangular) [5 marks]:")
    point(doc, "Check: d_min_trap = 2 × v_max²/(2a) = v_max²/a = 10000/500 = 20 mm > 15 mm → triangular", "1.5")
    point(doc, "v_peak = √(a × d/2) = √(500 × 7.5) = √3750 = 61.24 mm/s", "1.5")
    point(doc, "t_acc = v_peak/a = 61.24/500 = 0.1225 s; t_total = 2 × 0.1225 = 0.245 s", "2")
    sep(doc)

    # ---- Q7 ----
    heading(doc, "Q7  —  FBD + Difficult Numerical [20 marks]")

    heading(doc, "Q7 (a)  —  Spring-Mass-Damper Full Analysis [10 marks]", 2)
    point(doc, "Given: m = 2 kg, k = 500 N/m, c = 20 N·s/m, x₀ = 0.1 m, v₀ = 0")
    marks_note(doc, "FBD [3 marks]:")
    point(doc, "In displaced position: Spring force kx (restoring, toward equilibrium)", "1")
    point(doc, "Damping force cẋ (opposing velocity)", "1")
    point(doc, "Weight and normal cancel about equilibrium point", "1")
    marks_note(doc, "EOM and classification [4 marks]:")
    point(doc, "mẍ + cẋ + kx = 0 → 2ẍ + 20ẋ + 500x = 0 → ẍ + 10ẋ + 250x = 0", "1")
    point(doc, "ω_n = √(250) = 15.81 rad/s; c_cr = 2√(km) = 63.25 N·s/m", "1")
    point(doc, "ζ = 20/63.25 = 0.316 → Underdamped", "1")
    point(doc, "ω_d = 15.81√(1 − 0.0999) = 15.81 × 0.9487 = 15.00 rad/s", "1")
    marks_note(doc, "Free response [3 marks]:")
    point(doc, "x(t) = e^(−5t)[A cos(15t) + B sin(15t)]", "1")
    point(doc, "IC: x(0) = 0.1 → A = 0.1; ẋ(0) = 0 → B = (ζω_n/ω_d) × A = (5/15)(0.1) = 0.0333", "1")
    point(doc, "x(t) = e^(−5t)[0.1 cos(15t) + 0.0333 sin(15t)] m", "1")

    heading(doc, "Q7 (b)  —  Double Pendulum [10 marks]", 2)
    point(doc, "Given: m₁ = m₂ = 1 kg, L₁ = L₂ = 1 m")
    marks_note(doc, "Kinetic and Potential Energy [4 marks]:")
    point(doc, "T = ½m₁L₁²θ̇₁² + ½m₂[L₁²θ̇₁² + L₂²θ̇₂² + 2L₁L₂θ̇₁θ̇₂ cos(θ₁ − θ₂)]", "2")
    point(doc, "V = −m₁gL₁ cos θ₁ − m₂g(L₁ cos θ₁ + L₂ cos θ₂)", "2")
    marks_note(doc, "Euler-Lagrange equations [4 marks]:")
    point(doc, "For θ₁: (m₁ + m₂)L₁²θ̈₁ + m₂L₁L₂θ̈₂ cos(θ₁ − θ₂) + m₂L₁L₂θ̇₂² sin(θ₁ − θ₂) + (m₁ + m₂)gL₁ sin θ₁ = 0", "2")
    point(doc, "For θ₂: m₂L₂²θ̈₂ + m₂L₁L₂θ̈₁ cos(θ₁ − θ₂) − m₂L₁L₂θ̇₁² sin(θ₁ − θ₂) + m₂gL₂ sin θ₂ = 0", "2")
    marks_note(doc, "Small-angle linearization and chaos [2 marks]:")
    point(doc, "Linearized form (sin θ ≈ θ, cos Δθ ≈ 1): 2θ̈₁ + θ̈₂ + 2g θ₁ = 0; θ̈₂ + θ̈₁ + g θ₂ = 0", "1")
    point(doc, "The double pendulum is a canonical example of deterministic chaos: sensitive to initial "
          "conditions, large-angle motion appears random despite deterministic equations", "1")
    sep(doc)


# ═══════════════ SET 3 ANSWERS ═══════════════

def answers_set3(doc):
    # ---- Q1(a) ----
    heading(doc, "Q1 (a)  —  Damping Regimes [5 marks]")
    point(doc, "Underdamped (ζ < 1): Oscillates with exponentially decaying amplitude; ω_d = ω_n√(1 − ζ²)", "1.5")
    point(doc, "Critically damped (ζ = 1): Returns to equilibrium fastest without oscillation; x = (A + Bt)e^(−ω_n t)", "1.5")
    point(doc, "Overdamped (ζ > 1): Slow return, no oscillation; two real negative roots; sluggish response", "1")
    point(doc, "Automotive: slightly underdamped (ζ ≈ 0.2–0.4) — fast return with minimal overshoot (ride quality + control)", "1")
    sep(doc)

    # ---- Q1(b) ----
    heading(doc, "Q1 (b)  —  Rocket Kinematics [5 marks]")
    point(doc, "Given: y(t) = 60t − 10t²")
    point(doc, "v(t) = dy/dt = 60 − 20t → v₀ = v(0) = 60 m/s", "1")
    point(doc, "a = dv/dt = −20 m/s² (constant deceleration)", "1")
    point(doc, "At max height: v = 0 → 60 − 20t = 0 → t = 3 s", "1.5")
    point(doc, "y_max = 60(3) − 10(9) = 180 − 90 = 90 m", "1.5")
    sep(doc)

    # ---- Q1(c) ----
    heading(doc, "Q1 (c)  —  Atwood 6 kg and 4 kg [5 marks]")
    point(doc, "Given: m₁ = 6 kg, m₂ = 4 kg, g = 9.81 m/s²")
    point(doc, "a = (6 − 4)(9.81)/(6 + 4) = 2(9.81)/10 = 1.962 m/s²", "1.5")
    point(doc, "T = 2(6)(4)(9.81)/10 = 470.88/10 = 47.09 N", "1.5")
    point(doc, "v(1 s) = 1.962 × 1 = 1.962 m/s", "1")
    point(doc, "Check: m₂g = 39.24 < T = 47.09 < m₁g = 58.86 ✓", "1")
    sep(doc)

    # ---- Q1(d) ----
    heading(doc, "Q1 (d)  —  Coordinate Conversions [5 marks]")
    point(doc, "(6, 8) → polar: r = √(36 + 64) = √100 = 10; θ = arctan(8/6) = 53.13°", "2")
    point(doc, "(ρ = 3, φ = 45°, z = 4) → Cartesian: x = 3 cos 45° = 2.121 m, y = 3 sin 45° = 2.121 m, z = 4 m", "2")
    point(doc, "Verification: |r| = √(2.121² + 2.121² + 4²) = √(4.5 + 4.5 + 16) = √25 = 5 m; also √(ρ² + z²) = √(9 + 16) = 5 ✓", "1")
    sep(doc)

    # ---- Q2 ----
    heading(doc, "Q2  —  Incline & Pulley Systems [20 marks]")

    heading(doc, "Q2 (a)  —  Block Up Rough Incline [5 marks]", 2)
    point(doc, "Given: m = 5 kg, θ = 30°, μ_k = 0.2, μ_s = 0.3, v₀ = 10 m/s up")
    point(doc, "Going up (friction opposes = acts down): a = −g(sin 30° + μ_k cos 30°) = −9.81(0.5 + 0.1732) = −6.604 m/s²", "1")
    point(doc, "d = v₀²/(2 × 6.604) = 100/13.208 = 7.57 m", "1")
    point(doc, "t = v₀/6.604 = 10/6.604 = 1.514 s", "1")
    point(doc, "Slide back check: g sin 30° = 4.905 N/kg vs μ_s g cos 30° = 0.3 × 8.496 = 2.549 N/kg", "1")
    point(doc, "4.905 > 2.549 → block slides back; a_down = g(sin 30° − μ_k cos 30°) = 9.81 × 0.3268 = 3.206 m/s²", "1")

    heading(doc, "Q2 (b)  —  Elevator-Counterweight [5 marks]", 2)
    point(doc, "Given: m_car = 1200 kg, m_cw = 1000 kg, 4 passengers × 75 = 300 kg")
    point(doc, "Total car side: 1200 + 300 = 1500 kg", "1")
    point(doc, "a = (m_car_side − m_cw)g/(m_car_side + m_cw) = (1500 − 1000)(9.81)/(1500 + 1000) = 500(9.81)/2500 = 1.962 m/s² (car descends)", "2")
    point(doc, "T = 2(1500)(1000)(9.81)/2500 = 29430000/2500 = 11772 N", "2")
    note(doc, "This is an Atwood-type problem. Car side is heavier so it descends.")

    heading(doc, "Q2 (c)  —  Frictionless 30° Incline [5 marks]", 2)
    point(doc, "Given: m = 1 kg, θ = 30°, frictionless, from rest")
    point(doc, "N = mg cos 30° = 1(9.81)(0.866) = 8.496 N", "1")
    point(doc, "a = g sin 30° = 9.81(0.5) = 4.905 m/s²", "1.5")
    point(doc, "v(3) = at = 4.905 × 3 = 14.72 m/s", "1")
    point(doc, "s = ½at² = ½(4.905)(9) = 22.07 m", "1.5")

    heading(doc, "Q2 (d)  —  CNC Trapezoidal/Triangular Profile [5 marks]", 2)
    point(doc, "Given: d = 250 mm, v_max = 150 mm/s, a = 1000 mm/s²")
    point(doc, "t_acc = 150/1000 = 0.15 s; d_acc = ½(1000)(0.0225) = 11.25 mm; 2 × d_acc = 22.5 mm < 250 → trapezoidal", "1")
    point(doc, "d_cruise = 250 − 22.5 = 227.5 mm; t_cruise = 227.5/150 = 1.517 s", "1.5")
    point(doc, "t_total = 0.15 + 1.517 + 0.15 = 1.817 s", "0.5")
    point(doc, "For 10 mm: d_min_trap = v_max²/a = 22.5 mm > 10 → triangular", "0.5")
    point(doc, "v_peak = √(a × d/2) = √(1000 × 5) = √5000 = 70.71 mm/s; t_total = 2(70.71/1000) = 0.1414 s", "1.5")
    sep(doc)

    # ---- Q3 ----
    heading(doc, "Q3  —  Rigid Body & Energy [20 marks]")

    heading(doc, "Q3 (a)  —  Rolling Disk on 20° Incline [5 marks]", 2)
    point(doc, "Given: m = 10 kg, R = 0.15 m, θ = 20°, from rest, s = 3 m")
    point(doc, "Rolling disk: β = 1/2 → a = g sin θ/(1 + 0.5) = 9.81 sin 20°/1.5 = 9.81(0.342)/1.5 = 2.236 m/s²", "1.5")
    point(doc, "v = √(2as) = √(2 × 2.236 × 3) = √13.42 = 3.663 m/s", "1")
    point(doc, "t = √(2s/a) = √(6/2.236) = √2.683 = 1.638 s", "1")
    point(doc, "Frictionless sliding: a = g sin 20° = 3.355 m/s² (50% faster — no rotational KE)", "1.5")

    heading(doc, "Q3 (b)  —  Pendulum Released from 60° [5 marks]", 2)
    point(doc, "Given: m = 2 kg, L = 1.5 m, θ₀ = 60°")
    point(doc, "Height: h = L(1 − cos 60°) = 1.5(1 − 0.5) = 0.75 m", "1")
    point(doc, "mgh = ½mv² → v = √(2gh) = √(2 × 9.81 × 0.75) = √14.715 = 3.836 m/s", "2")
    point(doc, "At bottom: T − mg = mv²/L → T = mg + mv²/L = 2(9.81) + 2(14.715)/1.5 = 19.62 + 19.62 = 39.24 N", "2")

    heading(doc, "Q3 (c)  —  Car Braking [5 marks]", 2)
    point(doc, "Given: m = 1500 kg, v = 100 km/h = 27.78 m/s, F_brake = 8000 N")
    point(doc, "KE = ½mv² = ½(1500)(771.73) = 578796 J ≈ 578.8 kJ", "1.5")
    point(doc, "W = F × d = KE → d = 578796/8000 = 72.35 m", "2")
    point(doc, "All KE is dissipated as heat = 578.8 kJ", "1.5")

    heading(doc, "Q3 (d)  —  Damped Vibration Parameters [5 marks]", 2)
    point(doc, "Given: m = 2 kg, k = 800 N/m, ζ = 0.2")
    point(doc, "ω_n = √(k/m) = √(800/2) = √400 = 20 rad/s", "1")
    point(doc, "ω_d = ω_n√(1 − ζ²) = 20√(1 − 0.04) = 20 × 0.9798 = 19.60 rad/s", "1")
    point(doc, "c_cr = 2mω_n = 2(2)(20) = 80 N·s/m", "1")
    point(doc, "c = ζ × c_cr = 0.2 × 80 = 16 N·s/m", "1")
    point(doc, "Decay to half: e^(−ζω_n nT_d) = 0.5 → n = ln 2/(ζω_n T_d) = 0.6931/(0.2 × 20 × 2π/19.60) = 0.6931/(4 × 0.3206) = 0.6931/1.2824 = 0.54 → ~1 cycle (amplitude halves in about 1 full oscillation)", "1")
    note(doc, "More precisely: per cycle decay = e^(−ζω_n T_d) = e^(−2πζ/√(1−ζ²)) = e^(−1.283) = 0.277, so < 1 cycle needed to halve")
    sep(doc)

    # ---- Q4 ----
    heading(doc, "Q4  —  Rigid Body Properties [20 marks]")

    heading(doc, "Q4 (a)  —  MOI & Parallel Axis Theorem [5 marks]", 2)
    point(doc, "MOI: I = ∫r² dm — measures resistance to angular acceleration", "1")
    point(doc, "Parallel Axis: I_O = I_cm + md², where d = distance from CM to new axis", "1")
    point(doc, "Rod about center: I = ∫_{−L/2}^{L/2} x² (m/L) dx = (m/L)[x³/3]_{−L/2}^{L/2} = mL²/12", "1.5")
    point(doc, "About end (d = L/2): I = mL²/12 + m(L/2)² = mL²/12 + mL²/4 = mL²/3", "1.5")

    heading(doc, "Q4 (b)  —  Slider-Crank Theory [5 marks]", 2)
    point(doc, "Components: crank (radius r, angle θ), connecting rod (length L), slider (piston)")
    point(doc, "Constraint: L sin φ = r sin θ → φ = arcsin[(r/L) sin θ]", "1.5")
    point(doc, "Piston position: x = r cos θ + L cos φ", "1")
    point(doc, "Asymmetric motion: piston spends more time near TDC than BDC due to rod geometry", "1")
    point(doc, "Higher L/r → more symmetric motion, less side thrust, smoother operation", "1.5")

    heading(doc, "Q4 (c)  —  MOI Calculations [5 marks]", 2)
    point(doc, "Given: m = 5 kg for each shape")
    point(doc, "Solid disk (R = 0.2): I = ½mR² = ½(5)(0.04) = 0.1 kg·m²", "1")
    point(doc, "Rod (L = 1) about center: I = mL²/12 = 5(1)/12 = 0.417 kg·m²", "1")
    point(doc, "Rod about end: I = mL²/3 = 5/3 = 1.667 kg·m²", "1")
    point(doc, "Solid sphere (R = 0.2): I = 2mR²/5 = 2(5)(0.04)/5 = 0.08 kg·m²", "1")
    point(doc, "PAT check for rod: 0.417 + 5(0.5²) = 0.417 + 1.25 = 1.667 ✓", "1")

    heading(doc, "Q4 (d)  —  Slider-Crank at 6000 RPM, θ = 90° [5 marks]", 2)
    point(doc, "Given: r = 50 mm, L = 150 mm, n = 6000 RPM, θ = 90°")
    point(doc, "ω = 2π(6000)/60 = 628.32 rad/s", "1")
    point(doc, "sin φ = (r/L) sin 90° = 50/150 = 0.3333 → φ = 19.47°", "1")
    point(doc, "v_piston = −rω sin θ − Lφ̇ sin φ where φ̇ = rω cos θ/(L cos φ) = 50(628.32)(0)/(150 × 0.9428) = 0", "1")
    point(doc, "v_piston = −50(628.32)(1) − 0 = −31416 mm/s = −31.42 m/s", "0.5")
    point(doc, "At TDC (θ = 0): a_piston ≈ rω²(1 + r/L) = 0.05(628.32²)(1 + 1/3) = 0.05(394826)(1.333) = 26322 m/s² = 2683 g", "1")
    point(doc, "2683 g >> 500 g → YES, acceleration exceeds 500 g by a large margin at 6000 RPM", "0.5")
    sep(doc)

    # ---- Q5 ----
    heading(doc, "Q5  —  Lagrangian & Advanced [20 marks]")

    heading(doc, "Q5 (a)  —  Simple Pendulum via Lagrange [5 marks]", 2)
    point(doc, "Generalized coordinate: θ (angle from vertical)", "0.5")
    point(doc, "T = ½m(Lθ̇)² = ½mL²θ̇²; V = −mgL cos θ (datum at pivot)", "1.5")
    point(doc, "L = T − V = ½mL²θ̇² + mgL cos θ", "1")
    point(doc, "d/dt(∂L/∂θ̇) = mL²θ̈; ∂L/∂θ = −mgL sin θ", "1")
    point(doc, "EOM: mL²θ̈ + mgL sin θ = 0 → θ̈ + (g/L) sin θ = 0 — same as Newton (ma_t = −mg sin θ → Lα = −g sin θ)", "1")

    heading(doc, "Q5 (b)  —  Normal Modes & Modal Analysis [5 marks]", 2)
    point(doc, "For n-DOF: [M]ẍ + [K]x = 0 → assume x = X e^(iωt)", "1")
    point(doc, "Eigenvalue problem: [K]X = ω²[M]X → det([K] − ω²[M]) = 0", "1.5")
    point(doc, "Solving gives n natural frequencies ω₁, ω₂, ..., ω_n", "0.5")
    point(doc, "Mode shapes: eigenvectors X_i describe the relative motion of each DOF at that frequency", "1")
    point(doc, "Physical meaning: at a natural frequency, all parts of the system oscillate in a fixed pattern (mode shape)", "1")

    heading(doc, "Q5 (c)  —  Simple vs Physical Pendulum [5 marks]", 2)
    point(doc, "Given: m = 2 kg, L = 1.5 m")
    point(doc, "Simple: L = ½mL²θ̇² + mgL cos θ; EOM: θ̈ + (g/L) sin θ = 0", "1")
    point(doc, "ω_n = √(g/L) = √(9.81/1.5) = √6.54 = 2.557 rad/s; T = 2π/2.557 = 2.457 s", "1.5")
    point(doc, "Physical pendulum (rod, same L, pivot at end):")
    point(doc, "I_O = mL²/3, d = L/2: ω_n = √(mgd/I) = √(2(9.81)(0.75)/(2(2.25)/3)) = √(14.715/1.5) = √9.81 = 3.132 rad/s", "1.5")
    point(doc, "T_phys = 2π/3.132 = 2.006 s → physical pendulum is faster (shorter period)", "1")

    heading(doc, "Q5 (d)  —  ODE Solvers Theory [5 marks]", 2)
    point(doc, "Forward Euler: y_{n+1} = y_n + h f(t_n, y_n); Order: O(h), 1st order; Simple but unstable for stiff systems", "1.5")
    point(doc, "RK2 (Midpoint/Heun): uses slope at midpoint; Order: O(h²); Better accuracy, moderate cost", "1.5")
    point(doc, "RK4 (Classical): 4 slope evaluations per step; Order: O(h⁴); Excellent balance of accuracy and efficiency", "1.5")
    point(doc, "RK4 preferred: 4th-order accuracy means doubling step size only reduces accuracy by factor 16; "
          "no Jacobian needed (unlike implicit methods); widely used in MATLAB ode45, PyBullet, etc.", "0.5")
    sep(doc)

    # ---- Q6 ----
    heading(doc, "Q6  —  FBD + Difficult Numerical [20 marks]")

    heading(doc, "Q6 (a)  —  Rotating Sensor FBD [10 marks]", 2)
    point(doc, "Given: ω = 2.0 rad/s, r = 0.5 m, ṙ = 0.3 m/s, m = 0.2 kg")
    marks_note(doc, "FBD in rotating frame [4 marks]:")
    point(doc, "Weight W = mg = 0.2(9.81) = 1.962 N (downward)", "1")
    point(doc, "Centrifugal force: mω²r = 0.2(4)(0.5) = 0.4 N (outward)", "1")
    point(doc, "Coriolis force: 2mωṙ = 2(0.2)(2.0)(0.3) = 0.24 N (perpendicular to radial motion)", "1")
    point(doc, "Constraint force from track/rail (provides centripetal and tangential forces)", "1")
    marks_note(doc, "Acceleration calculations [3 marks]:")
    point(doc, "Centripetal: a_c = ω²r = (2.0)²(0.5) = 2.0 m/s² (inward)", "1")
    point(doc, "Coriolis: a_Cor = 2ωṙ = 2(2.0)(0.3) = 1.2 m/s² (tangential, perpendicular to r)", "1")
    point(doc, "|a_total| = √(a_c² + a_Cor²) = √(4.0 + 1.44) = √5.44 = 2.33 m/s²", "1")
    marks_note(doc, "Constraint forces [3 marks]:")
    point(doc, "Radial constraint: F_r = m × a_centripetal = 0.2 × 2.0 = 0.4 N (inward) — maintains circular path", "1.5")
    point(doc, "Tangential constraint: F_t = m × a_Cor = 0.2 × 1.2 = 0.24 N — counters Coriolis", "1.5")

    heading(doc, "Q6 (b)  —  Incline then Horizontal [10 marks]", 2)
    point(doc, "Given: m = 3 kg, incline 35° μ_k = 0.25 (4 m), horizontal μ_k = 0.3, from rest")
    marks_note(doc, "On incline [5 marks]:")
    point(doc, "N₁ = mg cos 35° = 3(9.81)(0.8192) = 24.11 N", "0.5")
    point(doc, "h = 4 sin 35° = 4(0.5736) = 2.294 m", "0.5")
    point(doc, "W_gravity = mgh = 3(9.81)(2.294) = 67.52 J", "1")
    point(doc, "W_friction_incline = −μ_k N₁ s = −0.25(24.11)(4) = −24.11 J", "1")
    point(doc, "KE at bottom = 67.52 − 24.11 = 43.41 J; v = √(2 × 43.41/3) = √28.94 = 5.38 m/s", "2")
    marks_note(doc, "On horizontal [5 marks]:")
    point(doc, "N₂ = mg = 29.43 N; f₂ = 0.3(29.43) = 8.83 N", "1")
    point(doc, "½mv² = f₂ d → d = 43.41/8.83 = 4.92 m", "2")
    point(doc, "Total friction: 24.11 + 8.83 × 4.92 = 24.11 + 43.41 = 67.52 J", "1")
    point(doc, "Energy check: All PE (67.52 J) dissipated by friction ✓ (started and ended at rest)", "1")
    sep(doc)

    # ---- Q7 ----
    heading(doc, "Q7  —  FBD + Difficult Numerical [20 marks]")

    heading(doc, "Q7 (a)  —  Conveyor Transfer Energy Analysis [10 marks]", 2)
    point(doc, "Given: m = 2 kg, v₁ = 1.5 m/s, v₂ = 3.0 m/s, μ_s = 0.35, μ_k = 0.25")
    marks_note(doc, "FBD [3 marks]:")
    point(doc, "Weight: W = 2(9.81) = 19.62 N ↓; Normal: N = 19.62 N ↑", "1")
    point(doc, "Kinetic friction: f_k = 0.25(19.62) = 4.905 N → (forward, belt pulls package)", "1")
    point(doc, "Package is slower than belt: relative motion of package w.r.t. belt is backward → friction acts forward ✓", "1")
    marks_note(doc, "Kinematics [4 marks]:")
    point(doc, "a = f_k/m = 4.905/2 = 2.453 m/s²", "1")
    point(doc, "t = Δv/a = 1.5/2.453 = 0.612 s", "1")
    point(doc, "ΔKE = ½m(v₂² − v₁²) = ½(2)(9 − 2.25) = 6.75 J", "2")
    marks_note(doc, "Energy analysis [3 marks]:")
    point(doc, "Package displacement: s_p = v₁t + ½at² = 1.5(0.612) + ½(2.453)(0.374) = 0.918 + 0.459 = 1.377 m", "1")
    point(doc, "Belt displacement: s_b = v₂t = 3.0(0.612) = 1.836 m", "0.5")
    point(doc, "W_friction = f_k × s_b = 4.905 × 1.836 = 9.006 J (total work by friction on belt surface)", "0.5")
    note(doc, "Alternative: W_friction_on_package = f_k × s_p = 4.905 × 1.377 = 6.75 J = ΔKE ✓")
    point(doc, "Efficiency = ΔKE / W_total = 6.75 / 9.006 = 74.9% ≈ 75%", "1")

    heading(doc, "Q7 (b)  —  Compound Gear Train [10 marks]", 2)
    point(doc, "Given: Stage 1: 20T → 80T; Stage 2: 15T → 60T; n_in = 2000 RPM, τ_in = 2.0 N·m, η = 96%/stage")
    marks_note(doc, "Gear ratios [3 marks]:")
    point(doc, "GR₁ = 80/20 = 4; GR₂ = 60/15 = 4; Total GR = 4 × 4 = 16", "1")
    point(doc, "n_out = 2000/16 = 125 RPM = 13.09 rad/s", "1")
    point(doc, "η_total = 0.96 × 0.96 = 0.9216", "1")
    marks_note(doc, "Torques and power [4 marks]:")
    point(doc, "τ_out_ideal = τ_in × GR = 2.0 × 16 = 32.0 N·m", "1")
    point(doc, "τ_out_actual = τ_out_ideal × η_total = 32.0 × 0.9216 = 29.49 N·m", "1")
    point(doc, "P_in = τ_in × ω_in = 2.0 × 2π(2000)/60 = 2.0 × 209.44 = 418.88 W", "1")
    point(doc, "P_out = P_in × η_total = 418.88 × 0.9216 = 386.08 W; P_loss = 418.88 − 386.08 = 32.80 W", "1")
    marks_note(doc, "Load check [3 marks]:")
    point(doc, "Required: 25 N·m at 125 RPM → P_load = 25 × 13.09 = 327.25 W", "1")
    point(doc, "Available: τ_out = 29.49 N·m > 25 N·m ✓ and P_out = 386.08 W > 327.25 W ✓", "1")
    point(doc, "Motor IS sufficient with margin: torque surplus = 4.49 N·m (18%), power surplus = 58.83 W (15.2%)", "1")
    sep(doc)


# ═══════════════ MAIN ═══════════════

if __name__ == "__main__":
    os.makedirs("Final Papers", exist_ok=True)
    sets = {1: answers_set1, 2: answers_set2, 3: answers_set3}

    for num, writer in sets.items():
        doc = new_doc(num)
        writer(doc)
        path = f"Final Papers/DSM_Answers_Set{num}.docx"
        doc.save(path)
        print(f"  [OK] {path}")

    print("\nDone — 3 DSM answer booklets generated.")
