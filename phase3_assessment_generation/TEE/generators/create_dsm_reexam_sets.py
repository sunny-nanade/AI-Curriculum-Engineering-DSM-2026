"""
DSM Re-Examination Question Papers — Sets R1, R2
==================================================
Dynamic Systems Modeling (702MH0C025)
B.Tech Mechatronics, Semester IV
Old-semester / Re-exam students — same syllabus

Structure per set (identical to regular exam):
  Q1 (compulsory): 4 × 5 = 20 marks  (3 numericals + 1 theory)
  Q2–Q7: 20 marks each, solve any 4 = 80 marks
    Q2, Q3 — Pure Numericals (20 marks each)
    Q4, Q5 — Theory 10M + Numerical 10M
    Q6, Q7 — FBD 10M + Difficult Numerical 10M
  Total written: 140   Evaluated: 100

Questions are DIFFERENT from the 3 regular-exam sets.
Topics are jumbled across question slots.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Course Outcomes ──
COURSE_OUTCOMES = {
    "CO1": "Understand and apply theoretical concepts in kinematics",
    "CO2": "Analyze the given system using the Free Body Diagram",
    "CO3": "Understand and calculate the forces acting on dynamic systems",
    "CO4": "Apply the constitutive laws to analyze dynamic systems of particles and rigid bodies",
    "CO5": "Understand and calculate the energy of a dynamic system and power dissipated",
}

# ── Per-set synoptic rows (topics jumbled differently in each set) ──

MASTER_ROWS_R1 = [
    # (Q.No, Topic Area, Unit, CO, BL, Marks, Difficulty, Type)
    ("Q1 (a)", "Vibrations / Damping Theory",            "5",   "CO4",      "L2", "5",  "Easy",   "Theory"),
    ("Q1 (b)", "Rolling Rigid Body — Kinetic Energy",    "3,6", "CO1+CO5",  "L3", "5",  "Easy",   "Numerical"),
    ("Q1 (c)", "Energy / Spring Launch",                  "6",   "CO5",      "L3", "5",  "Medium", "Numerical"),
    ("Q1 (d)", "Constrained Two-Block System",            "4",   "CO3",      "L3", "5",  "Medium", "Numerical"),
    ("Q2",     "Rolling Dynamics & Energy",               "5,6", "CO4+CO5",  "L3", "20", "Medium", "Numerical"),
    ("Q3",     "Constrained Motion & Particle Kinetics",  "4",   "CO3+CO4",  "L3", "20", "Medium", "Numerical"),
    ("Q4",     "Lagrangian & Energy Methods",             "2,6", "CO4+CO5",  "L4", "20", "Medium", "Theory+Num"),
    ("Q5",     "Rotating Frames & Rigid Body Kinematics", "3",   "CO1+CO3",  "L4", "20", "Hard",   "Theory+Num"),
    ("Q6",     "FBD Analysis + Difficult Numerical",      "5,6", "CO2+CO5",  "L5", "20", "Hard",   "FBD+Num"),
    ("Q7",     "FBD Analysis + Difficult Numerical",      "4,5", "CO2+CO3",  "L5", "20", "Hard",   "FBD+Num"),
]

MASTER_ROWS_R2 = [
    ("Q1 (a)", "Newton's Laws — Theory",                  "1,4", "CO1",      "L2", "5",  "Easy",   "Theory"),
    ("Q1 (b)", "Moment of Inertia / PAT",                 "3",   "CO1",      "L3", "5",  "Easy",   "Numerical"),
    ("Q1 (c)", "Impulse, Momentum & Collisions",          "4",   "CO3",      "L3", "5",  "Medium", "Numerical"),
    ("Q1 (d)", "Friction & Block on Surface",              "4",   "CO3",      "L3", "5",  "Medium", "Numerical"),
    ("Q2",     "Rotating Machinery & Transmission",        "3,5", "CO3+CO4",  "L3", "20", "Medium", "Numerical"),
    ("Q3",     "Energy Methods & Circular Motion",         "5,6", "CO4+CO5",  "L3", "20", "Medium", "Numerical"),
    ("Q4",     "Rotating Frames & Rigid Body Properties",  "3,5", "CO1+CO3",  "L4", "20", "Medium", "Theory+Num"),
    ("Q5",     "Pendulums & Lagrangian Mechanics",         "2,5", "CO4+CO5",  "L4", "20", "Hard",   "Theory+Num"),
    ("Q6",     "FBD Analysis + Difficult Numerical",       "4,5", "CO2+CO5",  "L5", "20", "Hard",   "FBD+Num"),
    ("Q7",     "FBD Analysis + Difficult Numerical",       "3,4", "CO2+CO3",  "L5", "20", "Hard",   "FBD+Num"),
]

# ═══════════════ HELPER FUNCTIONS ═══════════════

def new_doc(set_number):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)
    for s in doc.sections:
        pgB = OxmlElement('w:pgBorders')
        pgB.set(qn('w:offsetFrom'), 'page')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '18')
            b.set(qn('w:space'), '24'); b.set(qn('w:color'), '000000')
            pgB.append(b)
        s._sectPr.append(pgB)

    def ctr(txt, sz=12, bold=True, color=None):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(0)
        r = p.add_run(txt); r.bold = bold; r.font.size = Pt(sz); r.font.name = 'Times New Roman'
        if color: r.font.color.rgb = color

    ctr("SVKM's NMIMS (Deemed-to-be University)", 14)
    ctr("School of Technology Management & Engineering (STME), Indore", 10)
    ctr("RE-EXAMINATION — MARCH 2026", 11)
    ctr(f"Dynamic Systems Modeling (702MH0C025) — SET R{set_number}", 13)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    for label, val in [("Program: ","B.Tech Mechatronics"), ("  Semester: ","IV"),
                       ("  Duration: ","3 Hours"), ("  Max Marks: ","100")]:
        r = p.add_run(label); r.bold = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'
        r = p.add_run(val); r.font.size = Pt(9); r.font.name = 'Times New Roman'

    line(doc)

    inst = doc.add_paragraph()
    inst.paragraph_format.space_after = Pt(4)
    r = inst.add_run("Instructions: "); r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'
    r = inst.add_run("Q1 is compulsory (20 marks). Attempt any FOUR from Q2–Q7 (4 × 20 = 80 marks). "
                     "Draw neat Free Body Diagrams wherever applicable. No programmable calculators. "
                     "All numerical answers must include proper units. "
                     "Assume g = 9.81 m/s² unless stated otherwise.")
    r.font.size = Pt(9); r.font.name = 'Times New Roman'

    line(doc)
    return doc


def line(doc):
    p = doc.add_paragraph("─" * 95)
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    for r in p.runs:
        r.font.size = Pt(6); r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0x99,0x99,0x99)


def q_head(doc, number, marks, text=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"Q{number}  [{marks} marks]"); r.bold = True; r.font.size = Pt(11); r.font.name = 'Times New Roman'
    if text:
        r = p.add_run(f"  —  {text}"); r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0x33,0x33,0x33)


def sub_q(doc, label, text, marks=5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(f"({label})  "); r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'
    r = p.add_run(text); r.font.size = Pt(10); r.font.name = 'Times New Roman'
    r = p.add_run(f"  [{marks}]"); r.bold = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0x00,0x00,0x99)


def big_q(doc, number, q_type, intro, subs):
    q_head(doc, number, 20, q_type)
    if intro:
        p = doc.add_paragraph(intro)
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(10); r.font.name = 'Times New Roman'; r.italic = True
    for label, text, marks in subs:
        sub_q(doc, label, text, marks)
    line(doc)


def add_synoptic_table(doc, rows):
    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SYNOPTIC TABLE — QUESTION PAPER MAPPING"); r.bold = True
    r.font.size = Pt(12); r.font.name = 'Times New Roman'

    headers = ["Q. No.", "Topic Area", "Unit", "CO", "BL", "Marks", "Difficulty", "Type"]
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]; cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(8); r.font.name = 'Times New Roman'
    for ri, row_data in enumerate(rows, 1):
        for ci, val in enumerate(row_data):
            cell = tbl.rows[ri].cells[ci]; cell.text = val
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(8); r.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Coverage Summary\n"); r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'
    r = p.add_run(
        "Units: 1, 2, 3, 4, 5, 6 (all theory units covered; Unit 7 has no numerical)\n"
        "COs: CO1–CO5 (all Course Outcomes assessed)\n"
        "Bloom's Levels: L2–L5 (Understand → Evaluate)\n"
        "Difficulty: Easy → Medium → Hard (50 marks easy, 80+ difficult)\n"
        "Q2–Q7: 2 Pure Numerical + 2 Theory+Numerical + 2 FBD+Difficult Numerical"
    )
    r.font.size = Pt(9); r.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Course Outcomes:\n"); r.bold = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'
    for co, desc in COURSE_OUTCOMES.items():
        r = p.add_run(f"  {co}: {desc}\n"); r.font.size = Pt(8); r.font.name = 'Times New Roman'


# ═══════════════ SET R1 ═══════════════
# Jumbled: Q1 mixes vibrations+rigid body+energy+forces
# Q2 = rolling/energy (was Q3-type), Q3 = pulleys/constrained (was Q2-type)
# Q4 = Lagrangian/energy (was Q5-type), Q5 = rotating frames (was Q4-type)
# Q6/Q7 = FBD swapped topics

def write_set_r1(doc):
    # ── Q1 (Compulsory) ──
    q_head(doc, 1, 20, "Compulsory — Answer ALL parts")

    sub_q(doc, "a",
          "Define the following terms related to vibration analysis: "
          "(i) Natural frequency, (ii) Damping ratio, (iii) Logarithmic decrement, and "
          "(iv) Critical damping. State the physical significance of each and give one "
          "practical example where each concept is important.", 5)

    sub_q(doc, "b",
          "A solid cylinder of mass 8 kg and radius 0.25 m rolls without slipping on a "
          "horizontal surface at a velocity of 4 m/s. "
          "Calculate: (i) the angular velocity, "
          "(ii) the translational kinetic energy, "
          "(iii) the rotational kinetic energy, and "
          "(iv) the total kinetic energy.", 5)

    sub_q(doc, "c",
          "A spring of stiffness 500 N/m is compressed by 0.15 m and used to launch a "
          "0.5 kg ball vertically upward. Using energy conservation, find: "
          "(i) the velocity of the ball at the moment of release, and "
          "(ii) the maximum height reached by the ball above the release point.", 5)

    sub_q(doc, "d",
          "Two blocks are connected by a light string over a frictionless pulley. "
          "Block A (8 kg) hangs vertically and Block B (6 kg) sits on a rough horizontal "
          "surface with μ_k = 0.25. Find: "
          "(i) the acceleration of the system, and "
          "(ii) the tension in the string.", 5)

    line(doc)

    # ── Q2 — Pure Numerical: Rolling & Energy (jumbled from Q3-type) ──
    big_q(doc, 2, "Pure Numerical — Rolling Dynamics & Energy",
          "Attempt all four parts.",
          [("a",
            "A solid sphere of mass 4 kg and radius 0.12 m rolls without slipping up a "
            "25° incline. Its initial velocity at the base is 6 m/s. "
            "Calculate: (i) the deceleration of the center of mass, "
            "(ii) the distance traveled up the incline before stopping, and "
            "(iii) the time to stop.",
            5),
           ("b",
            "A flywheel modeled as a solid disk of mass 25 kg and radius 0.4 m is spinning "
            "at 600 RPM. A brake applies a constant friction torque of 15 N·m. "
            "Find: (i) the moment of inertia, "
            "(ii) the angular deceleration, "
            "(iii) the time to stop, and "
            "(iv) the total energy dissipated by the brake.",
            5),
           ("c",
            "A 3 kg block at the top of a frictionless circular track of radius 5 m is "
            "released from rest. Find: (i) the speed at the bottom of the track, "
            "(ii) the normal force at the bottom, and "
            "(iii) the height at which the block leaves the track if the track is only "
            "a quarter circle (90° arc).",
            5),
           ("d",
            "A torsional pendulum consists of a disk of mass 1.5 kg and radius 0.1 m "
            "suspended by a wire of torsional stiffness k_t = 0.8 N·m/rad. "
            "Find: (i) the moment of inertia of the disk, "
            "(ii) the natural frequency of torsional oscillation in rad/s and Hz, and "
            "(iii) the period. "
            "If the initial angular displacement is 15°, find the maximum angular velocity.",
            5)])

    # ── Q3 — Pure Numerical: Pulleys & Constrained Motion (jumbled from Q2-type) ──
    big_q(doc, 3, "Pure Numerical — Constrained Motion & Particle Kinetics",
          "Attempt all four parts.",
          [("a",
            "A 3-pulley block-and-tackle system with individual pulley efficiency of 90% "
            "is used to lift a 500 kg crate through 4 m. "
            "Calculate: (i) the ideal mechanical advantage, "
            "(ii) the overall efficiency, "
            "(iii) the input force required, and "
            "(iv) the total heat generated.",
            5),
           ("b",
            "A conical pendulum has a bob of mass 0.8 kg on a string of length 1.2 m. "
            "The bob traces a horizontal circle at a constant speed such that the string "
            "makes an angle of 30° with the vertical. "
            "Find: (i) the tension in the string, "
            "(ii) the radius of the circular path, "
            "(iii) the speed of the bob, and "
            "(iv) the period of revolution.",
            5),
           ("c",
            "A 10 kg block on a frictionless 45° incline is connected by a light string "
            "over a pulley to a 6 kg block hanging vertically. "
            "Find: (i) the acceleration of the system, "
            "(ii) the tension in the string, and "
            "(iii) the velocity of each block after 2 seconds starting from rest.",
            5),
           ("d",
            "A CNC machine performs a circular interpolation with radius 40 mm at a "
            "feedrate of 60 mm/s, followed by a 100 mm straight-line cut at 80 mm/s. "
            "Calculate: (i) the time for the quarter-circle arc, "
            "(ii) the centripetal acceleration during the arc, "
            "(iii) the time for the straight-line cut, and "
            "(iv) the total machining time.",
            5)])

    # ── Q4 — Theory + Numerical: Lagrangian & Energy (jumbled from Q5-type) ──
    big_q(doc, 4, "Theory 10M + Numerical 10M — Lagrangian & Energy Methods",
          None,
          [("a",
            "Define the Lagrangian L = T − V and state Hamilton's principle of least action. "
            "What are generalized coordinates? Explain why Lagrangian mechanics is preferred "
            "over Newtonian mechanics for systems with holonomic constraints.",
            5),
           ("b",
            "Derive the equation of motion for a mass m on a spring (stiffness k) using the "
            "Lagrangian method. Extend this to include viscous damping c by using the Rayleigh "
            "dissipation function D = ½cq̇². Show that the resulting EOM is mẍ + cẋ + kx = 0.",
            5),
           ("c",
            "A 6 kg block slides down a smooth curved surface from a height of 5 m and then "
            "enters a rough horizontal patch (μ_k = 0.4) of length 3 m. After the rough patch, "
            "it hits a spring of stiffness 2000 N/m. "
            "Find: (i) the speed just before entering the rough patch, "
            "(ii) the speed just after the rough patch, and "
            "(iii) the maximum compression of the spring.",
            5),
           ("d",
            "A ballistic pendulum consists of a 5 kg wooden block suspended by a string. "
            "A 20 g bullet traveling at 400 m/s embeds itself in the block. "
            "Find: (i) the velocity of the block+bullet just after impact (using conservation "
            "of momentum), (ii) the maximum height reached by the block+bullet combination "
            "(using conservation of energy), and (iii) the percentage of kinetic energy lost "
            "in the collision.",
            5)])

    # ── Q5 — Theory + Numerical: Rotating Frames / Rigid Body (jumbled from Q4-type) ──
    big_q(doc, 5, "Theory 10M + Numerical 10M — Rotating Frames & Rigid Body Kinematics",
          None,
          [("a",
            "Explain the concept of an Instantaneous Center of Rotation (ICR) for a rigid body "
            "in planar motion. How is it determined for (i) pure rolling, and "
            "(ii) general planar motion? State two engineering applications of the ICR concept.",
            5),
           ("b",
            "Derive the acceleration of a point P on a rigid body performing general plane motion: "
            "a_P = a_cm + α × r + ω × (ω × r). Identify and explain the tangential and normal "
            "components of the acceleration due to rotation.",
            5),
           ("c",
            "A turntable rotates at Ω = 4 rad/s. An ant walks radially outward at a constant "
            "speed of 0.2 m/s starting from r = 0.1 m (position at t = 0). "
            "At t = 2 s, calculate: (i) the radial position of the ant, "
            "(ii) the Coriolis acceleration, "
            "(iii) the centripetal acceleration, and "
            "(iv) the total acceleration magnitude.",
            5),
           ("d",
            "A helicopter rotor blade of length 6 m has a uniform mass distribution. "
            "It rotates at 300 RPM. "
            "Find: (i) the angular velocity in rad/s, "
            "(ii) the tip speed, "
            "(iii) the centripetal acceleration at the tip, and "
            "(iv) the ratio of centripetal acceleration to g. "
            "Also calculate the velocity and centripetal acceleration at the midpoint.",
            5)])

    # ── Q6 — FBD + Difficult Numerical (jumbled: energy/vibrations focus) ──
    big_q(doc, 6, "FBD 10M + Difficult Numerical 10M",
          None,
          [("a",
            "A 4 kg mass hanging from a spring (k = 800 N/m) is pulled 0.08 m below its "
            "equilibrium position and released with an upward velocity of 0.5 m/s. "
            "The system has viscous damping with c = 16 N·s/m. "
            "Draw the complete Free Body Diagram of the mass in the displaced position "
            "(label all forces including spring, damper, and weight about equilibrium). "
            "Derive the equation of motion. "
            "Calculate: (i) ω_n, (ii) ζ, (iii) ω_d, (iv) classify the system, and "
            "(v) write the particular solution using the given initial conditions.",
            10),
           ("b",
            "A compound gear train has three stages: "
            "Stage 1 (18-tooth driver meshing with 72-tooth driven), "
            "Stage 2 (16-tooth driver meshing with 64-tooth driven), "
            "Stage 3 (20-tooth driver meshing with 80-tooth driven). "
            "The input motor runs at 3600 RPM, delivering 1.5 N·m torque. "
            "Each stage has efficiency 94%. "
            "Find: (i) the overall gear ratio, (ii) the output RPM, "
            "(iii) the output torque accounting for losses, "
            "(iv) the total power loss in watts, and "
            "(v) the overall efficiency.",
            10)])

    # ── Q7 — FBD + Difficult Numerical (jumbled: constrained/particle focus) ──
    big_q(doc, 7, "FBD 10M + Difficult Numerical 10M",
          None,
          [("a",
            "A 3 kg block is on a 40° rough incline (μ_s = 0.4, μ_k = 0.3). "
            "It is connected by a light string over a frictionless pulley at the top "
            "of the incline to a 5 kg block hanging vertically. "
            "Draw the complete Free Body Diagram for BOTH blocks, labeling all forces. "
            "Determine: (i) whether the system moves and in which direction, "
            "(ii) the acceleration of the system, and "
            "(iii) the tension in the string.",
            10),
           ("b",
            "A slider-crank mechanism has crank radius r = 60 mm, connecting rod length "
            "L = 200 mm, and operates at 4500 RPM. At crank angle θ = 60° from TDC, "
            "calculate: (i) the angular velocity of the crank, "
            "(ii) the connecting rod angle φ, "
            "(iii) the piston velocity, and "
            "(iv) the piston acceleration (approximate formula). "
            "Also determine the maximum possible piston velocity and the crank angle "
            "at which it occurs.",
            10)])


# ═══════════════ SET R2 ═══════════════
# Further jumbled: Q1 mixes rotation+collisions+gears+incline
# Q2 = gears/viscous (was Q3-d type), Q3 = energy/loop (was Q2-c type)
# Q4 = transport theorem + MOI (mixed), Q5 = pendulums/Lagrange (mixed)
# Q6 = conveyor + multi-DOF, Q7 = manipulator + crank

def write_set_r2(doc):
    # ── Q1 (Compulsory) ──
    q_head(doc, 1, 20, "Compulsory — Answer ALL parts")

    sub_q(doc, "a",
          "State Newton's three laws of motion. For each law, give one example from "
          "the field of mechatronics or robotics. Explain how the second law is "
          "generalized to rotational motion.", 5)

    sub_q(doc, "b",
          "A uniform rod of mass 3 kg and length 1.2 m is pivoted at a point 0.3 m "
          "from one end. Find: "
          "(i) the moment of inertia about the center of mass, "
          "(ii) the distance from the pivot to the center of mass, and "
          "(iii) the moment of inertia about the pivot using the Parallel Axis Theorem.", 5)

    sub_q(doc, "c",
          "A 2 kg ball traveling at 8 m/s collides head-on with a stationary 3 kg ball. "
          "The coefficient of restitution is e = 0.6. "
          "Find: (i) the velocities of both balls after collision, "
          "(ii) the kinetic energy lost, and "
          "(iii) the impulse experienced by each ball.", 5)

    sub_q(doc, "d",
          "A 4 kg block is placed on a rough horizontal surface with μ_s = 0.4 "
          "and μ_k = 0.35. A horizontal force of 25 N is applied. "
          "Find: (i) the normal force and maximum static friction force, "
          "(ii) whether the block moves, and "
          "(iii) the acceleration of the block.", 5)

    line(doc)

    # ── Q2 — Pure Numerical: Gears & Viscous/Rotating Systems (jumbled) ──
    big_q(doc, 2, "Pure Numerical — Rotating Machinery & Transmission",
          "Attempt all four parts.",
          [("a",
            "A two-stage gear train has Stage 1 (20-tooth driver meshing with 60-tooth driven) "
            "and Stage 2 (18-tooth driver meshing with 90-tooth driven). "
            "The input shaft runs at 2400 RPM delivering a torque of 1.2 N·m. "
            "Each stage has efficiency 95%. "
            "Find: (i) the overall gear ratio, "
            "(ii) the output RPM, "
            "(iii) the output torque including efficiency losses, and "
            "(iv) the total power loss.",
            5),
           ("b",
            "A solid disk flywheel of mass 50 kg and radius 0.5 m is rotating at 1200 RPM. "
            "A braking torque brings it to rest in 30 seconds. "
            "Find: (i) the moment of inertia, "
            "(ii) the initial angular velocity in rad/s, "
            "(iii) the angular deceleration, "
            "(iv) the braking torque, and "
            "(v) the total kinetic energy absorbed by the brake.",
            5),
           ("c",
            "A motor drives a load through a worm gear (GR = 40:1, η = 75%). "
            "The motor output is 2 kW at 1440 RPM. "
            "Find: (i) the motor torque, "
            "(ii) the output speed, "
            "(iii) the output torque, and "
            "(iv) the power lost as heat.",
            5),
           ("d",
            "A solid cylinder of mass 3 kg and radius 0.1 m is driven by a constant "
            "torque of 0.6 N·m against viscous friction (c = 0.2 N·m·s/rad). "
            "Find: (i) the moment of inertia, "
            "(ii) the steady-state angular velocity, "
            "(iii) the time constant, and "
            "(iv) the angular velocity after one time constant.",
            5)])

    # ── Q3 — Pure Numerical: Energy & Circular Motion (jumbled) ──
    big_q(doc, 3, "Pure Numerical — Energy Methods & Circular Motion",
          "Attempt all four parts.",
          [("a",
            "A roller-coaster car of mass 800 kg is at the top of a 25 m hill traveling "
            "at 5 m/s. It goes down and through a vertical loop of radius 10 m. "
            "Using energy conservation, find: "
            "(i) the speed at the bottom of the hill, "
            "(ii) the speed at the top of the loop (height 20 m from bottom), and "
            "(iii) the normal force at the top of the loop.",
            5),
           ("b",
            "A damped vibration system with mass 5 kg has two successive peaks: "
            "x₁ = 25 mm at t₁ = 0.1 s and x₂ = 18 mm at t₂ = 0.38 s. "
            "Find: (i) the damped period, "
            "(ii) the logarithmic decrement, "
            "(iii) the damping ratio, and "
            "(iv) the undamped natural frequency and spring stiffness.",
            5),
           ("c",
            "A 1.5 kg pendulum bob on a 2.0 m string is released from a horizontal "
            "position (θ = 90°). "
            "Find: (i) the speed at the lowest point, "
            "(ii) the tension in the string at the lowest point, and "
            "(iii) the height at which the tension equals the weight (if it exists).",
            5),
           ("d",
            "A 70 kg athlete jumps vertically upward from a crouching position. "
            "During the push-off phase (0.3 m), the average leg force is 2500 N. "
            "Find: (i) the net upward force, "
            "(ii) the velocity at take-off using work-energy, "
            "(iii) the maximum jump height, and "
            "(iv) the average power during push-off if the phase lasts 0.15 s.",
            5)])

    # ── Q4 — Theory + Numerical: Transport Theorem & MOI (jumbled) ──
    big_q(doc, 4, "Theory 10M + Numerical 10M — Rotating Frames & Rigid Body Properties",
          None,
          [("a",
            "State and derive the Transport Theorem for the velocity and acceleration "
            "of a point observed from a rotating reference frame. "
            "Define each acceleration term (Coriolis, centripetal, Euler, relative) "
            "with its physical origin.",
            5),
           ("b",
            "Define the radius of gyration k and its relation to the moment of inertia "
            "(I = mk²). State the perpendicular axis theorem and the parallel axis theorem. "
            "Give one application of each in engineering design.",
            5),
           ("c",
            "A centrifuge arm of length r = 2.5 m rotates at 500 RPM. "
            "A test sample at the tip of the arm experiences a small vibration in the "
            "radial direction with velocity ṙ = 0.05 m/s. "
            "Calculate: (i) the centripetal acceleration, "
            "(ii) the Coriolis acceleration, "
            "(iii) the total acceleration magnitude, and "
            "(iv) express the centripetal acceleration in units of g.",
            5),
           ("d",
            "A compound object consists of a solid disk (m₁ = 4 kg, R = 0.15 m) "
            "welded to a uniform rod (m₂ = 2 kg, L = 0.6 m) at the rod's end. "
            "The rod extends radially from the edge of the disk. "
            "Find the total moment of inertia about the disk center.",
            5)])

    # ── Q5 — Theory + Numerical: Pendulums & Lagrangian (jumbled) ──
    big_q(doc, 5, "Theory 10M + Numerical 10M — Pendulums & Lagrangian Mechanics",
          None,
          [("a",
            "Compare a simple pendulum and a physical (compound) pendulum. "
            "Derive the period of each for small oscillations. "
            "Under what condition does a physical pendulum have the same period "
            "as a simple pendulum?",
            5),
           ("b",
            "Explain the concept of normal modes in a 2-DOF vibrating system. "
            "What information do the eigenvalues and eigenvectors of the system provide? "
            "Describe the physical meaning of the two mode shapes for two masses "
            "connected by springs.",
            5),
           ("c",
            "A U-tube manometer contains a liquid column of total length 0.6 m. "
            "When displaced slightly, the liquid oscillates. "
            "Derive the equation of motion using the Lagrangian approach "
            "(generalized coordinate = displacement x from equilibrium). "
            "Find the natural frequency and period of oscillation.",
            5),
           ("d",
            "A 2-DOF system has masses m₁ = 2 kg and m₂ = 3 kg coupled by springs: "
            "Wall — k₁(200 N/m) — m₁ — k₂(300 N/m) — m₂ — k₃(200 N/m) — Wall. "
            "Write the [M] and [K] matrices. Solve for the two natural frequencies. "
            "Describe the mode shapes.",
            5)])

    # ── Q6 — FBD + Difficult Numerical (jumbled: conveyor + vibration isolation) ──
    big_q(doc, 6, "FBD 10M + Difficult Numerical 10M",
          None,
          [("a",
            "A 5 kg package slides from a chute onto a conveyor belt moving at 4.0 m/s. "
            "The package's initial horizontal speed is 1.0 m/s (in the belt direction). "
            "Coefficients of friction: μ_s = 0.4, μ_k = 0.3. "
            "Draw the complete Free Body Diagram showing all forces during the sliding phase. "
            "Determine: (i) the friction force and acceleration, "
            "(ii) the time for the package to reach belt speed, "
            "(iii) the distances traveled by the package and the belt surface, and "
            "(iv) the energy dissipated as heat.",
            10),
           ("b",
            "A machine of mass 100 kg is mounted on springs with total stiffness "
            "k = 40000 N/m and dashpots with total damping c = 800 N·s/m. "
            "The machine operates at 1500 RPM and generates a rotating unbalance "
            "force of 500 N amplitude. "
            "Find: (i) the natural frequency, "
            "(ii) the damping ratio, "
            "(iii) the frequency ratio, "
            "(iv) the steady-state vibration amplitude, and "
            "(v) the force transmitted to the foundation (transmissibility).",
            10)])

    # ── Q7 — FBD + Difficult Numerical (jumbled: manipulator + crank) ──
    big_q(doc, 7, "FBD 10M + Difficult Numerical 10M",
          None,
          [("a",
            "A two-link planar robot arm has link 1 (L₁ = 0.5 m, m₁ = 3 kg) rotating "
            "at ω₁ = 2 rad/s with α₁ = 0, and link 2 (L₂ = 0.4 m, m₂ = 2 kg) at "
            "ω₂ = 3 rad/s (relative to link 1) with α₂ = 1 rad/s². "
            "Both links are in a horizontal plane. At the instant when θ₁ = 0° and θ₂ = 90°, "
            "draw the Free Body Diagram of the end-effector mass (1 kg at the tip of link 2). "
            "Calculate the absolute velocity and acceleration of the tip using the transport theorem.",
            10),
           ("b",
            "A Scotch yoke mechanism converts rotary to pure sinusoidal linear motion. "
            "The crank radius is r = 40 mm and it rotates at 1800 RPM. "
            "The slider has mass 0.5 kg. "
            "Find: (i) the maximum displacement, velocity, and acceleration of the slider, "
            "(ii) the maximum inertia force on the slider, "
            "(iii) the power required to accelerate the slider at the position of maximum acceleration, and "
            "(iv) the slider position and velocity when the crank is at θ = 45°.",
            10)])


# ═══════════════ MAIN ═══════════════

if __name__ == "__main__":
    out = os.path.join(SCRIPT_DIR, "Re-Exam")
    os.makedirs(out, exist_ok=True)

    sets = {1: (write_set_r1, MASTER_ROWS_R1), 2: (write_set_r2, MASTER_ROWS_R2)}

    for num, (writer, rows) in sets.items():
        doc = new_doc(num)
        writer(doc)
        add_synoptic_table(doc, rows)
        path = os.path.join(out, f"DSM_ReExam_QP_SetR{num}.docx")
        doc.save(path)
        print(f"[OK] {path}")

    print("\nDone — 2 DSM Re-Exam question papers generated.")
