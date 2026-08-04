"""
DSM Re-Exam Answer Booklets — Sets R1, R2
===========================================
Dynamic Systems Modeling (702MH0C025)
B.Tech Mechatronics, Semester IV

Comprehensive solutions with step-by-step working,
mark allocation, and key notes.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# ═══════════════ HELPER FUNCTIONS ═══════════════

def new_doc(set_number):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.5); s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)
    for s in doc.sections:
        pgB = OxmlElement('w:pgBorders')
        pgB.set(qn('w:offsetFrom'), 'page')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '18')
            b.set(qn('w:space'), '24'); b.set(qn('w:color'), '000000')
            pgB.append(b)
        s._sectPr.append(pgB)

    def ctr(txt, sz=12, bold=True, color=None):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(0)
        r = p.add_run(txt); r.bold = bold; r.font.size = Pt(sz); r.font.name = 'Times New Roman'
        if color: r.font.color.rgb = color

    ctr("SVKM's NMIMS — STME Indore", 13)
    ctr(f"DSM (702MH0C025) — RE-EXAM ANSWER BOOKLET — SET R{set_number}", 12)
    ctr("Re-Examination — March 2026", 10, bold=False)
    sep(doc)
    return doc


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'
    r.font.size = Pt(12 if level == 1 else 10)
    r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)


def point(doc, text, marks=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(9); r.font.name = 'Times New Roman'
    if marks:
        r = p.add_run(f"  [{marks}]"); r.bold = True; r.font.size = Pt(8)
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(8); r.font.name = 'Times New Roman'
    r.italic = True; r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def marks_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0x00, 0x00, 0xCC)


def sep(doc):
    p = doc.add_paragraph("─" * 95)
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    for r in p.runs:
        r.font.size = Pt(6); r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0x99,0x99,0x99)


# ═══════════════ SET R1 ANSWERS ═══════════════

def answers_set_r1(doc):
    # ── Q1(a) ──
    heading(doc, "Q1 (a)  —  Vibration Definitions [5 marks]")
    point(doc, "(i) Natural frequency ω_n = √(k/m): the frequency at which an undamped system oscillates freely. "
          "Example: tuning fork vibrations.", "1.25")
    point(doc, "(ii) Damping ratio ζ = c/c_cr = c/(2√(km)): dimensionless measure of damping relative to critical. "
          "Example: automotive shock absorber design.", "1.25")
    point(doc, "(iii) Logarithmic decrement δ = ln(x_n/x_{n+1}): natural log of ratio of successive peaks. "
          "Used to experimentally determine ζ. Example: vibration testing of structures.", "1.25")
    point(doc, "(iv) Critical damping (ζ = 1, c_cr = 2√(km)): fastest return to equilibrium without oscillation. "
          "Example: door closer mechanism.", "1.25")
    sep(doc)

    # ── Q1(b) ──
    heading(doc, "Q1 (b)  —  Rolling Cylinder KE [5 marks]")
    point(doc, "Given: m = 8 kg, R = 0.25 m, v = 4 m/s, solid cylinder I = ½mR²")
    point(doc, "ω = v/R = 4/0.25 = 16 rad/s", "1")
    point(doc, "KE_trans = ½mv² = ½(8)(16) = 64 J", "1")
    point(doc, "I = ½(8)(0.0625) = 0.25 kg·m²; KE_rot = ½Iω² = ½(0.25)(256) = 32 J", "1.5")
    point(doc, "KE_total = 64 + 32 = 96 J", "1")
    note(doc, "For a rolling cylinder: KE_rot/KE_total = 1/3 ≈ 33.3%. Check: 32/96 = 0.333 ")
    point(doc, "Alternatively: KE = ½mv²(1 + β) = ½(8)(16)(1 + 0.5) = 96 J ", "0.5")
    sep(doc)

    # ── Q1(c) ──
    heading(doc, "Q1 (c)  —  Spring Launch [5 marks]")
    point(doc, "Given: k = 500 N/m, x = 0.15 m, m = 0.5 kg")
    point(doc, "PE_spring = ½kx² = ½(500)(0.0225) = 5.625 J", "1")
    point(doc, "At release: ½mv² = PE_spring → v = √(2 × 5.625/0.5) = √22.5 = 4.743 m/s", "2")
    point(doc, "Max height: ½mv² = mgh → h = v²/(2g) = 22.5/19.62 = 1.147 m", "2")
    sep(doc)

    # ── Q1(d) ──
    heading(doc, "Q1 (d)  —  Two-Block Pulley System [5 marks]")
    point(doc, "Given: m_A = 8 kg (hanging), m_B = 6 kg (horizontal, μ_k = 0.25)")
    point(doc, "Forces on B: f_k = μ_k m_B g = 0.25(6)(9.81) = 14.72 N (opposing motion)")
    point(doc, "Net force = m_A g − f_k = 8(9.81) − 14.72 = 78.48 − 14.72 = 63.76 N", "1.5")
    point(doc, "a = 63.76/(m_A + m_B) = 63.76/14 = 4.554 m/s²", "2")
    point(doc, "T = m_A(g − a) = 8(9.81 − 4.554) = 8(5.256) = 42.05 N", "1.5")
    note(doc, "Check: T − f_k = 42.05 − 14.72 = 27.33 N = m_B × a = 6 × 4.554 = 27.32 N ")
    sep(doc)

    # ── Q2 ──
    heading(doc, "Q2  —  Rolling Dynamics & Energy [20 marks]")

    heading(doc, "Q2 (a)  —  Sphere Rolling Up Incline [5 marks]", 2)
    point(doc, "Given: m = 4 kg, R = 0.12 m, θ = 25°, v₀ = 6 m/s, solid sphere β = 2/5")
    point(doc, "a = −g sin θ/(1 + β) = −9.81 sin 25°/1.4 = −9.81(0.4226)/1.4 = −2.960 m/s²", "2")
    point(doc, "d = v₀²/(2|a|) = 36/(5.92) = 6.08 m", "1.5")
    point(doc, "t = v₀/|a| = 6/2.960 = 2.027 s", "1.5")

    heading(doc, "Q2 (b)  —  Flywheel Braking [5 marks]", 2)
    point(doc, "Given: m = 25 kg, R = 0.4 m, n = 600 RPM, τ_brake = 15 N·m")
    point(doc, "I = ½mR² = ½(25)(0.16) = 2.0 kg·m²", "1")
    point(doc, "ω₀ = 2π(600)/60 = 62.83 rad/s", "0.5")
    point(doc, "α = τ/I = 15/2 = 7.5 rad/s²", "1")
    point(doc, "t_stop = ω₀/α = 62.83/7.5 = 8.38 s", "1")
    point(doc, "E = ½Iω₀² = ½(2)(3947.8) = 3947.8 J ≈ 3.95 kJ", "1.5")

    heading(doc, "Q2 (c)  —  Circular Track [5 marks]", 2)
    point(doc, "Given: m = 3 kg, R = 5 m, from rest at top (quarter circle)")
    point(doc, "h = R = 5 m; v_bottom = √(2gR) = √(2 × 9.81 × 5) = √98.1 = 9.905 m/s", "2")
    point(doc, "N − mg = mv²/R → N = m(g + v²/R) = 3(9.81 + 98.1/5) = 3(9.81 + 19.62) = 88.29 N", "2")
    point(doc, "Quarter circle ends at bottom — block stays on track throughout (normal force always positive)", "1")
    note(doc, "If track were inverted/loop: would need v > √(gR) at top to maintain contact")

    heading(doc, "Q2 (d)  —  Torsional Pendulum [5 marks]", 2)
    point(doc, "Given: m = 1.5 kg, R = 0.1 m, k_t = 0.8 N·m/rad, θ₀ = 15° = 0.2618 rad")
    point(doc, "I = ½mR² = ½(1.5)(0.01) = 0.0075 kg·m²", "1")
    point(doc, "ω_n = √(k_t/I) = √(0.8/0.0075) = √106.67 = 10.33 rad/s; f_n = 1.644 Hz", "2")
    point(doc, "T = 2π/10.33 = 0.608 s", "1")
    point(doc, "ω_max = ω_n × θ₀ = 10.33 × 0.2618 = 2.704 rad/s", "1")
    sep(doc)

    # ── Q3 ──
    heading(doc, "Q3  —  Constrained Motion & Particle Kinetics [20 marks]")

    heading(doc, "Q3 (a)  —  3-Pulley Block & Tackle [5 marks]", 2)
    point(doc, "Given: 3 pulleys, η = 0.90 each, m = 500 kg, h = 4 m")
    point(doc, "MA = 3; η_total = 0.90³ = 0.729", "1")
    point(doc, "F_ideal = 500(9.81)/3 = 1635 N", "0.5")
    point(doc, "F_actual = 1635/0.729 = 2242.8 N", "1.5")
    point(doc, "W_out = mgh = 500(9.81)(4) = 19620 J", "0.5")
    point(doc, "W_in = F_actual × 3 × 4 = 2242.8 × 12 = 26913.6 J", "0.5")
    point(doc, "Heat = 26913.6 − 19620 = 7293.6 J ≈ 7.29 kJ", "1")

    heading(doc, "Q3 (b)  —  Conical Pendulum [5 marks]", 2)
    point(doc, "Given: m = 0.8 kg, L = 1.2 m, angle = 30° from vertical")
    point(doc, "T cos 30° = mg → T = mg/cos 30° = 0.8(9.81)/0.866 = 9.064 N", "1.5")
    point(doc, "r = L sin 30° = 1.2(0.5) = 0.6 m", "0.5")
    point(doc, "T sin 30° = mv²/r → v = √(Tr sin 30°/m) = √(9.064 × 0.6 × 0.5/0.8) = √3.399 = 1.844 m/s", "1.5")
    point(doc, "P = 2πr/v = 2π(0.6)/1.844 = 2.043 s", "1.5")

    heading(doc, "Q3 (c)  —  Incline-Pulley System [5 marks]", 2)
    point(doc, "Given: m₁ = 10 kg on 45° frictionless, m₂ = 6 kg hanging")
    point(doc, "m₁ component down incline: 10(9.81) sin 45° = 69.37 N")
    point(doc, "m₂ weight: 6(9.81) = 58.86 N")
    point(doc, "Since 69.37 > 58.86: m₁ slides down, m₂ goes up", "0.5")
    point(doc, "a = (m₁ g sin 45° − m₂ g)/(m₁ + m₂) = (69.37 − 58.86)/16 = 10.51/16 = 0.657 m/s²", "2")
    point(doc, "T = m₂(g + a) = 6(9.81 + 0.657) = 6(10.467) = 62.80 N", "1.5")
    point(doc, "v(2 s) = 0.657 × 2 = 1.314 m/s", "1")

    heading(doc, "Q3 (d)  —  CNC Circular + Linear [5 marks]", 2)
    point(doc, "Given: circle R = 40 mm, F = 60 mm/s, quarter arc; line = 100 mm at 80 mm/s")
    point(doc, "Arc length = πR/2 = π(40)/2 = 62.83 mm; t_arc = 62.83/60 = 1.047 s", "2")
    point(doc, "a_c = v²/R = 3600/40 = 90 mm/s² = 0.09 m/s²", "1")
    point(doc, "t_line = 100/80 = 1.25 s", "1")
    point(doc, "t_total = 1.047 + 1.25 = 2.297 s", "1")
    sep(doc)

    # ── Q4 ──
    heading(doc, "Q4  —  Lagrangian & Energy Methods [20 marks]")

    heading(doc, "Q4 (a)  —  Lagrangian Theory [5 marks]", 2)
    point(doc, "L = T − V (Lagrangian = Kinetic − Potential energy)", "1")
    point(doc, "Hamilton's principle: δ∫L dt = 0 — the actual path minimizes the action integral", "1")
    point(doc, "Generalized coordinates: minimum set of independent variables describing configuration", "1")
    point(doc, "Advantage: constraint forces never appear in the equations of motion", "1")
    point(doc, "Holonomic constraints reduce complexity; Lagrangian auto-eliminates internal forces", "1")

    heading(doc, "Q4 (b)  —  Spring-Mass-Damper via Lagrange [5 marks]", 2)
    point(doc, "T = ½mẋ², V = ½kx² → L = ½mẋ² − ½kx²", "1")
    point(doc, "d/dt(∂L/∂ẋ) = mẍ; ∂L/∂x = −kx → mẍ + kx = 0 (undamped)", "1.5")
    point(doc, "Rayleigh dissipation: D = ½cẋ²", "0.5")
    point(doc, "Modified E-L: d/dt(∂L/∂ẋ) − ∂L/∂x + ∂D/∂ẋ = 0 → mẍ + cẋ + kx = 0", "2")

    heading(doc, "Q4 (c)  —  Block + Rough Patch + Spring [5 marks]", 2)
    point(doc, "Given: m = 6 kg, h = 5 m, μ_k = 0.4, d_rough = 3 m, k = 2000 N/m")
    point(doc, "Before rough: v₁ = √(2gh) = √(98.1) = 9.905 m/s", "1")
    point(doc, "Friction work: W_f = μ_k mg d = 0.4(6)(9.81)(3) = 70.63 J", "1")
    point(doc, "KE after rough: ½mv₂² = ½mv₁² − W_f = ½(6)(98.1) − 70.63 = 294.3 − 70.63 = 223.67 J", "1")
    point(doc, "v₂ = √(2 × 223.67/6) = √74.56 = 8.635 m/s", "0.5")
    point(doc, "Spring compression: ½kδ² = 223.67 → δ = √(2 × 223.67/2000) = √0.2237 = 0.4729 m", "1.5")
    note(doc, "Assuming friction does not act during spring compression (spring on smooth surface beyond rough patch)")

    heading(doc, "Q4 (d)  —  Ballistic Pendulum [5 marks]", 2)
    point(doc, "Given: m_block = 5 kg, m_bullet = 0.02 kg, v_bullet = 400 m/s")
    point(doc, "Momentum: 0.02(400) = (5.02)v → v = 8.0/5.02 = 1.594 m/s", "1.5")
    point(doc, "Height: ½mv² = mgh → h = v²/(2g) = 2.541/(19.62) = 0.1295 m ≈ 12.95 cm", "1.5")
    point(doc, "KE_before = ½(0.02)(160000) = 1600 J; KE_after = ½(5.02)(2.541) = 6.377 J", "1")
    point(doc, "% lost = (1600 − 6.377)/1600 × 100 = 99.60%", "1")
    note(doc, "Perfectly inelastic collision — maximum possible KE loss while conserving momentum")
    sep(doc)

    # ── Q5 ──
    heading(doc, "Q5  —  Rotating Frames & Rigid Body Kinematics [20 marks]")

    heading(doc, "Q5 (a)  —  ICR Theory [5 marks]", 2)
    point(doc, "ICR: the point about which all points of a rigid body appear to rotate at a given instant", "1")
    point(doc, "Pure rolling: ICR is at the contact point (zero velocity point)", "1")
    point(doc, "General planar motion: ICR found by (i) intersection of velocity perpendiculars, or "
          "(ii) if v and ω known: ICR lies at distance v/ω from the point", "1.5")
    point(doc, "Application 1: velocity analysis of mechanisms (slider-crank, four-bar)", "0.75")
    point(doc, "Application 2: path curvature analysis for cam profiles / gear tooth design", "0.75")

    heading(doc, "Q5 (b)  —  Rigid Body Acceleration Derivation [5 marks]", 2)
    point(doc, "a_P = a_cm + α × r_{P/cm} + ω × (ω × r_{P/cm})", "1.5")
    point(doc, "α × r: tangential component — perpendicular to r, in the plane of rotation", "1")
    point(doc, "ω × (ω × r): centripetal/normal component — directed from P toward rotation axis (−ω²r radially)", "1")
    point(doc, "|a_tangential| = αr; |a_normal| = ω²r; |a_total| = r√(α² + ω⁴)", "1.5")

    heading(doc, "Q5 (c)  —  Ant on Turntable [5 marks]", 2)
    point(doc, "Given: Ω = 4 rad/s, ṙ = 0.2 m/s, r₀ = 0.1 m, at t = 2 s")
    point(doc, "r(2) = r₀ + ṙ × t = 0.1 + 0.2(2) = 0.5 m", "1")
    point(doc, "a_Coriolis = 2Ωṙ = 2(4)(0.2) = 1.6 m/s²", "1")
    point(doc, "a_centripetal = Ω²r = 16(0.5) = 8.0 m/s²", "1.5")
    point(doc, "|a_total| = √(a_Cor² + a_cen²) = √(2.56 + 64) = √66.56 = 8.159 m/s²", "1.5")
    note(doc, "Euler term is zero since Ω is constant; a_rel = 0 since ṙ constant (r̈ = 0)")

    heading(doc, "Q5 (d)  —  Helicopter Rotor [5 marks]", 2)
    point(doc, "Given: L = 6 m, n = 300 RPM")
    point(doc, "ω = 2π(300)/60 = 10π = 31.42 rad/s", "1")
    point(doc, "v_tip = ωL = 31.42 × 6 = 188.5 m/s = 678.6 km/h", "1")
    point(doc, "a_c_tip = ω²L = 986.96 × 6 = 5921.8 m/s²", "1")
    point(doc, "a_c/g = 5921.8/9.81 = 603.6 g", "0.5")
    point(doc, "Midpoint (r = 3): v_mid = 31.42 × 3 = 94.25 m/s; a_mid = 986.96 × 3 = 2960.9 m/s²", "1.5")
    sep(doc)

    # ── Q6 ──
    heading(doc, "Q6  —  FBD + Difficult Numerical [20 marks]")

    heading(doc, "Q6 (a)  —  Spring-Mass-Damper Full Analysis [10 marks]", 2)
    point(doc, "Given: m = 4 kg, k = 800 N/m, c = 16 N·s/m, x₀ = −0.08 m (below eq), ẋ₀ = +0.5 m/s (upward)")
    marks_note(doc, "FBD [3 marks]:")
    point(doc, "Spring force: kx = 800(0.08) = 64 N (upward, restoring toward equilibrium)", "1")
    point(doc, "Damping force: cẋ (opposing velocity, so downward when moving up)", "1")
    point(doc, "About static equilibrium: weight and spring preload cancel — only dynamic forces matter", "1")
    marks_note(doc, "EOM and Parameters [4 marks]:")
    point(doc, "4ẍ + 16ẋ + 800x = 0 → ẍ + 4ẋ + 200x = 0", "1")
    point(doc, "ω_n = √(200) = 14.14 rad/s", "0.75")
    point(doc, "c_cr = 2mω_n = 2(4)(14.14) = 113.14 N·s/m → ζ = 16/113.14 = 0.1414", "0.75")
    point(doc, "ω_d = ω_n√(1 − ζ²) = 14.14√(1 − 0.02) = 14.14 × 0.9899 = 13.997 ≈ 14.0 rad/s", "0.75")
    point(doc, "System: Underdamped (ζ = 0.1414 < 1)", "0.75")
    marks_note(doc, "Solution with ICs [3 marks]:")
    point(doc, "x(t) = e^(−2t)[A cos(14t) + B sin(14t)]", "1")
    point(doc, "x(0) = A = −0.08 m", "0.5")
    point(doc, "ẋ(0) = −2A + 14B = 0.5 → 0.16 + 14B = 0.5 → B = 0.0243 m", "0.75")
    point(doc, "x(t) = e^(−2t)[−0.08 cos(14t) + 0.0243 sin(14t)] m", "0.75")

    heading(doc, "Q6 (b)  —  3-Stage Gear Train [10 marks]", 2)
    point(doc, "Given: GR₁ = 72/18 = 4, GR₂ = 64/16 = 4, GR₃ = 80/20 = 4; η = 0.94/stage; n_in = 3600 RPM, τ_in = 1.5 N·m")
    point(doc, "GR_total = 4 × 4 × 4 = 64", "1.5")
    point(doc, "n_out = 3600/64 = 56.25 RPM = 5.89 rad/s", "1.5")
    point(doc, "η_total = 0.94³ = 0.8306", "1")
    point(doc, "τ_out = τ_in × GR × η_total = 1.5 × 64 × 0.8306 = 79.74 N·m", "2")
    point(doc, "P_in = τ_in × ω_in = 1.5 × 2π(3600)/60 = 1.5 × 376.99 = 565.49 W", "1.5")
    point(doc, "P_out = P_in × η_total = 565.49 × 0.8306 = 469.67 W", "1")
    point(doc, "P_loss = 565.49 − 469.67 = 95.82 W", "0.75")
    point(doc, "Overall η = 83.06%", "0.75")
    sep(doc)

    # ── Q7 ──
    heading(doc, "Q7  —  FBD + Difficult Numerical [20 marks]")

    heading(doc, "Q7 (a)  —  Incline + Hanging Block FBD [10 marks]", 2)
    point(doc, "Given: m₃ = 3 kg on 40° incline (μ_s = 0.4, μ_k = 0.3), m₅ = 5 kg hanging")
    marks_note(doc, "FBD — Block on incline [2 marks]:")
    point(doc, "Weight W₃ = 29.43 N; N = W₃ cos 40° = 29.43(0.766) = 22.54 N", "1")
    point(doc, "Friction (up the incline, opposing relative motion): f", "0.5")
    point(doc, "Tension T (up the incline through string)", "0.5")
    marks_note(doc, "FBD — Hanging block [2 marks]:")
    point(doc, "Weight W₅ = 49.05 N (down); Tension T (up)", "2")
    marks_note(doc, "Motion check [3 marks]:")
    point(doc, "If system moves (m₅ down, m₃ up incline):")
    point(doc, "Driving force = m₅g = 49.05 N", "0.5")
    point(doc, "Opposing = m₃g sin 40° + f_s_max = 29.43(0.6428) + 0.4(22.54) = 18.92 + 9.02 = 27.94 N", "1")
    point(doc, "49.05 > 27.94 → YES, system moves. m₅ goes down, m₃ goes up the incline.", "1.5")
    marks_note(doc, "Acceleration & Tension [3 marks]:")
    point(doc, "Kinetic friction: f_k = μ_k N = 0.3(22.54) = 6.762 N", "0.5")
    point(doc, "ΣF = m₅g − m₃g sin 40° − f_k = 49.05 − 18.92 − 6.762 = 23.368 N", "1")
    point(doc, "a = 23.368/(3 + 5) = 23.368/8 = 2.921 m/s²", "0.75")
    point(doc, "T = m₅(g − a) = 5(9.81 − 2.921) = 5(6.889) = 34.45 N", "0.75")

    heading(doc, "Q7 (b)  —  Slider-Crank at 4500 RPM [10 marks]", 2)
    point(doc, "Given: r = 60 mm, L = 200 mm, n = 4500 RPM, θ = 60°")
    point(doc, "ω = 2π(4500)/60 = 471.24 rad/s", "1.5")
    point(doc, "sin φ = (r/L) sin θ = (60/200)(0.866) = 0.2598 → φ = 15.05°; cos φ = 0.9659", "1.5")
    point(doc, "v_piston = −rω[sin θ + (r sin 2θ)/(2L)] = −60(471.24)[0.866 + 0.3 × sin 120°/2]", "1")
    note(doc, "Exact: v = −rω sin θ − Lφ̇ sin φ")
    point(doc, "φ̇ = rω cos θ/(L cos φ) = 60(471.24)(0.5)/(200 × 0.9659) = 14137.2/193.18 = 73.18 rad/s", "1")
    point(doc, "v = −60(471.24)(0.866) − 200(73.18)(0.2598) = −24485 − 3802 = −28287 mm/s = −28.29 m/s", "1.5")
    point(doc, "Approximate acceleration: a ≈ −rω²[cos θ + r cos 2θ/L] = −60(471.24²)[0.5 + 0.3(−0.5)]", "1")
    point(doc, "= −60(222027)[0.35] = −4662567 mm/s² = −4662.6 m/s² ≈ 475 g", "1")
    point(doc, "Max velocity occurs near θ ≈ 76–80° for this L/r ratio (numerical max at sin θ peak of v expression)", "0.5")
    sep(doc)


# ═══════════════ SET R2 ANSWERS ═══════════════

def answers_set_r2(doc):
    # ── Q1(a) ──
    heading(doc, "Q1 (a)  —  Newton's Laws [5 marks]")
    point(doc, "1st Law (Inertia): A body at rest stays at rest unless acted upon by an external force. "
          "Example: robot arm maintains position until motors activate.", "1.5")
    point(doc, "2nd Law: F = ma (ΣF = dp/dt). Rotational: τ = Iα. "
          "Example: motor torque determines angular acceleration of a joint.", "2")
    point(doc, "3rd Law: Action = −Reaction. Example: robot gripper grips object — object pushes "
          "back on gripper with equal force.", "1.5")
    sep(doc)

    # ── Q1(b) ──
    heading(doc, "Q1 (b)  —  Rod MOI with PAT [5 marks]")
    point(doc, "Given: m = 3 kg, L = 1.2 m, pivot 0.3 m from one end")
    point(doc, "I_cm = mL²/12 = 3(1.44)/12 = 0.36 kg·m²", "1.5")
    point(doc, "CM is at L/2 = 0.6 m from each end; pivot is at 0.3 m from one end = 0.3 m from CM", "1")
    point(doc, "d = |0.6 − 0.3| = 0.3 m (from one end: 0.3 m; CM from that end: 0.6 m → d = 0.3 m)", "0.5")
    point(doc, "I_pivot = I_cm + md² = 0.36 + 3(0.09) = 0.36 + 0.27 = 0.63 kg·m²", "2")
    sep(doc)

    # ── Q1(c) ──
    heading(doc, "Q1 (c)  —  Collision [5 marks]")
    point(doc, "Given: m₁ = 2 kg, v₁ = 8 m/s; m₂ = 3 kg, v₂ = 0; e = 0.6")
    point(doc, "Momentum: 2(8) = 2v₁' + 3v₂' → 2v₁' + 3v₂' = 16", "0.5")
    point(doc, "Restitution: v₂' − v₁' = e(v₁ − v₂) = 0.6(8) = 4.8", "0.5")
    point(doc, "From eq2: v₂' = v₁' + 4.8. Sub into eq1: 2v₁' + 3(v₁' + 4.8) = 16 → 5v₁' = 1.6 → v₁' = 0.32 m/s", "1.5")
    point(doc, "v₂' = 0.32 + 4.8 = 5.12 m/s", "0.5")
    point(doc, "KE_before = ½(2)(64) = 64 J; KE_after = ½(2)(0.1024) + ½(3)(26.21) = 0.1024 + 39.32 = 39.42 J", "1")
    point(doc, "KE lost = 64 − 39.42 = 24.58 J; Impulse J = m₂v₂' = 3(5.12) = 15.36 N·s", "1")
    sep(doc)

    # ── Q1(d) ──
    heading(doc, "Q1 (d)  —  Block on Rough Surface [5 marks]")
    point(doc, "Given: m = 4 kg, μ_k = 0.35, μ_s = 0.4, F = 25 N horizontal")
    point(doc, "N = mg = 4(9.81) = 39.24 N", "0.5")
    point(doc, "f_s_max = μ_s N = 0.4(39.24) = 15.70 N", "1")
    point(doc, "Since F = 25 N > f_s_max = 15.70 N → block moves", "1")
    point(doc, "f_k = μ_k N = 0.35(39.24) = 13.73 N", "1")
    point(doc, "a = (F − f_k)/m = (25 − 13.73)/4 = 11.27/4 = 2.818 m/s²", "1.5")
    sep(doc)

    # ── Q2 ──
    heading(doc, "Q2  —  Rotating Machinery & Transmission [20 marks]")

    heading(doc, "Q2 (a)  —  Two-Stage Gear Train [5 marks]", 2)
    point(doc, "Given: GR₁ = 60/20 = 3, GR₂ = 90/18 = 5; n_in = 2400 RPM, τ_in = 1.2 N·m, η = 0.95/stage")
    point(doc, "GR_total = 3 × 5 = 15; η_total = 0.95² = 0.9025", "1")
    point(doc, "n_out = 2400/15 = 160 RPM", "1")
    point(doc, "τ_out = 1.2 × 15 × 0.9025 = 16.245 N·m", "1.5")
    point(doc, "P_in = 1.2 × 2π(2400)/60 = 1.2 × 251.33 = 301.6 W; P_loss = 301.6(1 − 0.9025) = 29.38 W", "1.5")

    heading(doc, "Q2 (b)  —  Flywheel Braking [5 marks]", 2)
    point(doc, "Given: m = 50 kg, R = 0.5 m, n₀ = 1200 RPM, t_stop = 30 s")
    point(doc, "I = ½(50)(0.25) = 6.25 kg·m²; ω₀ = 2π(1200)/60 = 125.66 rad/s", "1")
    point(doc, "α = ω₀/t = 125.66/30 = 4.189 rad/s²", "1")
    point(doc, "τ_brake = Iα = 6.25 × 4.189 = 26.18 N·m", "1")
    point(doc, "E = ½Iω₀² = ½(6.25)(15790.4) = 49345 J ≈ 49.35 kJ", "2")

    heading(doc, "Q2 (c)  —  Worm Gear Drive [5 marks]", 2)
    point(doc, "Given: GR = 40, η = 0.75, P_in = 2000 W, n_in = 1440 RPM")
    point(doc, "τ_motor = P_in/ω_in = 2000/(2π × 1440/60) = 2000/150.80 = 13.26 N·m", "1")
    point(doc, "n_out = 1440/40 = 36 RPM = 3.77 rad/s", "1")
    point(doc, "τ_out = τ_motor × GR × η = 13.26 × 40 × 0.75 = 397.8 N·m", "1.5")
    point(doc, "P_lost = P_in(1 − η) = 2000(0.25) = 500 W", "1.5")

    heading(doc, "Q2 (d)  —  Cylinder with Viscous Friction [5 marks]", 2)
    point(doc, "Given: m = 3 kg, R = 0.1 m, τ = 0.6 N·m, c = 0.2 N·m·s/rad")
    point(doc, "I = ½mR² = ½(3)(0.01) = 0.015 kg·m²", "1")
    point(doc, "ω_ss = τ/c = 0.6/0.2 = 3.0 rad/s", "1.5")
    point(doc, "Time constant τ_d = I/c = 0.015/0.2 = 0.075 s", "1.5")
    point(doc, "ω(τ_d) = ω_ss(1 − e⁻¹) = 3.0(0.632) = 1.896 rad/s", "1")
    sep(doc)

    # ── Q3 ──
    heading(doc, "Q3  —  Energy Methods & Circular Motion [20 marks]")

    heading(doc, "Q3 (a)  —  Roller Coaster [5 marks]", 2)
    point(doc, "Given: m = 800 kg, h₁ = 25 m, v₁ = 5 m/s, R_loop = 10 m, h_top = 20 m")
    point(doc, "½mv₁² + mgh₁ = ½mv_bot²; v_bot = √(v₁² + 2gh₁) = √(25 + 490.5) = √515.5 = 22.71 m/s", "1.5")
    point(doc, "At top of loop (h = 20 m): v_top = √(v₁² + 2g(h₁ − 20)) = √(25 + 2(9.81)(5)) = √(25 + 98.1) = √123.1 = 11.10 m/s", "1.5")
    point(doc, "N_top = mv²/R − mg = 800(123.1)/10 − 800(9.81) = 9848 − 7848 = 2000 N", "2")

    heading(doc, "Q3 (b)  —  Damped Vibration Analysis [5 marks]", 2)
    point(doc, "Given: m = 5 kg, x₁ = 25 mm (t₁ = 0.1 s), x₂ = 18 mm (t₂ = 0.38 s)")
    point(doc, "T_d = t₂ − t₁ = 0.38 − 0.1 = 0.28 s", "0.5")
    point(doc, "δ = ln(25/18) = ln(1.389) = 0.3285", "1")
    point(doc, "ζ = δ/√(4π² + δ²) = 0.3285/√(39.478 + 0.1079) = 0.3285/6.293 = 0.05220", "1")
    point(doc, "ω_d = 2π/T_d = 2π/0.28 = 22.44 rad/s; ω_n = ω_d/√(1 − ζ²) = 22.44/0.9986 = 22.47 rad/s", "1.5")
    point(doc, "k = mω_n² = 5(504.9) = 2524.5 N/m", "1")

    heading(doc, "Q3 (c)  —  Pendulum from Horizontal [5 marks]", 2)
    point(doc, "Given: m = 1.5 kg, L = 2.0 m, θ₀ = 90° (horizontal)")
    point(doc, "h = L(1 − cos 90°) = L = 2.0 m", "0.5")
    point(doc, "v_bottom = √(2gL) = √(2 × 9.81 × 2) = √39.24 = 6.264 m/s", "1.5")
    point(doc, "T_bottom = mg + mv²/L = 1.5(9.81) + 1.5(39.24)/2 = 14.715 + 29.43 = 44.145 N = 3mg", "2")
    point(doc, "T(θ) = 3mg cos θ (for release from horizontal). "
          "T = mg when 3 cos θ = 1 → θ = arccos(1/3) ≈ 70.5° from vertical.", "1")
    note(doc, "At the bottom T = 3mg; at the release point (θ = 90°, v = 0) T = 0; T = mg at θ ≈ 70.5°")

    heading(doc, "Q3 (d)  —  Vertical Jump [5 marks]", 2)
    point(doc, "Given: m = 70 kg, d = 0.3 m, F_avg = 2500 N, t_push = 0.15 s")
    point(doc, "F_net = F_avg − mg = 2500 − 70(9.81) = 2500 − 686.7 = 1813.3 N", "1")
    point(doc, "W_net = F_net × d = 1813.3 × 0.3 = 544.0 J = ½mv²", "1")
    point(doc, "v = √(2 × 544/70) = √15.54 = 3.943 m/s", "1")
    point(doc, "h = v²/(2g) = 15.54/19.62 = 0.792 m ≈ 79.2 cm", "1")
    point(doc, "P_avg = W_total/t = (F_avg × d)/t = (2500 × 0.3)/0.15 = 750/0.15 = 5000 W = 5 kW", "1")
    sep(doc)

    # ── Q4 ──
    heading(doc, "Q4  —  Rotating Frames & Rigid Body Properties [20 marks]")

    heading(doc, "Q4 (a)  —  Transport Theorem Derivation [5 marks]", 2)
    point(doc, "(dA/dt)_fixed = (dA/dt)_rot + Ω × A → applies to any vector A", "1")
    point(doc, "Velocity: v_fixed = v_rel + Ω × r + v_origin", "1")
    point(doc, "Acceleration: a_fixed = a_rel + 2Ω × v_rel + Ω × (Ω × r) + α × r + a_origin", "1")
    point(doc, "Coriolis (2Ω × v_rel): deflection due to velocity in rotating frame", "0.5")
    point(doc, "Centripetal (Ω × (Ω × r)): always radially inward, = −Ω²r", "0.5")
    point(doc, "Euler (α × r): exists only when angular velocity changes with time", "0.5")
    point(doc, "Relative (a_rel): acceleration measured by observer in the rotating frame", "0.5")

    heading(doc, "Q4 (b)  —  Radius of Gyration & Axis Theorems [5 marks]", 2)
    point(doc, "Radius of gyration: k = √(I/m) such that I = mk² — equivalent point mass distance", "1")
    point(doc, "Perpendicular Axis (2D bodies only): I_z = I_x + I_y (for a lamina in xy-plane)", "1")
    point(doc, "Parallel Axis: I_O = I_cm + md² where d = distance between parallel axes", "1")
    point(doc, "Application PAT: Finding MOI of a gear about its shaft when COM offset is known", "1")
    point(doc, "Application Perp: Finding I_z for a flat plate/disk from two perpendicular in-plane moments", "1")

    heading(doc, "Q4 (c)  —  Centrifuge [5 marks]", 2)
    point(doc, "Given: r = 2.5 m, n = 500 RPM, ṙ = 0.05 m/s")
    point(doc, "ω = 2π(500)/60 = 52.36 rad/s", "0.5")
    point(doc, "a_centripetal = ω²r = (52.36)²(2.5) = 2741.6 × 2.5 = 6853.9 m/s²", "1.5")
    point(doc, "a_Coriolis = 2ωṙ = 2(52.36)(0.05) = 5.236 m/s²", "1")
    point(doc, "|a_total| = √(6853.9² + 5.236²) ≈ 6853.9 m/s² (Coriolis negligible)", "1")
    point(doc, "a_c/g = 6853.9/9.81 = 698.7 g", "1")

    heading(doc, "Q4 (d)  —  Compound Object MOI [5 marks]", 2)
    point(doc, "Given: Disk m₁ = 4 kg, R = 0.15 m; Rod m₂ = 2 kg, L = 0.6 m, attached at disk edge")
    point(doc, "I_disk about its center: ½m₁R² = ½(4)(0.0225) = 0.045 kg·m²", "1")
    point(doc, "Rod CM is at R + L/2 = 0.15 + 0.30 = 0.45 m from disk center", "1")
    point(doc, "I_rod about disk center (PAT): m₂L²/12 + m₂(0.45)² = 2(0.36)/12 + 2(0.2025) = 0.06 + 0.405 = 0.465 kg·m²", "2")
    point(doc, "I_total = 0.045 + 0.465 = 0.510 kg·m²", "1")
    sep(doc)

    # ── Q5 ──
    heading(doc, "Q5  —  Pendulums & Lagrangian Mechanics [20 marks]")

    heading(doc, "Q5 (a)  —  Simple vs Physical Pendulum [5 marks]", 2)
    point(doc, "Simple pendulum: T = 2π√(L/g) — point mass m on massless string of length L", "1")
    point(doc, "Physical pendulum: T = 2π√(I_O/(mgd)) — rigid body, I_O = MOI about pivot, d = pivot-to-CM distance", "1.5")
    point(doc, "Derivation: τ = I_O α → −mgd sin θ = I_O θ̈ → θ̈ + (mgd/I_O)θ = 0 for small angles", "1")
    point(doc, "Same period when: L_equiv = I_O/(md) = L → if I_O/(md) equals the simple pendulum length", "1.5")

    heading(doc, "Q5 (b)  —  Normal Modes Theory [5 marks]", 2)
    point(doc, "Normal mode: a pattern where all parts oscillate at the same frequency with fixed amplitude ratios", "1")
    point(doc, "Eigenvalues ω² from det([K] − ω²[M]) = 0 → give natural frequencies", "1")
    point(doc, "Eigenvectors X from ([K] − ω²[M])X = 0 → give mode shapes (relative amplitudes)", "1")
    point(doc, "Mode 1 (lowest ω): typically both masses move in-phase (same direction)", "1")
    point(doc, "Mode 2 (highest ω): typically masses move out-of-phase (opposite directions)", "1")

    heading(doc, "Q5 (c)  —  U-Tube Manometer via Lagrange [5 marks]", 2)
    point(doc, "Let x = displacement from equilibrium. Total liquid mass m = ρAL_total with L_total = 0.6 m", "0.5")
    point(doc, "T = ½mẋ² (entire column moves at ẋ)", "1")
    point(doc, "V = ρA g x × x = ρAg x² (one side rises x, other drops x → net restoring head = 2x)", "1")
    note(doc, "More precisely: V = 2ρAgx² (from both columns). L = ½mẋ² − ρAgx²")
    point(doc, "EOM: mẍ + 2ρAgx = 0 → ẍ + (2g/L_total)x = 0, since m = ρA × L_total", "1")
    point(doc, "ω_n = √(2g/L) = √(2 × 9.81/0.6) = √32.7 = 5.718 rad/s; T = 2π/5.718 = 1.099 s", "1.5")

    heading(doc, "Q5 (d)  —  2-DOF System [5 marks]", 2)
    point(doc, "Given: m₁ = 2, m₂ = 3 kg; k₁ = 200, k₂ = 300, k₃ = 200 N/m")
    point(doc, "[M] = [[2, 0], [0, 3]]; [K] = [[500, −300], [−300, 500]]", "1.5")
    point(doc, "det([K] − ω²[M]) = (500 − 2ω²)(500 − 3ω²) − 90000 = 0", "0.5")
    point(doc, "6ω⁴ − 2500ω² + 160000 = 0; ω² = (2500 ± √(6250000 − 3840000))/12 = (2500 ± √2410000)/12", "1")
    point(doc, "√2410000 = 1552.4; ω₁² = (2500 − 1552.4)/12 = 78.97 → ω₁ = 8.887 rad/s", "0.5")
    point(doc, "ω₂² = (2500 + 1552.4)/12 = 337.7 → ω₂ = 18.377 rad/s", "0.5")
    point(doc, "Mode 1: in-phase; Mode 2: out-of-phase", "1")
    sep(doc)

    # ── Q6 ──
    heading(doc, "Q6  —  FBD + Difficult Numerical [20 marks]")

    heading(doc, "Q6 (a)  —  Package on Conveyor Belt [10 marks]", 2)
    point(doc, "Given: m = 5 kg, v_belt = 4.0 m/s, v_pkg = 1.0 m/s, μ_s = 0.4, μ_k = 0.3")
    marks_note(doc, "FBD [3 marks]:")
    point(doc, "Weight W = 5(9.81) = 49.05 N ↓; Normal N = 49.05 N ↑", "1")
    point(doc, "Kinetic friction f_k = μ_k N = 0.3(49.05) = 14.715 N → (forward — belt drags package)", "1")
    point(doc, "Package is slower than belt → relative motion of package w.r.t. belt is backward → friction forward ", "1")
    marks_note(doc, "Calculations [7 marks]:")
    point(doc, "a = f_k/m = 14.715/5 = 2.943 m/s²", "1")
    point(doc, "t = Δv/a = (4.0 − 1.0)/2.943 = 3.0/2.943 = 1.019 s", "1.5")
    point(doc, "s_pkg = v₁t + ½at² = 1.0(1.019) + ½(2.943)(1.038) = 1.019 + 1.528 = 2.547 m", "1.5")
    point(doc, "s_belt = v_belt × t = 4.0(1.019) = 4.076 m", "0.5")
    point(doc, "Relative sliding = s_belt − s_pkg = 4.076 − 2.547 = 1.529 m", "0.5")
    point(doc, "Heat = f_k × (relative sliding) = 14.715 × 1.529 = 22.50 J", "2")
    note(doc, "Check: ΔKE_pkg = ½(5)(16 − 1) = 37.5 J; W_friction_on_pkg = f_k × s_pkg = 14.715 × 2.547 = 37.49 J ")

    heading(doc, "Q6 (b)  —  Forced Vibration & Transmissibility [10 marks]", 2)
    point(doc, "Given: m = 100 kg, k = 40000 N/m, c = 800 N·s/m, n = 1500 RPM, F₀ = 500 N")
    point(doc, "ω_n = √(k/m) = √(400) = 20 rad/s; f_n = 3.183 Hz", "1")
    point(doc, "c_cr = 2mω_n = 2(100)(20) = 4000 N·s/m; ζ = 800/4000 = 0.2", "1")
    point(doc, "ω = 2π(1500)/60 = 157.08 rad/s; r = ω/ω_n = 157.08/20 = 7.854", "1")
    point(doc, "Amplitude X = F₀/k / √((1−r²)² + (2ζr)²)", "0.5")
    point(doc, "= (500/40000) / √((1 − 61.69)² + (2 × 0.2 × 7.854)²)", "0.5")
    point(doc, "= 0.0125 / √(3682.7 + 9.87) = 0.0125 / √3692.6 = 0.0125 / 60.77 = 0.0002057 m ≈ 0.206 mm", "2")
    point(doc, "Transmissibility: TR = √(1 + (2ζr)²) / √((1−r²)² + (2ζr)²) = √(1 + 9.87) / 60.77 = √10.87/60.77 = 3.297/60.77 = 0.0543", "2")
    point(doc, "F_transmitted = TR × F₀ = 0.0543 × 500 = 27.1 N", "2")
    note(doc, "At r >> 1, system provides excellent isolation — only 5.4% of force transmitted")
    sep(doc)

    # ── Q7 ──
    heading(doc, "Q7  —  FBD + Difficult Numerical [20 marks]")

    heading(doc, "Q7 (a)  —  Two-Link Robot Arm [10 marks]", 2)
    point(doc, "Given: L₁ = 0.5 m, m₁ = 3 kg, ω₁ = 2 rad/s, α₁ = 0; L₂ = 0.4 m, m₂ = 2 kg, ω₂_rel = 3 rad/s, α₂_rel = 1 rad/s²")
    point(doc, "At θ₁ = 0°, θ₂ = 90° (link 2 perpendicular to link 1)")
    marks_note(doc, "FBD of tip mass (1 kg) [3 marks]:")
    point(doc, "Weight W = 1(9.81) = 9.81 N (downward, out of horizontal plane → irrelevant for planar analysis if horizontal)", "1")
    point(doc, "Joint reaction force at link 2 tip → from link 2 structure", "1")
    point(doc, "Centripetal and Coriolis forces in the rotating frame of link 1", "1")
    marks_note(doc, "Velocity calculation [3 marks]:")
    point(doc, "Tip of link 1 (joint 2): v₁ = ω₁ × L₁ = 2 × 0.5 = 1.0 m/s (ĵ direction, if link 1 along x̂)", "1")
    point(doc, "Link 2 at 90° from link 1 → extends in ĵ direction; ω₂_abs = ω₁ + ω₂_rel = 2 + 3 = 5 rad/s", "0.5")
    point(doc, "v_tip_rel = ω₂_abs × L₂ component + v_joint: v_tip = v₁ + ω₂_abs × r₂", "0.5")
    point(doc, "v₂ contribution: 5 × 0.4 = 2.0 m/s (perpendicular to link 2, i.e. in −x̂ direction since link 2 is along ĵ)", "0.5")
    point(doc, "v_tip = (−2.0 î + 1.0 ĵ) m/s; |v_tip| = √(4 + 1) = √5 = 2.236 m/s", "0.5")
    marks_note(doc, "Acceleration calculation [4 marks]:")
    point(doc, "a_joint2 = −ω₁²L₁ (centripetal, toward pivot) = −4(0.5) = −2.0 m/s² (−x̂ direction)", "1")
    point(doc, "a_tip centripetal (about joint 2): −ω₂_abs²L₂ (toward joint 2) = −25(0.4) = −10.0 m/s² (−ĵ)", "1")
    point(doc, "a_tip tangential (about joint 2): α₂_abs L₂ where α₂_abs includes frame effects", "0.5")
    point(doc, "Coriolis contribution: 2ω₁ × v₂_rel = 2(2)(3 × 0.4) = 4.8 m/s² (perpendicular component)", "0.5")
    point(doc, "Total acceleration is the vector sum of all components — magnitude ≈ √(centripetal² + tangential² + Coriolis²)", "1")
    note(doc, "Full vector acceleration requires careful bookkeeping of rotating frame terms — detailed calculation depends on convention")

    heading(doc, "Q7 (b)  —  Scotch Yoke Mechanism [10 marks]", 2)
    point(doc, "Given: r = 40 mm = 0.04 m, n = 1800 RPM, m_slider = 0.5 kg")
    point(doc, "ω = 2π(1800)/60 = 188.50 rad/s", "1")
    point(doc, "x(t) = r cos(ωt); v(t) = −rω sin(ωt); a(t) = −rω² cos(ωt)", "1")
    point(doc, "x_max = r = 40 mm = 0.04 m", "0.5")
    point(doc, "v_max = rω = 0.04 × 188.50 = 7.540 m/s", "1.5")
    point(doc, "a_max = rω² = 0.04 × 35532.25 = 1421.3 m/s² = 144.9 g", "1.5")
    point(doc, "F_inertia_max = m × a_max = 0.5 × 1421.3 = 710.65 N", "1")
    point(doc, "P_max = F × v at position of max acceleration: at x = ±r, v = 0 → power = 0 at that instant", "1")
    note(doc, "Max power occurs at maximum velocity (midstroke): P = F × v, but F = 0 at midstroke for pure inertia load. "
         "Power is max when F × v peaks, at θ = 45°/135°: v = rω sin 45° = 5.332 m/s, a = rω² cos 45° = 1005 m/s², "
         "P = m × a × v = 0.5 × 1005 × 5.332 = 2679 W")
    point(doc, "At θ = 45°: x = r cos 45° = 28.28 mm; v = rω sin 45° = 5.332 m/s", "2.5")
    sep(doc)


# ═══════════════ MAIN ═══════════════

if __name__ == "__main__":
    out = os.path.join(SCRIPT_DIR, "Re-Exam")
    os.makedirs(out, exist_ok=True)
    sets = {1: answers_set_r1, 2: answers_set_r2}

    for num, writer in sets.items():
        doc = new_doc(num)
        writer(doc)
        path = os.path.join(out, f"DSM_ReExam_Answers_SetR{num}.docx")
        doc.save(path)
        print(f"[OK] {path}")

    print("\nDone — 2 DSM Re-Exam answer booklets generated.")
