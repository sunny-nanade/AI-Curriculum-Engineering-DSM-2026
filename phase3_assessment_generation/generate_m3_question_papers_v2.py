"""
M3 (Re-Examination) Question Paper Generator — Dynamic Systems Modeling
=======================================================================
Generates 2 sets of QP + 2 sets of MODEL ANSWERS in Word (.docx) format,
plus separate QP PDF and Solutions PDF.

Syllabus Scope : U1 – U5 (complete)
Marking Scheme : Q1 compulsory (4 marks) + Solve any 2 from Q2/Q3/Q4 (3 marks each) = 10 marks
Duration       : 45 minutes

Question Design (balanced across all 5 units):
  Q1(a) [2M, 4×0.5] — Pure theory spanning U1-U5
  Q1(b) [2M, 4×0.5] — Mixed theory + short numerical spanning U1-U5
  Q2    [3M, 6×0.5] — Lagrangian Mechanics (U2) — key topic (easy, formulas given)
  Q3    [3M, 6×0.5] — Rotating Reference Frames / Coriolis (U3)
  Q4    [3M, 6×0.5] — Vibrations / Rigid Body Kinetics (U5)

Course Outcomes
  CO1 — Describe theoretical concepts in mathematical modelling
  CO2 — Analyze dynamic systems using constitutive laws
  CO3 — Synthesize dynamic systems via modelling, simulation, and control design
"""

import os, math
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle,
)
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


# ═══════════════════════ CONFIGURATION ═══════════════════════

INSTITUTION = "Mukesh Patel School of Technology Management & Engineering"
PROGRAM     = "B.Tech Mechatronics Engineering, Semester IV"
COURSE      = "Dynamic Systems Modeling (702MH0C023)"
EXAM_TITLE  = "Re-Examination"
EXAM_DATE   = ""            # to be filled by the instructor
DURATION    = "45 minutes"
TOTAL_MARKS = "10"
NUM_SETS    = 2
OUTPUT_DIR  = os.path.join(os.path.dirname(os.path.abspath(__file__)), "question_papers_m3")

g = 9.81


# ═══════════════════════ QUESTION + SOLUTION BANKS ═══════════

# ─── Q1(a) Pure Theory ──────────────────────────────────────
# Bloom L1–L2 (Remember / Understand)  |  CO1
# 4 sub-questions × 0.5 marks = 2 marks

Q1A_BANKS = {
    1: {
        "questions": [
            # U1 — Coordinate systems
            "Write the expressions for velocity in polar coordinates "
            "(v_r and v_θ components) and state when polar coordinates "
            "are preferred over Cartesian.",

            # U2 — Lagrangian
            "Define the Lagrangian L of a mechanical system and state "
            "the Euler-Lagrange equation in terms of a generalized "
            "coordinate q.",

            # U3 — Rotating frames
            "State the transport theorem for a vector quantity as observed "
            "from a fixed (inertial) frame and a rotating frame. "
            "Identify the correction term.",

            # U5 — Vibrations
            "Define the damping ratio ζ and state the three response "
            "regimes (underdamped, critically damped, overdamped) with "
            "their conditions.",
        ],
        "solutions": [
            "In polar coordinates:\n"
            "  v_r = ṙ  (radial component)\n"
            "  v_θ = rθ̇  (transverse component)\n"
            "  → v⃗ = ṙ ê_r + rθ̇ ê_θ\n"
            "Polar coordinates are preferred when the motion has radial "
            "symmetry or when dealing with central-force problems "
            "(e.g. planetary motion, pendulums), because the constraint "
            "is naturally expressed in r and θ.",

            "The Lagrangian is defined as L = T − V, where T = kinetic "
            "energy and V = potential energy.\n"
            "The Euler-Lagrange equation is:\n"
            "  d/dt (∂L/∂q̇) − ∂L/∂q = 0\n"
            "This yields the equation(s) of motion for the system in "
            "terms of the generalized coordinate q.",

            "The transport theorem states:\n"
            "  (dA⃗/dt)_fixed = (dA⃗/dt)_rotating + Ω⃗ × A⃗\n"
            "where Ω⃗ is the angular velocity of the rotating frame.\n"
            "The correction term Ω⃗ × A⃗ accounts for the rotation of "
            "the basis vectors themselves as seen from the fixed frame.",

            "The damping ratio ζ = c / c_cr = c / (2√(km)), where c = "
            "damping coefficient, k = stiffness, m = mass.\n"
            "Three regimes:\n"
            "  ζ < 1 → Underdamped: oscillatory with exponentially "
            "decaying amplitude.\n"
            "  ζ = 1 → Critically damped: fastest return to equilibrium "
            "without overshoot.\n"
            "  ζ > 1 → Overdamped: slow, non-oscillatory return to "
            "equilibrium.",
        ],
    },
    2: {
        "questions": [
            # U1 — Acceleration in polar coordinates
            "Write the expressions for the radial and transverse "
            "components of acceleration in polar coordinates "
            "(a_r and a_θ) and identify the centripetal term.",

            # U2 — Energy methods
            "State Hamilton's principle (principle of least action) "
            "and explain its relationship to the Euler-Lagrange equation.",

            # U4 — Constraints
            "Differentiate between holonomic and non-holonomic constraints. "
            "Give one example of each from a mechanical system.",

            # U5 — Rigid body
            "State the parallel axis theorem for moment of inertia "
            "and give its mathematical expression.",
        ],
        "solutions": [
            "In polar coordinates the acceleration components are:\n"
            "  a_r = r̈ − rθ̇²  (radial component)\n"
            "  a_θ = rθ̈ + 2ṙθ̇  (transverse component)\n"
            "The term −rθ̇² in a_r is the centripetal acceleration, "
            "directed radially inward. The term 2ṙθ̇ in a_θ is "
            "the Coriolis component arising from simultaneous "
            "radial and angular motion.",

            "Hamilton's principle states that the actual path taken by "
            "a mechanical system between two time instants t₁ and t₂ "
            "is the one that makes the action integral stationary:\n"
            "  δS = δ ∫(t₁ to t₂) L dt = 0, where L = T − V.\n"
            "Applying the calculus of variations to this condition "
            "yields the Euler-Lagrange equation:\n"
            "  d/dt(∂L/∂q̇) − ∂L/∂q = 0.",

            "Holonomic constraint: can be expressed as f(q₁, q₂, …, t) = 0 "
            "(depends only on coordinates and time, not velocities).\n"
            "Example: A bead on a wire — its position is restricted to "
            "the wire curve.\n"
            "Non-holonomic constraint: involves velocities or inequalities "
            "and cannot be reduced to a coordinate-only equation.\n"
            "Example: A wheel rolling without slipping — the constraint "
            "dx = R dθ relates velocity components but cannot be "
            "integrated into a position-only equation for all paths.",

            "The parallel axis theorem states that the moment of inertia "
            "about any axis parallel to one through the center of mass is:\n"
            "  I = I_cm + md²\n"
            "where I_cm = moment of inertia about the CM axis, "
            "m = total mass, d = perpendicular distance between "
            "the two axes. I is always smallest about the CM axis.",
        ],
    },
}


# ─── Q1(b) Mixed Theory + Short Numerical ──────────────────
# Bloom L2–L3 (Understand / Apply)  |  CO1, CO2
# 4 sub-questions × 0.5 marks = 2 marks

Q1B_BANKS = {
    1: {
        "questions": [
            # U1 — short numerical (coordinate transform)
            "A particle moves in a circle of radius r = 2 m at constant "
            "angular speed θ̇ = 3 rad/s. Calculate the radial (centripetal) "
            "acceleration a_r and the transverse acceleration a_θ.",

            # U2 — conceptual (Lagrangian concept)
            "Write the kinetic energy T and potential energy V for a "
            "simple pendulum of mass m and length L, using θ as the "
            "generalized coordinate.",

            # U4 — conceptual (constraints)
            "Define a holonomic constraint and state how the number of "
            "degrees of freedom (DOF) is determined for a system of N "
            "particles with k holonomic constraints.",

            # U5 — short numerical (natural frequency)
            "A spring-mass system has m = 4 kg and k = 256 N/m. "
            "Calculate the natural frequency ω_n (rad/s) and the "
            "time period T of free oscillation.",
        ],
        "solutions": [
            "For circular motion at constant angular speed (ṙ = 0, r̈ = 0, θ̈ = 0):\n"
            "  a_r = r̈ − rθ̇² = 0 − 2 × 3² = −18 m/s²  (directed radially inward)\n"
            "  a_θ = rθ̈ + 2ṙθ̇ = 0 + 0 = 0 m/s²\n"
            "The magnitude of centripetal acceleration is 18 m/s².",

            "For a simple pendulum with angle θ:\n"
            "  T = ½ m (Lθ̇)² = ½ m L² θ̇²\n"
            "  V = −mgL cos θ  (taking the pivot as the datum)\n"
            "  or equivalently V = mgL(1 − cos θ) with datum at the "
            "lowest point.\n"
            "The Lagrangian is L = T − V = ½mL²θ̇² + mgL cos θ.",

            "A holonomic constraint is one that can be expressed as an "
            "equation relating only coordinates (and possibly time), "
            "not velocities: f(q₁, q₂, …, t) = 0.\n"
            "DOF = 3N − k, where N = number of particles, k = number "
            "of independent holonomic constraints.\n"
            "Example: A bead on a wire (k = 2 of 3 coordinates "
            "constrained → 1 DOF).",

            "ω_n = √(k/m) = √(256/4) = √64 = 8 rad/s.\n"
            "T = 2π/ω_n = 2π/8 = π/4 ≈ 0.785 s.",
        ],
    },
    2: {
        "questions": [
            # U1 — short numerical (polar acceleration)
            "A particle moves along a spiral with r = 3 m, ṙ = 0, "
            "and θ̇ = 2 rad/s (constant). Calculate the centripetal "
            "acceleration and state its direction.",

            # U2 — conceptual (generalized coordinates)
            "Explain the concept of generalized coordinates. For a "
            "double pendulum in a plane, state the number of DOF "
            "and identify the generalized coordinates.",

            # U3 — conceptual (Coriolis)
            "Define the Coriolis acceleration. State the expression "
            "for its magnitude and explain in which physical "
            "situations it becomes significant.",

            # U5 — short numerical (moment of inertia)
            "A uniform rod of mass m = 3 kg and length L = 1.2 m "
            "is pivoted at one end. Calculate its moment of inertia "
            "about the pivot using the parallel axis theorem.",
        ],
        "solutions": [
            "For circular motion with r = 3 m, ṙ = 0, θ̇ = 2 rad/s:\n"
            "  a_centripetal = rθ̇² = 3 × 2² = 3 × 4 = 12 m/s²\n"
            "  (In polar form: a_r = r̈ − rθ̇² = 0 − 12 = −12 m/s²)\n"
            "Direction: radially inward, toward the center of curvature.",

            "Generalized coordinates are the minimum set of independent "
            "coordinates needed to completely describe a system's "
            "configuration. They need not be Cartesian — they can be "
            "angles, arc lengths, etc.\n"
            "A double pendulum in a plane has 2 DOF.\n"
            "Generalized coordinates: θ₁ (angle of first link) and "
            "θ₂ (angle of second link), both measured from the vertical.",

            "The Coriolis acceleration arises when an object moves "
            "with velocity v_rel in a reference frame rotating at "
            "angular velocity Ω:\n"
            "  a_Coriolis = 2Ω × v_rel,  magnitude = 2Ωv_rel\n"
            "It is perpendicular to both Ω and v_rel. It becomes "
            "significant in large-scale meteorology (weather systems), "
            "long-range ballistics, and fast-rotating machinery "
            "where the product Ωv_rel is comparable to other "
            "acceleration terms.",

            "For a uniform rod: I_cm = (1/12)mL² = (1/12) × 3 × 1.2² "
            "= (1/12) × 3 × 1.44 = 0.36 kg·m².\n"
            "Parallel axis theorem: I_pivot = I_cm + m(L/2)²\n"
            "= 0.36 + 3 × (0.6)² = 0.36 + 3 × 0.36 = 0.36 + 1.08\n"
            "= 1.44 kg·m².\n"
            "(Check: I_pivot = (1/3)mL² = (1/3) × 3 × 1.44 = 1.44 kg·m² ✓)",
        ],
    },
}


# ─── Q2: Lagrangian Mechanics (U2) ─────────────────────────
# Bloom L3 (Apply)  |  CO1, CO2
# 6 sub-questions × 0.5 marks = 3 marks
# Formulas given in stem (easy question)

Q2_BANKS = {
    1: {
        "stem": (
            "A simple pendulum consists of a point mass m = 1 kg attached "
            "to a massless, inextensible string of length L = 1 m. "
            "The bob swings in a vertical plane under gravity (g = 9.81 m/s²). "
            "Using the angle θ from the vertical as the generalized "
            "coordinate, analyze the system using the Lagrangian method.\n"
            "Use the following relations: Lagrangian L = T − V; "
            "Euler-Lagrange equation: d/dt(∂L/∂q̇) − ∂L/∂q = 0; "
            "for small angles: sin θ ≈ θ."
        ),
        "bloom": "BL 3 (Apply)", "co": "CO1, CO2",
        "m": 1.0, "L": 1.0,
        "subs": [
            "State the number of degrees of freedom of this system "
            "and identify the generalized coordinate.",
            "Write the position of the mass in Cartesian coordinates "
            "(x, y) in terms of θ and L.",
            "Derive the kinetic energy T in terms of θ, θ̇, m, and L.",
            "Write the potential energy V in terms of θ, m, g, and L.",
            "Form the Lagrangian L = T − V and apply the Euler-Lagrange "
            "equation to derive the equation of motion.",
            "For small oscillations (sin θ ≈ θ), determine the natural "
            "frequency of the pendulum.",
        ],
    },
    2: {
        "stem": (
            "A simple pendulum consists of a point mass m = 2 kg attached "
            "to a massless, inextensible string of length L = 0.5 m. "
            "The bob swings in a vertical plane under gravity (g = 9.81 m/s²). "
            "Using the angle θ from the vertical as the generalized "
            "coordinate, analyze the system using the Lagrangian method.\n"
            "Use the following relations: Lagrangian L = T − V; "
            "Euler-Lagrange equation: d/dt(∂L/∂q̇) − ∂L/∂q = 0; "
            "for small angles: sin θ ≈ θ."
        ),
        "bloom": "BL 3 (Apply)", "co": "CO1, CO2",
        "m": 2.0, "L": 0.5,
        "subs": [
            "State the number of degrees of freedom of this system "
            "and identify the generalized coordinate.",
            "Write the position of the mass in Cartesian coordinates "
            "(x, y) in terms of θ and L.",
            "Derive the kinetic energy T in terms of θ, θ̇, m, and L.",
            "Write the potential energy V in terms of θ, m, g, and L.",
            "Form the Lagrangian L = T − V and apply the Euler-Lagrange "
            "equation to derive the equation of motion.",
            "For small oscillations (sin θ ≈ θ), determine the natural "
            "frequency of the pendulum.",
        ],
    },
}


def solve_q2(set_num):
    """Compute Lagrangian pendulum solutions."""
    b = Q2_BANKS[set_num]
    m, L_p = b["m"], b["L"]
    omega_n = math.sqrt(g / L_p)
    T_period = 2 * math.pi / omega_n
    return [
        "The pendulum has 1 degree of freedom (DOF = 1).\n"
        "Generalized coordinate: θ — the angle of the string "
        "measured from the vertical.",

        f"Taking the pivot as the origin, with x horizontal and y "
        f"vertically downward:\n"
        f"  x = L sin θ = {L_p} sin θ\n"
        f"  y = L cos θ = {L_p} cos θ\n"
        f"Differentiating: ẋ = Lθ̇ cos θ,  ẏ = −Lθ̇ sin θ",

        f"v² = ẋ² + ẏ² = L²θ̇²cos²θ + L²θ̇²sin²θ = L²θ̇²\n"
        f"T = ½mv² = ½ × {m} × {L_p}² × θ̇² = ½mL²θ̇²\n"
        f"T = ½({m})({L_p}²)θ̇² = {0.5 * m * L_p**2} θ̇²  J",

        f"Taking the pivot as the datum (y positive downward):\n"
        f"V = −mgy = −mgL cos θ = −{m} × {g} × {L_p} × cos θ\n"
        f"V = −{m*g*L_p:.2f} cos θ  J",

        f"L = T − V = ½mL²θ̇² + mgL cos θ\n"
        f"Euler-Lagrange: d/dt(∂L/∂θ̇) − ∂L/∂θ = 0\n"
        f"  ∂L/∂θ̇ = mL²θ̇ = {m} × {L_p}² × θ̇ = {m * L_p**2} θ̇\n"
        f"  d/dt(mL²θ̇) = mL²θ̈ = {m * L_p**2} θ̈\n"
        f"  ∂L/∂θ = −mgL sin θ = −{m*g*L_p:.2f} sin θ\n"
        f"EOM: {m * L_p**2} θ̈ + {m*g*L_p:.2f} sin θ = 0\n"
        f"Simplifying (divide by mL²): θ̈ + (g/L) sin θ = 0\n"
        f"  → θ̈ + ({g}/{L_p}) sin θ = 0  →  θ̈ + {g/L_p:.2f} sin θ = 0",

        f"For small angles, sin θ ≈ θ → θ̈ + (g/L)θ = 0\n"
        f"This is SHM with ω_n² = g/L = {g}/{L_p} = {g/L_p:.2f}\n"
        f"ω_n = √({g/L_p:.2f}) = {omega_n:.3f} rad/s\n"
        f"Time period T = 2π/ω_n = 2π/{omega_n:.3f} = {T_period:.3f} s",
    ]


# ─── Q3: Rotating Reference Frames / Coriolis (U3) ─────────
# Bloom L3–L4 (Apply / Analyze)  |  CO2
# 6 sub-questions × 0.5 marks = 3 marks

Q3_BANKS = {
    1: {
        "stem": (
            "A horizontal platform rotates at a constant angular velocity "
            "Ω = 4 rad/s about a vertical axis. An object at a radial "
            "distance r = 1.5 m from the center moves radially outward "
            "at a constant speed v_rel = 2 m/s (ṙ = 2 m/s, r̈ = 0). "
            "The angular acceleration of the platform is zero (α = 0)."
        ),
        "bloom": "BL 3–4 (Apply / Analyze)", "co": "CO2",
        "omega": 4.0, "r": 1.5, "v_rel": 2.0, "r2": 3.0,
        "subs": [
            "Write the general expression relating the acceleration observed "
            "from the fixed (inertial) frame to quantities measured in the "
            "rotating frame. Identify each term.",
            "Calculate the centripetal acceleration magnitude at r = 1.5 m.",
            "Calculate the Coriolis acceleration magnitude.",
            "State the direction of the Coriolis acceleration relative "
            "to the radial velocity.",
            "Calculate the total acceleration magnitude of the object "
            "as observed from the fixed (inertial) frame.",
            "If the object continues at the same radial speed, find "
            "the centripetal acceleration when r = 3 m and comment on "
            "how it changes.",
        ],
    },
    2: {
        "stem": (
            "A horizontal turntable rotates at a constant angular velocity "
            "Ω = 5 rad/s about a vertical axis. A small ball at r = 2 m "
            "from the center is pushed radially outward at a constant "
            "relative speed v_rel = 3 m/s (ṙ = 3 m/s, r̈ = 0). "
            "The angular acceleration of the turntable is zero (α = 0)."
        ),
        "bloom": "BL 3–4 (Apply / Analyze)", "co": "CO2",
        "omega": 5.0, "r": 2.0, "v_rel": 3.0, "r2": 4.0,
        "subs": [
            "Write the general expression relating the acceleration observed "
            "from the fixed (inertial) frame to quantities measured in the "
            "rotating frame. Identify each term.",
            "Calculate the centripetal acceleration magnitude at r = 2 m.",
            "Calculate the Coriolis acceleration magnitude.",
            "State the direction of the Coriolis acceleration relative "
            "to the radial velocity.",
            "Calculate the total acceleration magnitude of the object "
            "as observed from the fixed (inertial) frame.",
            "If the ball continues at the same radial speed, find "
            "the centripetal acceleration when r = 4 m and comment on "
            "how it changes.",
        ],
    },
}


def solve_q3(set_num):
    """Compute Coriolis / rotating frame solutions."""
    b = Q3_BANKS[set_num]
    omega, r_pos, v_rel, r2 = b["omega"], b["r"], b["v_rel"], b["r2"]
    a_cent = omega**2 * r_pos
    a_cor = 2 * omega * v_rel
    a_total = math.sqrt(a_cent**2 + a_cor**2)
    a_cent_2 = omega**2 * r2
    return [
        "The fixed-frame (inertial) acceleration, expressed in terms of "
        "rotating-frame quantities, is:\n"
        "  a⃗_fixed = a⃗_rel + 2Ω⃗ × v⃗_rel + α⃗ × r⃗ + Ω⃗ × (Ω⃗ × r⃗)\n"
        "where:\n"
        "  a⃗_rel = acceleration of the object in the rotating frame (r̈ = 0 here),\n"
        "  2Ω⃗ × v⃗_rel = Coriolis acceleration,\n"
        "  α⃗ × r⃗ = Euler (angular acceleration) term (zero here since α = 0),\n"
        "  Ω⃗ × (Ω⃗ × r⃗) = centripetal acceleration (directed radially inward).",

        f"a_centripetal = Ω²r = {omega}² × {r_pos} = {omega**2:.0f} × {r_pos} "
        f"= {a_cent:.1f} m/s²\n"
        f"Direction: radially inward (toward the axis of rotation).",

        f"a_Coriolis = 2Ωv_rel = 2 × {omega} × {v_rel} = {a_cor:.1f} m/s²",

        "By the right-hand rule for 2Ω⃗ × v⃗_rel (Ω along vertical, v_rel "
        "radially outward): the Coriolis acceleration is perpendicular to "
        "v_rel, directed tangentially in the direction opposite to the "
        "platform's rotation.",

        f"|a_total| = √(a_cent² + a_Cor²) = √({a_cent:.1f}² + {a_cor:.1f}²)\n"
        f"= √({a_cent**2:.0f} + {a_cor**2:.0f}) = √{a_cent**2 + a_cor**2:.0f} "
        f"= {a_total:.2f} m/s²",

        f"At r = {r2:.0f} m: a_centripetal = Ω²r = {omega}² × {r2:.0f} = {a_cent_2:.1f} m/s²\n"
        f"Compared to {a_cent:.1f} m/s² at r = {r_pos} m, the centripetal "
        f"acceleration has {'doubled' if abs(r2/r_pos - 2) < 0.01 else f'increased by a factor of {r2/r_pos:.1f}'}. "
        f"It increases linearly with radial "
        f"distance (a_cent ∝ r), while the Coriolis acceleration "
        f"({a_cor:.1f} m/s²) remains unchanged since v_rel is constant.",
    ]


# ─── Q4: Vibrations / Rigid Body Kinetics (U5) ─────────────
# Bloom L4–L5 (Analyze / Evaluate)  |  CO2, CO3
# 6 sub-questions × 0.5 marks = 3 marks

Q4_BANKS = {
    # Set 1 — Damped Vibrations (U5 L4)
    1: {
        "stem": (
            "A mass m = 3 kg is attached to a spring of stiffness "
            "k = 300 N/m and a viscous damper with damping coefficient "
            "c = 12 N·s/m. The system undergoes free vibration after "
            "being displaced 0.05 m from equilibrium and released from "
            "rest. Take the equilibrium position as the origin."
        ),
        "bloom": "BL 4–5 (Analyze / Evaluate)", "co": "CO2, CO3",
        "type": "damped_vibration",
        "m": 3.0, "k": 300.0, "c": 12.0, "x0": 0.05,
        "subs": [
            "Write the equation of motion for this damped spring-mass system.",
            "Calculate the natural frequency ω_n and the critical damping "
            "coefficient c_cr.",
            "Calculate the damping ratio ζ and state the type of damping.",
            "Calculate the damped natural frequency ω_d.",
            "Calculate the logarithmic decrement δ and find the ratio of "
            "two successive peak amplitudes.",
            "If the damper is removed (c = 0), write the displacement "
            "as a function of time x(t) for the given initial conditions.",
        ],
    },
    # Set 2 — Physical (Compound) Pendulum (U5 L1)
    2: {
        "stem": (
            "A uniform rod of mass m = 2 kg and length L = 1.2 m is "
            "pivoted at one end and oscillates as a physical (compound) "
            "pendulum under gravity. Take g = 9.81 m/s².\n"
            "Use: I_rod_cm = (1/12)mL²; parallel axis theorem: "
            "I = I_cm + md²; period of physical pendulum: "
            "T = 2π√(I_pivot / (mgd)), where d = distance from "
            "pivot to center of mass."
        ),
        "bloom": "BL 4–5 (Analyze / Evaluate)", "co": "CO2, CO3",
        "type": "physical_pendulum",
        "m": 2.0, "L": 1.2,
        "subs": [
            "Calculate the moment of inertia of the rod about its "
            "center of mass (I_cm).",
            "Using the parallel axis theorem, calculate the moment of "
            "inertia about the pivot at one end (I_pivot).",
            "Find the distance d of the center of mass from the pivot.",
            "Write the equation of motion for small-angle oscillations "
            "and determine the natural frequency ω_n.",
            "Calculate the time period of oscillation T.",
            "Compare this period with that of a simple pendulum of the "
            "same length (L = 1.2 m) and explain the difference.",
        ],
    },
}


def solve_q4(set_num):
    """Compute Q4 solutions."""
    b = Q4_BANKS[set_num]

    if b["type"] == "damped_vibration":
        m_val, k_val, c_val, x0 = b["m"], b["k"], b["c"], b["x0"]
        omega_n = math.sqrt(k_val / m_val)
        c_cr = 2 * math.sqrt(k_val * m_val)
        zeta = c_val / c_cr
        omega_d = omega_n * math.sqrt(1 - zeta**2)
        delta = 2 * math.pi * zeta / math.sqrt(1 - zeta**2)
        ratio = math.exp(-delta)
        return [
            f"The equation of motion for a damped SDOF system:\n"
            f"  mẍ + cẋ + kx = 0\n"
            f"  {m_val}ẍ + {c_val}ẋ + {k_val}x = 0\n"
            f"Dividing by m: ẍ + {c_val/m_val:.1f}ẋ + {k_val/m_val:.1f}x = 0\n"
            f"Or: ẍ + 2ζω_n ẋ + ω_n² x = 0",

            f"ω_n = √(k/m) = √({k_val}/{m_val}) = √{k_val/m_val:.0f} "
            f"= {omega_n:.2f} rad/s\n"
            f"c_cr = 2√(km) = 2√({k_val} × {m_val}) = 2√{k_val*m_val:.0f} "
            f"= 2 × {math.sqrt(k_val*m_val):.2f} = {c_cr:.2f} N·s/m",

            f"ζ = c / c_cr = {c_val} / {c_cr:.2f} = {zeta:.4f}\n"
            f"Since ζ = {zeta:.4f} < 1 → the system is UNDERDAMPED.\n"
            f"It will oscillate with exponentially decaying amplitude.",

            f"ω_d = ω_n √(1 − ζ²) = {omega_n:.2f} × √(1 − {zeta:.4f}²)\n"
            f"= {omega_n:.2f} × √(1 − {zeta**2:.6f})\n"
            f"= {omega_n:.2f} × {math.sqrt(1 - zeta**2):.6f}\n"
            f"= {omega_d:.4f} rad/s ≈ {omega_d:.2f} rad/s",

            f"δ = 2πζ / √(1 − ζ²) = 2π × {zeta:.4f} / √(1 − {zeta**2:.6f})\n"
            f"= {2*math.pi*zeta:.4f} / {math.sqrt(1 - zeta**2):.6f}\n"
            f"= {delta:.4f}\n"
            f"Ratio of successive peaks: x_n / x_(n+1) = e^δ = e^{delta:.4f} "
            f"= {1/ratio:.4f}\n"
            f"Or equivalently x_(n+1) / x_n = e^(−δ) = {ratio:.4f}\n"
            f"Each successive peak is ≈ {ratio*100:.2f}% of the previous one.",

            f"With c = 0 (undamped), EOM: mẍ + kx = 0\n"
            f"x(t) = x₀ cos(ω_n t) = {x0} cos({omega_n:.2f} t)  m\n"
            f"(Since initial conditions are x(0) = {x0} m, ẋ(0) = 0, "
            f"the solution is a pure cosine with amplitude {x0} m "
            f"and period T = 2π/{omega_n:.2f} = {2*math.pi/omega_n:.4f} s)",
        ]

    else:  # physical_pendulum
        m_val, L_val = b["m"], b["L"]
        I_cm = m_val * L_val**2 / 12
        d = L_val / 2
        I_piv = I_cm + m_val * d**2
        omega_n = math.sqrt(m_val * g * d / I_piv)
        T_phys = 2 * math.pi / omega_n
        T_simp = 2 * math.pi * math.sqrt(L_val / g)
        return [
            f"I_cm = (1/12)mL² = (1/12) × {m_val} × {L_val}²\n"
            f"= (1/12) × {m_val} × {L_val**2:.2f} = {I_cm:.4f} kg·m²",

            f"I_pivot = I_cm + m(L/2)² = {I_cm:.4f} + {m_val} × ({L_val}/2)²\n"
            f"= {I_cm:.4f} + {m_val} × {d}² = {I_cm:.4f} + {m_val * d**2:.4f}\n"
            f"= {I_piv:.4f} kg·m²\n"
            f"(Check: I_pivot = (1/3)mL² = (1/3) × {m_val} × {L_val**2:.2f} "
            f"= {m_val * L_val**2 / 3:.4f} kg·m² ✓)",

            f"d = L/2 = {L_val}/2 = {d} m\n"
            f"(The center of mass of a uniform rod is at its midpoint.)",

            f"Restoring torque: τ = −mgd sin θ ≈ −mgdθ  (small angle)\n"
            f"Equation of motion: I_pivot θ̈ + mgd θ = 0\n"
            f"  → θ̈ + (mgd / I_pivot) θ = 0\n"
            f"ω_n = √(mgd / I_pivot) = √({m_val} × {g} × {d} / {I_piv:.4f})\n"
            f"= √({m_val * g * d:.4f} / {I_piv:.4f}) = √{m_val * g * d / I_piv:.4f}\n"
            f"= {omega_n:.4f} rad/s ≈ {omega_n:.3f} rad/s",

            f"T = 2π / ω_n = 2π / {omega_n:.4f} = {T_phys:.4f} s ≈ {T_phys:.3f} s",

            f"Simple pendulum of same length L = {L_val} m:\n"
            f"T_simple = 2π√(L/g) = 2π√({L_val}/{g}) = 2π × {math.sqrt(L_val/g):.4f}\n"
            f"= {T_simp:.4f} s ≈ {T_simp:.3f} s\n"
            f"Physical pendulum period ({T_phys:.3f} s) < Simple pendulum period "
            f"({T_simp:.3f} s).\n"
            f"This is because the physical pendulum's mass is distributed along "
            f"its length, with some mass closer to the pivot than L. The effective "
            f"pendulum length L_eff = I_pivot/(md) = {I_piv:.4f}/({m_val}×{d}) "
            f"= {I_piv/(m_val*d):.4f} m < {L_val} m.",
        ]


# ═══════════════════════ WORD DOCUMENT BUILDERS ══════════════

def _set_no_borders(table):
    """Remove all borders from a Word table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        from docx.oxml import OxmlElement
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for child in list(tblPr):
        if child.tag == qn('w:tblBorders'):
            tblPr.remove(child)
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    ))


def _cell_run(cell, text, bold=False, size=Pt(11), first=True, align=None):
    """Add a formatted run to a cell paragraph."""
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = "Times New Roman"
    return p, run


def _build_qp_docx(set_num):
    """Build QP .docx in SVKM'S NMIMS template layout."""
    doc = Document()

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    q1a = Q1A_BANKS[set_num]
    q1b = Q1B_BANKS[set_num]
    q2  = Q2_BANKS[set_num]
    q3  = Q3_BANKS[set_num]
    q4  = Q4_BANKS[set_num]

    # ═══ TABLE 0 — Header ═══
    ht = doc.add_table(rows=6, cols=2)
    _set_no_borders(ht)
    for row in ht.rows:
        row.cells[0].width = Cm(8.3)
        row.cells[1].width = Cm(7.65)

    ht.cell(0, 0).merge(ht.cell(0, 1))
    c = ht.cell(0, 0)
    c.text = ""
    _cell_run(c, "SVKM'S NMIMS", bold=True, size=Pt(12),
              align=WD_ALIGN_PARAGRAPH.CENTER)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run("MUKESH PATEL SCHOOL OF TECHNOLOGY MANAGEMENT & ENGINEERING")
    r.bold = True; r.font.size = Pt(11); r.font.name = "Times New Roman"
    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(0)
    r = p3.add_run("Academic Year: 2025-2026")
    r.font.size = Pt(11); r.font.name = "Times New Roman"

    ht.cell(1, 0).text = ""
    _cell_run(ht.cell(1, 0), "Program: B.Tech    Stream: Mechatronics")
    ht.cell(1, 1).text = ""
    _cell_run(ht.cell(1, 1), "Year: 2nd   Semester: IV")

    ht.cell(2, 0).text = ""
    _cell_run(ht.cell(2, 0), f"Subject: {COURSE}")
    ht.cell(2, 1).text = ""
    _cell_run(ht.cell(2, 1), "Time: 45 Minutes")

    ht.cell(3, 0).text = ""
    _cell_run(ht.cell(3, 0), f"Date: {EXAM_DATE}")
    ht.cell(3, 1).text = ""
    _cell_run(ht.cell(3, 1), "No. of Pages: 1")

    ht.cell(4, 0).text = ""
    _cell_run(ht.cell(4, 0), f"Marks: {TOTAL_MARKS}")
    ht.cell(4, 1).text = ""
    _cell_run(ht.cell(4, 1), f"Set: {set_num}")

    ht.cell(5, 0).merge(ht.cell(5, 1))
    c = ht.cell(5, 0)
    c.text = ""
    _cell_run(c, "Re-Examination", bold=True, size=Pt(14),
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # ═══ INSTRUCTIONS ═══
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Instructions: Candidates should read carefully the instructions.")
    r.bold = True; r.font.size = Pt(11); r.font.name = "Times New Roman"

    for txt in [
        "1) Figures in brackets on the right hand side indicate full marks.",
        "2) Syllabus scope: Unit 1 – Unit 5 (complete).",
        "3) Assume suitable data if necessary.",
        "4) Question 1 is compulsory.",
        "5) Answer any 2 from the remaining questions.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(txt)
        r.font.size = Pt(11); r.font.name = "Times New Roman"

    # ═══ TABLE 1 — Questions ═══
    qt = doc.add_table(rows=6, cols=4)
    _set_no_borders(qt)
    for row in qt.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(0.65)
        row.cells[2].width = Cm(12.75)
        row.cells[3].width = Cm(1.05)

    # Row 0 — Q1 header
    qt.cell(0, 0).text = ""
    _cell_run(qt.cell(0, 0), "Q1", bold=True)
    qt.cell(0, 1).text = ""
    qt.cell(0, 2).text = ""
    _cell_run(qt.cell(0, 2), "Answer briefly:", size=Pt(12))
    qt.cell(0, 3).text = ""

    # Row 1 — Q1(a)
    c0 = qt.cell(1, 0); c0.text = ""
    _cell_run(c0, "CO- 1;", size=Pt(9))
    _cell_run(c0, "SO- 1;", size=Pt(9), first=False)
    _cell_run(c0, "BL- 1,2", size=Pt(9), first=False)
    qt.cell(1, 1).text = ""
    _cell_run(qt.cell(1, 1), "a.")
    c2 = qt.cell(1, 2); c2.text = ""
    for i, q_text in enumerate(q1a["questions"]):
        roman = ["(i)", "(ii)", "(iii)", "(iv)"][i]
        _cell_run(c2, f"{roman} {q_text}", first=(i == 0))
    qt.cell(1, 3).text = ""
    _cell_run(qt.cell(1, 3), "[2]")

    # Row 2 — Q1(b)
    c0 = qt.cell(2, 0); c0.text = ""
    _cell_run(c0, "CO- 1,2;", size=Pt(9))
    _cell_run(c0, "SO- 1,2;", size=Pt(9), first=False)
    _cell_run(c0, "BL- 2,3", size=Pt(9), first=False)
    qt.cell(2, 1).text = ""
    _cell_run(qt.cell(2, 1), "b.")
    c2 = qt.cell(2, 2); c2.text = ""
    for i, q_text in enumerate(q1b["questions"]):
        roman = ["(i)", "(ii)", "(iii)", "(iv)"][i]
        _cell_run(c2, f"{roman} {q_text}", first=(i == 0))
    qt.cell(2, 3).text = ""
    _cell_run(qt.cell(2, 3), "[2]")

    # Rows 3–5 — Q2, Q3, Q4
    q_meta = [
        (3, q2, "Q2", "CO- 1,2;", "SO- 1,2;", "BL- 3"),
        (4, q3, "Q3", "CO- 2;",   "SO- 2;",   "BL- 3,4"),
        (5, q4, "Q4", "CO- 2,3;", "SO- 2,3;", "BL- 4,5"),
    ]
    for ri, qb, ql, co_t, so_t, bl_t in q_meta:
        c0 = qt.cell(ri, 0); c0.text = ""
        _cell_run(c0, ql, bold=True)
        _cell_run(c0, co_t, size=Pt(9), first=False)
        _cell_run(c0, so_t, size=Pt(9), first=False)
        _cell_run(c0, bl_t, size=Pt(9), first=False)
        qt.cell(ri, 1).text = ""
        c2 = qt.cell(ri, 2); c2.text = ""
        _cell_run(c2, qb["stem"])
        for j, sub in enumerate(qb["subs"]):
            letter = chr(ord("a") + j)
            _cell_run(c2, f"({letter}) {sub}", first=False)
        qt.cell(ri, 3).text = ""
        _cell_run(qt.cell(ri, 3), "[3]")

    return doc


# ─── Solution Builder ───────────────────────────────────────

def _add_header_sol(doc, set_num):
    """Add institution header for solution document."""
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.3)
        section.right_margin = Cm(1.3)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(9.5)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(INSTITUTION)
    run.bold = True; run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(PROGRAM)
    run.font.size = Pt(9.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"{EXAM_TITLE} — MODEL ANSWERS — Set {set_num}")
    run.bold = True; run.font.size = Pt(10.5)

    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info = [
        ("Course:", COURSE, "Date:", EXAM_DATE),
        ("Duration:", DURATION, "Total Marks:", TOTAL_MARKS),
    ]
    for i, row_data in enumerate(info):
        for j, text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ""
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            cp.paragraph_format.space_before = Pt(0)
            run = cp.add_run(text)
            run.font.size = Pt(8.5)
            if j % 2 == 0:
                run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("―" * 60)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _add_q_block(doc, label, marks, bloom, co, stem, subs, solutions=None):
    """Add one question block with optionally interleaved answers."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label} {marks}")
    run.bold = True; run.font.size = Pt(9.5)
    run2 = p.add_run(f"    Bloom: {bloom}  |  {co}")
    run2.italic = True; run2.font.size = Pt(7.5)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if stem:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        run3 = p3.add_run(stem)
        run3.font.size = Pt(9)

    for idx, sub in enumerate(subs):
        letter = chr(ord("a") + idx)
        p4 = doc.add_paragraph()
        p4.paragraph_format.left_indent = Cm(0.7)
        p4.paragraph_format.space_after = Pt(0)
        run4 = p4.add_run(f"({letter})  {sub}")
        run4.font.size = Pt(9)
        mark_run = p4.add_run("  [0.5]")
        mark_run.font.size = Pt(8); mark_run.bold = True

        if solutions and idx < len(solutions):
            p5 = doc.add_paragraph()
            p5.paragraph_format.left_indent = Cm(1.2)
            p5.paragraph_format.space_after = Pt(0)
            run5 = p5.add_run("Ans: ")
            run5.bold = True; run5.font.size = Pt(9)
            run5.font.color.rgb = RGBColor(0x00, 0x40, 0x80)
            for line in solutions[idx].split("\n"):
                run_sol = p5.add_run(line)
                run_sol.font.size = Pt(9)
                run_sol.font.color.rgb = RGBColor(0x00, 0x40, 0x80)
                if line != solutions[idx].split("\n")[-1]:
                    p5 = doc.add_paragraph()
                    p5.paragraph_format.left_indent = Cm(1.2)
                    p5.paragraph_format.space_after = Pt(0)


def _build_solutions_docx(set_num):
    """Build solution .docx for given set."""
    doc = Document()
    _add_header_sol(doc, set_num)

    q1a = Q1A_BANKS[set_num]
    q1b = Q1B_BANKS[set_num]
    q2  = Q2_BANKS[set_num]
    q3  = Q3_BANKS[set_num]
    q4  = Q4_BANKS[set_num]

    q2_sols = solve_q2(set_num)
    q3_sols = solve_q3(set_num)
    q4_sols = solve_q4(set_num)

    _add_q_block(doc, "Q1(a)", "[2 Marks]",
                 "BL 1–2 (Remember / Understand)", "CO1",
                 "Answer the following:", q1a["questions"], q1a["solutions"])

    _add_q_block(doc, "Q1(b)", "[2 Marks]",
                 "BL 2–3 (Understand / Apply)", "CO1, CO2",
                 None, q1b["questions"], q1b["solutions"])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Solve any TWO from Q2, Q3, Q4  [3 Marks each = 6 Marks]")
    run.bold = True; run.font.size = Pt(9.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_q_block(doc, "Q2.", "[3 Marks]", q2["bloom"], q2["co"],
                 q2["stem"], q2["subs"], q2_sols)
    _add_q_block(doc, "Q3.", "[3 Marks]", q3["bloom"], q3["co"],
                 q3["stem"], q3["subs"], q3_sols)
    _add_q_block(doc, "Q4.", "[3 Marks]", q4["bloom"], q4["co"],
                 q4["stem"], q4["subs"], q4_sols)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(f"— End of Set {set_num} Model Answers —")
    run.bold = True; run.font.size = Pt(9)

    return doc


# ═══════════════════════ PDF BUILDERS ════════════════════════

def _safe(text):
    """Escape XML special chars for ReportLab Paragraph."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _register_fonts():
    """Register TNR TTF for Unicode support (idempotent)."""
    FONT_DIR = r"C:\Windows\Fonts"
    try:
        pdfmetrics.getFont("TNR")
    except KeyError:
        pdfmetrics.registerFont(TTFont("TNR",   os.path.join(FONT_DIR, "times.ttf")))
        pdfmetrics.registerFont(TTFont("TNRB",  os.path.join(FONT_DIR, "timesbd.ttf")))
        pdfmetrics.registerFont(TTFont("TNRI",  os.path.join(FONT_DIR, "timesi.ttf")))
        pdfmetrics.registerFont(TTFont("TNRBI", os.path.join(FONT_DIR, "timesbi.ttf")))
        pdfmetrics.registerFontFamily("TNR", normal="TNR", bold="TNRB",
                                      italic="TNRI", boldItalic="TNRBI")


def _get_qp_styles():
    """Return dict of ReportLab styles for QP PDF."""
    return {
        "inst_head": ParagraphStyle("InstH", fontSize=11, fontName="TNRB",
                                    alignment=TA_CENTER, spaceAfter=0, spaceBefore=0, leading=13),
        "inst_sub":  ParagraphStyle("InstS", fontSize=10, fontName="TNRB",
                                    alignment=TA_CENTER, spaceAfter=0, leading=12),
        "inst_year": ParagraphStyle("InstY", fontSize=10, fontName="TNR",
                                    alignment=TA_CENTER, spaceAfter=0, leading=12),
        "hdr":       ParagraphStyle("Hdr", fontSize=9.5, fontName="TNR",
                                    leading=12, spaceAfter=0, spaceBefore=0),
        "test":      ParagraphStyle("Test", fontSize=13, fontName="TNRB",
                                    alignment=TA_CENTER, spaceAfter=1, spaceBefore=1, leading=15),
        "instr_hd":  ParagraphStyle("InstrHd", fontSize=9.5, fontName="TNRB",
                                    spaceAfter=0, spaceBefore=2*mm, leading=12),
        "instr":     ParagraphStyle("Instr", fontSize=9.5, fontName="TNR",
                                    leading=12, spaceAfter=0, spaceBefore=0),
        "q_label":   ParagraphStyle("Qlbl", fontSize=10, fontName="TNRB",
                                    spaceAfter=0, spaceBefore=0, leading=12),
        "co_so":     ParagraphStyle("CoSo", fontSize=7.5, fontName="TNR",
                                    leading=9, spaceAfter=0, spaceBefore=0),
        "q_text":    ParagraphStyle("Qtxt", fontSize=10, fontName="TNR",
                                    leading=12, spaceAfter=0, spaceBefore=0),
        "marks":     ParagraphStyle("Mrks", fontSize=10, fontName="TNR",
                                    leading=12, spaceAfter=0, spaceBefore=0),
        "q_brief":   ParagraphStyle("Qbrief", fontSize=10, fontName="TNR",
                                    leading=12, spaceAfter=0, spaceBefore=0),
    }


def _get_sol_styles():
    """Return dict of ReportLab styles for solution PDF."""
    return {
        "title":    ParagraphStyle("Title", fontSize=14, fontName="TNRB",
                                   alignment=TA_CENTER, spaceAfter=2*mm, spaceBefore=0, leading=17),
        "subtitle": ParagraphStyle("SubT", fontSize=11, fontName="TNR",
                                   alignment=TA_CENTER, spaceAfter=1*mm, leading=13),
        "qlabel":   ParagraphStyle("SQlbl", fontSize=10.5, fontName="TNRB",
                                   spaceAfter=0.5*mm, spaceBefore=2*mm, leading=13),
        "bloom":    ParagraphStyle("Bloom", fontSize=8, fontName="TNRI",
                                   textColor=colors.HexColor("#555555"),
                                   spaceAfter=0.5*mm, leading=10),
        "stem":     ParagraphStyle("Stem", fontSize=10, fontName="TNR",
                                   spaceAfter=0.5*mm, leading=12.5),
        "sub_q":    ParagraphStyle("SubQ", fontSize=10, fontName="TNR",
                                   leftIndent=6*mm, spaceAfter=0, leading=12.5),
        "ans":      ParagraphStyle("Ans", fontSize=10, fontName="TNR",
                                   leftIndent=10*mm, spaceAfter=0.5*mm, leading=12.5,
                                   textColor=colors.HexColor("#004080")),
        "ans_lbl":  ParagraphStyle("AnsL", fontSize=10, fontName="TNRB",
                                   leftIndent=10*mm, spaceAfter=0, leading=12.5,
                                   textColor=colors.HexColor("#004080")),
        "footer":   ParagraphStyle("Foot", fontSize=10, fontName="TNRB",
                                   alignment=TA_CENTER, spaceBefore=3*mm, leading=13),
        "sep":      ParagraphStyle("Sep", fontSize=10, fontName="TNRB",
                                   alignment=TA_CENTER, leading=13),
    }


def _build_qp_page(elements, qp_docx_path, sty, page_w):
    """Add one QP page to elements list from a .docx file."""
    docx = Document(qp_docx_path)

    ht = docx.tables[0]
    r0_text = ht.cell(0, 0).text.strip()
    lines = [l.strip() for l in r0_text.split('\n') if l.strip()]
    if len(lines) >= 1:
        elements.append(Paragraph(_safe(lines[0]), sty["inst_head"]))
    if len(lines) >= 2:
        elements.append(Paragraph(_safe(lines[1]), sty["inst_sub"]))
    if len(lines) >= 3:
        elements.append(Paragraph(_safe(lines[2]), sty["inst_year"]))

    info_data = []
    for ri in range(1, 5):
        c0 = ht.cell(ri, 0).text.strip()
        c1 = ht.cell(ri, 1).text.strip()
        info_data.append([
            Paragraph(_safe(c0), sty["hdr"]),
            Paragraph(_safe(c1), sty["hdr"]),
        ])
    info_table = Table(info_data, colWidths=[page_w * 0.52, page_w * 0.48])
    info_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 2),
        ("RIGHTPADDING", (0, 0), (-1, -1), 2),
        ("TOPPADDING", (0, 0), (-1, -1), 0.5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0.5),
    ]))
    elements.append(info_table)

    r5_text = ht.cell(5, 0).text.strip()
    elements.append(Paragraph(_safe(r5_text), sty["test"]))

    for elem in docx.element.body:
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag == 'p':
            from docx.text.paragraph import Paragraph as DocxPara
            p = DocxPara(elem, docx)
            text = p.text.strip()
            if not text:
                continue
            if "Instructions" in text:
                elements.append(Paragraph(_safe(text), sty["instr_hd"]))
            elif text[0].isdigit() and ")" in text[:3]:
                elements.append(Paragraph(_safe(text), sty["instr"]))

    qt = docx.tables[1]
    flat_data = []
    for ri in range(len(qt.rows)):
        c0_text = qt.cell(ri, 0).text.strip()
        c1_text = qt.cell(ri, 1).text.strip()
        c2_text = qt.cell(ri, 2).text.strip()
        c3_text = qt.cell(ri, 3).text.strip()

        c0_lines = [l.strip() for l in c0_text.split('\n') if l.strip()]
        c0_parts = []
        for line in c0_lines:
            if line.startswith("Q"):
                c0_parts.append(Paragraph(f"<b>{_safe(line)}</b>", sty["q_label"]))
            else:
                c0_parts.append(Paragraph(_safe(line), sty["co_so"]))
        c0_content = c0_parts if c0_parts else [Paragraph("", sty["q_label"])]

        c1_content = Paragraph(_safe(c1_text), sty["q_text"]) if c1_text else Paragraph("", sty["q_text"])

        c2_lines = [l.strip() for l in c2_text.split('\n') if l.strip()]
        if ri == 0 and c2_text:
            c2_content = [Paragraph(_safe(c2_text), sty["q_brief"])]
        else:
            c2_content = [Paragraph(_safe(l), sty["q_text"]) for l in c2_lines] if c2_lines else [Paragraph("", sty["q_text"])]

        c3_content = Paragraph(_safe(c3_text), sty["marks"]) if c3_text else Paragraph("", sty["marks"])
        flat_data.append([c0_content, c1_content, c2_content, c3_content])

    cw = [page_w * 0.09, page_w * 0.04, page_w * 0.80, page_w * 0.07]
    q_table = Table(flat_data, colWidths=cw)
    q_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 1),
        ("RIGHTPADDING", (0, 0), (-1, -1), 1),
        ("TOPPADDING", (0, 0), (-1, -1), 1),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
    ]))
    elements.append(Spacer(1, 1*mm))
    elements.append(q_table)


def _build_sol_page(elements, set_num, sty):
    """Add one solution set to elements list."""
    q1a = Q1A_BANKS[set_num]
    q1b = Q1B_BANKS[set_num]
    q2  = Q2_BANKS[set_num]
    q3  = Q3_BANKS[set_num]
    q4  = Q4_BANKS[set_num]
    q2_sols = solve_q2(set_num)
    q3_sols = solve_q3(set_num)
    q4_sols = solve_q4(set_num)

    elements.append(Paragraph(
        "SVKM'S NMIMS — MUKESH PATEL SCHOOL OF TECHNOLOGY MANAGEMENT &amp; ENGINEERING",
        sty["subtitle"]))
    elements.append(Paragraph(
        f"Re-Examination  |  Dynamic Systems Modeling  |  Set {set_num}",
        sty["subtitle"]))
    elements.append(Paragraph("MODEL ANSWERS", sty["title"]))

    def add_qa_block(label, marks, bloom, co, stem, subs, sols):
        elements.append(Paragraph(f"{_safe(label)}  {_safe(marks)}", sty["qlabel"]))
        elements.append(Paragraph(f"Bloom: {_safe(bloom)}  |  {_safe(co)}", sty["bloom"]))
        if stem:
            elements.append(Paragraph(_safe(stem), sty["stem"]))
        for i, sub in enumerate(subs):
            letter = chr(ord("a") + i)
            elements.append(Paragraph(
                f"({letter}) {_safe(sub)}  <b>[0.5]</b>", sty["sub_q"]))
            if sols and i < len(sols):
                elements.append(Paragraph("<b>Ans:</b>", sty["ans_lbl"]))
                for line in sols[i].split("\n"):
                    if line.strip():
                        elements.append(Paragraph(_safe(line), sty["ans"]))

    add_qa_block("Q1(a)", "[2 Marks]",
                 "BL 1–2 (Remember / Understand)", "CO1",
                 "Answer the following:", q1a["questions"], q1a["solutions"])
    add_qa_block("Q1(b)", "[2 Marks]",
                 "BL 2–3 (Understand / Apply)", "CO1, CO2",
                 None, q1b["questions"], q1b["solutions"])

    elements.append(Spacer(1, 2*mm))
    elements.append(Paragraph(
        "<b>Solve any TWO from Q2, Q3, Q4  [3 Marks each = 6 Marks]</b>",
        sty["sep"]))

    add_qa_block("Q2", "[3 Marks]", q2["bloom"], q2["co"],
                 q2["stem"], q2["subs"], q2_sols)
    add_qa_block("Q3", "[3 Marks]", q3["bloom"], q3["co"],
                 q3["stem"], q3["subs"], q3_sols)
    add_qa_block("Q4", "[3 Marks]", q4["bloom"], q4["co"],
                 q4["stem"], q4["subs"], q4_sols)

    elements.append(Paragraph(f"— End of Set {set_num} Model Answers —", sty["footer"]))


def build_qp_pdf(qp_docx_path, pdf_path):
    """Build QP PDF for a single set."""
    _register_fonts()
    sty = _get_qp_styles()
    doc = SimpleDocTemplate(
        pdf_path, pagesize=A4,
        topMargin=15*mm, bottomMargin=12*mm,
        leftMargin=18*mm, rightMargin=18*mm,
    )
    elements = []
    page_w = A4[0] - 36*mm
    _build_qp_page(elements, qp_docx_path, sty, page_w)
    doc.build(elements)


def build_sol_pdf(set_num, pdf_path):
    """Build Solutions PDF for a single set."""
    _register_fonts()
    sty = _get_sol_styles()
    doc = SimpleDocTemplate(
        pdf_path, pagesize=A4,
        topMargin=15*mm, bottomMargin=12*mm,
        leftMargin=18*mm, rightMargin=18*mm,
    )
    elements = []
    _build_sol_page(elements, set_num, sty)
    doc.build(elements)


# ═══════════════════════ MAIN ════════════════════════════════

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print("=" * 65)
    print("DYNAMIC SYSTEMS MODELING — M3 RE-EXAM QUESTION PAPER GENERATOR")
    print("=" * 65)
    print(f"  Exam     : {EXAM_TITLE}")
    print(f"  Scope    : Unit 1 – Unit 5 (complete)")
    print(f"  Sets     : {NUM_SETS}")
    print(f"  Format   : Word (.docx) + Per-set QP PDF + Per-set Solutions PDF")
    print()

    qp_paths = []
    # Generate .docx files
    for s in range(1, NUM_SETS + 1):
        qp_doc = _build_qp_docx(s)
        qp_path = os.path.join(OUTPUT_DIR, f"DSM_M3_Re_Exam_Set_{s:02d}.docx")
        qp_doc.save(qp_path)
        qp_paths.append(qp_path)
        print(f"  ✓ QP Set {s}:        {os.path.basename(qp_path)}")

        sol_doc = _build_solutions_docx(s)
        sol_path = os.path.join(OUTPUT_DIR, f"DSM_M3_Re_Exam_Solutions_Set_{s:02d}.docx")
        sol_doc.save(sol_path)
        print(f"  ✓ Solutions Set {s}: {os.path.basename(sol_path)}")

    print()

    # Per-set PDFs
    try:
        from PyPDF2 import PdfReader
        has_pypdf = True
    except ImportError:
        has_pypdf = False

    for s in range(1, NUM_SETS + 1):
        qp_pdf = os.path.join(OUTPUT_DIR, f"DSM_M3_Re_Exam_QP_Set_{s:02d}.pdf")
        build_qp_pdf(qp_paths[s - 1], qp_pdf)
        if has_pypdf:
            pages = len(PdfReader(qp_pdf).pages)
            print(f"  ✓ QP PDF Set {s}:       {os.path.basename(qp_pdf)} ({pages} pages)")
        else:
            print(f"  ✓ QP PDF Set {s}:       {os.path.basename(qp_pdf)}")

        sol_pdf = os.path.join(OUTPUT_DIR, f"DSM_M3_Re_Exam_Solutions_Set_{s:02d}.pdf")
        build_sol_pdf(s, sol_pdf)
        if has_pypdf:
            pages = len(PdfReader(sol_pdf).pages)
            print(f"  ✓ Sol PDF Set {s}:       {os.path.basename(sol_pdf)} ({pages} pages)")
        else:
            print(f"  ✓ Sol PDF Set {s}:       {os.path.basename(sol_pdf)}")

    print()
    print(f"  Output dir: {OUTPUT_DIR}")
    print("=" * 65)


if __name__ == "__main__":
    main()
